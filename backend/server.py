from dotenv import load_dotenv
from pathlib import Path

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

import asyncio
import os
import pickle
import re
import signal
import zlib
import hmac
import hashlib
import html as _html
import io
import json as _json
import secrets
import uuid
import logging
import time
import resource
import bcrypt
import jwt
import unicodedata
from collections import OrderedDict, defaultdict, deque
from datetime import datetime, timezone, timedelta
from typing import Any, Deque, Dict, List, Optional, Tuple
from urllib.parse import unquote, urlencode, urlparse
from urllib.request import Request as UrlRequest, urlopen

from fastapi import FastAPI, APIRouter, HTTPException, Depends, Request, Response, Cookie, UploadFile, File
from fastapi.exceptions import RequestValidationError
from fastapi.responses import JSONResponse, RedirectResponse, StreamingResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from starlette.middleware.cors import CORSMiddleware
from starlette.middleware.gzip import GZipMiddleware
from contextlib import asynccontextmanager
from motor.motor_asyncio import AsyncIOMotorClient, AsyncIOMotorGridFSBucket
from pydantic import BaseModel, Field, EmailStr, ConfigDict
from pymongo.errors import AutoReconnect, ServerSelectionTimeoutError

try:
    from rights_engine import RIGHTS_REPORT_FILENAMES, rights_publish_blockers, rights_report_csv, rights_report_rows
except ImportError:  # pragma: no cover - supports package-style test imports
    from backend.rights_engine import RIGHTS_REPORT_FILENAMES, rights_publish_blockers, rights_report_csv, rights_report_rows

try:
    from catalog_truth import (
        CONTROLLED_LIVE_BOOK_SLUGS as CATALOG_TRUTH_LIVE_BOOK_SLUGS,
        AUDIO_ENABLED_SLUGS as CATALOG_TRUTH_AUDIO_ENABLED_SLUGS,
        PIPELINE_CANDIDATE_SLUGS as CATALOG_TRUTH_PIPELINE_SLUGS,
        audio_release_qa_status,
        can_expose_audio,
        can_expose_reader,
        controlled_artifact_status,
        dracula_artifact_status,
        explicit_preview_chapter_ids,
        live_approved_mongo_query,
        load_controlled_artifact_book,
        load_dracula_artifact_book,
        public_book_projection,
    )
except ImportError:  # pragma: no cover - supports package-style test imports
    from backend.catalog_truth import (
        CONTROLLED_LIVE_BOOK_SLUGS as CATALOG_TRUTH_LIVE_BOOK_SLUGS,
        AUDIO_ENABLED_SLUGS as CATALOG_TRUTH_AUDIO_ENABLED_SLUGS,
        PIPELINE_CANDIDATE_SLUGS as CATALOG_TRUTH_PIPELINE_SLUGS,
        audio_release_qa_status,
        can_expose_audio,
        can_expose_reader,
        controlled_artifact_status,
        dracula_artifact_status,
        explicit_preview_chapter_ids,
        live_approved_mongo_query,
        load_controlled_artifact_book,
        load_dracula_artifact_book,
        public_book_projection,
    )

try:
    from home_curation import build_home_curated_payload
except ImportError:  # pragma: no cover - supports package-style test imports
    from backend.home_curation import build_home_curated_payload


# ---------- Environment / DB ----------
ENVIRONMENT = os.environ.get("ENVIRONMENT", "production").strip().lower()
LOG_LEVEL = os.environ.get("LOG_LEVEL", "INFO").strip().upper()
logging.basicConfig(level=getattr(logging, LOG_LEVEL, logging.INFO), format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)
_PROCESS_STARTED_AT = time.monotonic()


class UTF8JSONResponse(JSONResponse):
    media_type = "application/json; charset=utf-8"

    def render(self, content) -> bytes:
        return _json.dumps(
            content,
            ensure_ascii=False,
            allow_nan=False,
            separators=(",", ":"),
        ).encode("utf-8")


def _env_bool(name: str, default: bool) -> bool:
    value = os.environ.get(name)
    if value is None:
        return default
    return value.strip().lower() not in {"0", "false", "no", "off"}


def _env_int(name: str, default: int, minimum: int = 1) -> int:
    try:
        return max(minimum, int(os.environ.get(name, default)))
    except (TypeError, ValueError):
        return default


def _env_float(name: str, default: float, minimum: float = 0.1) -> float:
    try:
        return max(minimum, float(os.environ.get(name, default)))
    except (TypeError, ValueError):
        return default


# Railway cost-control defaults are fail-closed: the web service serves API
# traffic only unless an operator explicitly enables expensive jobs.
COST_CONTROL_MODE = _env_bool("COST_CONTROL_MODE", ENVIRONMENT == "production")
ENABLE_BACKGROUND_WORKERS = _env_bool("ENABLE_BACKGROUND_WORKERS", False)
ENABLE_AUDIOBOOK_PIPELINE = _env_bool("ENABLE_AUDIOBOOK_PIPELINE", False)
ENABLE_BOOK_RENDERING_JOBS = _env_bool("ENABLE_BOOK_RENDERING_JOBS", False)
ENABLE_COVER_GENERATION = _env_bool("ENABLE_COVER_GENERATION", False)
ENABLE_SCHEDULED_JOBS = _env_bool("ENABLE_SCHEDULED_JOBS", False)
ENABLE_QUEUE_CONSUMER = _env_bool("ENABLE_QUEUE_CONSUMER", False)
ENABLE_ADMIN_MEDIA_UPLOADS = _env_bool("ENABLE_ADMIN_MEDIA_UPLOADS", False)
ENABLE_STARTUP_DB_MAINTENANCE = _env_bool("ENABLE_STARTUP_DB_MAINTENANCE", not COST_CONTROL_MODE)
MAX_CONCURRENT_JOBS = _env_int("MAX_CONCURRENT_JOBS", 1)
REQUEST_BODY_LIMIT_BYTES = _env_int(
    "REQUEST_BODY_LIMIT_BYTES",
    2 * 1024 * 1024 if COST_CONTROL_MODE else 8 * 1024 * 1024,
)
DOCX_UPLOAD_MAX_BYTES = _env_int(
    "DOCX_UPLOAD_MAX_BYTES",
    8 * 1024 * 1024 if COST_CONTROL_MODE else 50 * 1024 * 1024,
)
CHAPTER_UPLOAD_MAX_BYTES = _env_int(
    "CHAPTER_UPLOAD_MAX_BYTES",
    8 * 1024 * 1024 if COST_CONTROL_MODE else 50 * 1024 * 1024,
)
ADMIN_MEDIA_UPLOAD_MAX_BYTES = _env_int(
    "ADMIN_MEDIA_UPLOAD_MAX_BYTES",
    4 * 1024 * 1024 if COST_CONTROL_MODE else 10 * 1024 * 1024,
)
_expensive_job_state: Dict[str, Any] = {"active": 0, "started": defaultdict(int), "blocked": defaultdict(int)}
_expensive_job_lock: Optional[asyncio.Lock] = None


mongo_url = os.environ.get("MONGODB_URL") or os.environ.get("MONGO_URL")
if not mongo_url:
    raise RuntimeError("MONGODB_URL is required")

def _database_name_from_mongo_url(url: str) -> str:
    parsed = urlparse(url)
    db_name = parsed.path.lstrip("/").split("/", 1)[0]
    return db_name or "earnalism"

MONGODB_MAX_POOL_SIZE = _env_int("MONGODB_MAX_POOL_SIZE", 8 if COST_CONTROL_MODE else 25)
MONGODB_MIN_POOL_SIZE = _env_int("MONGODB_MIN_POOL_SIZE", 1, minimum=0)
MONGODB_MAX_CONNECTING = _env_int("MONGODB_MAX_CONNECTING", 2)
MONGODB_SERVER_SELECTION_TIMEOUT_MS = _env_int("MONGODB_SERVER_SELECTION_TIMEOUT_MS", 15000)
MONGODB_WAIT_QUEUE_TIMEOUT_MS = _env_int("MONGODB_WAIT_QUEUE_TIMEOUT_MS", 5000)
client = AsyncIOMotorClient(
    mongo_url,
    maxPoolSize=MONGODB_MAX_POOL_SIZE,
    minPoolSize=min(MONGODB_MIN_POOL_SIZE, MONGODB_MAX_POOL_SIZE),
    maxConnecting=MONGODB_MAX_CONNECTING,
    serverSelectionTimeoutMS=MONGODB_SERVER_SELECTION_TIMEOUT_MS,
    waitQueueTimeoutMS=MONGODB_WAIT_QUEUE_TIMEOUT_MS,
    uuidRepresentation="standard",
)
db = client[os.environ.get("DB_NAME") or _database_name_from_mongo_url(mongo_url)]
admin_upload_files = AsyncIOMotorGridFSBucket(db, bucket_name="admin_upload_files")

JWT_SECRET = os.environ.get("JWT_SECRET")
if not JWT_SECRET:
    raise RuntimeError("JWT_SECRET is required")
JWT_ALG = "HS256"
JWT_EXPIRE_MINUTES = _env_int("JWT_EXPIRE_MINUTES", 1440)
USER_ACCESS_TOKEN_EXPIRE_MINUTES = _env_int("USER_ACCESS_TOKEN_EXPIRE_MINUTES", 30)
USER_REFRESH_IDLE_MINUTES = _env_int("USER_REFRESH_IDLE_MINUTES", 30)
USER_REFRESH_TOTAL_HOURS = _env_int("USER_REFRESH_TOTAL_HOURS", 12)
TRUSTED_DEVICE_MAX_ACTIVE_SESSIONS = _env_int("TRUSTED_DEVICE_MAX_ACTIVE_SESSIONS", 1)
READING_HEARTBEAT_EARLY_GRACE_SECONDS = _env_int("READING_HEARTBEAT_EARLY_GRACE_SECONDS", 5)
READING_SESSION_IDLE_GRACE_SECONDS = _env_int("READING_SESSION_IDLE_GRACE_SECONDS", 120)
SESSION_TOUCH_INTERVAL_SECONDS = _env_int("SESSION_TOUCH_INTERVAL_SECONDS", 60)
if SESSION_TOUCH_INTERVAL_SECONDS > 0:
    SESSION_TOUCH_INTERVAL_SECONDS = min(
        SESSION_TOUCH_INTERVAL_SECONDS,
        max(1, (USER_REFRESH_IDLE_MINUTES * 60) // 2),
    )
ADMIN_EMAIL = os.environ.get("ADMIN_EMAIL", "admin@theearnalism.com").strip()
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "").strip()
SEED_TEST_READER = os.environ.get("SEED_TEST_READER", "false").strip().lower() == "true"
SEED_TEST_READER_EMAIL = os.environ.get("SEED_TEST_READER_EMAIL", "reader@earnalism.com").strip()
SEED_TEST_READER_PASSWORD = os.environ.get("SEED_TEST_READER_PASSWORD", "").strip()

# Cookie config — httpOnly session cookie. SECURE flag is on by default (HTTPS in prod);
# can be disabled via env for plain-HTTP local dev only.
SESSION_COOKIE = "ear_session"
USER_REFRESH_COOKIE = "ear_user_refresh"
SESSION_TTL_SECONDS = JWT_EXPIRE_MINUTES * 60
COOKIE_SECURE = os.environ.get("COOKIE_SECURE", "true").lower() != "false"
COOKIE_SAMESITE = os.environ.get("COOKIE_SAMESITE", "lax")

# Trial-safe horizontal-scaling switch. Keep false on Railway Trial. When Pro is
# active, set MULTI_REPLICA_ENABLED=true with REDIS_URL so shared request state
# moves out of each process.
MULTI_REPLICA_ENABLED = _env_bool("MULTI_REPLICA_ENABLED", False)
REDIS_URL = os.environ.get("REDIS_URL", "").strip()
REDIS_KEY_PREFIX = os.environ.get("REDIS_KEY_PREFIX", "earnalism").strip() or "earnalism"
REDIS_CACHE_ENABLED = _env_bool("REDIS_CACHE_ENABLED", bool(REDIS_URL) or MULTI_REPLICA_ENABLED)
REDIS_CACHE_FAIL_FAST = _env_bool("REDIS_CACHE_FAIL_FAST", MULTI_REPLICA_ENABLED)
REDIS_SOCKET_CONNECT_TIMEOUT_SECONDS = _env_float("REDIS_SOCKET_CONNECT_TIMEOUT_SECONDS", 2.0)
REDIS_SOCKET_TIMEOUT_SECONDS = _env_float("REDIS_SOCKET_TIMEOUT_SECONDS", 2.0)
REDIS_CACHE_COMPRESS_MIN_BYTES = _env_int("REDIS_CACHE_COMPRESS_MIN_BYTES", 4096)
REDIS_CACHE_TTL_JITTER_SECONDS = _env_int("REDIS_CACHE_TTL_JITTER_SECONDS", 30, minimum=0)
REDIS_CONFIGURE_ON_STARTUP = _env_bool("REDIS_CONFIGURE_ON_STARTUP", False)
REDIS_MAXMEMORY = os.environ.get("REDIS_MAXMEMORY", "").strip()
REDIS_MAXMEMORY_POLICY = os.environ.get("REDIS_MAXMEMORY_POLICY", "volatile-lfu").strip()
USER_AUTH_CACHE_TTL_SECONDS = _env_int("USER_AUTH_CACHE_TTL_SECONDS", 20)
USER_SESSION_CACHE_TTL_SECONDS = _env_int("USER_SESSION_CACHE_TTL_SECONDS", 20)
USER_WALLET_CACHE_TTL_SECONDS = _env_int("USER_WALLET_CACHE_TTL_SECONDS", 8)
USER_TRANSACTIONS_CACHE_TTL_SECONDS = _env_int("USER_TRANSACTIONS_CACHE_TTL_SECONDS", 20)
USER_PAYMENT_INTENTS_CACHE_TTL_SECONDS = _env_int("USER_PAYMENT_INTENTS_CACHE_TTL_SECONDS", 15)
READER_MANIFEST_CACHE_TTL_SECONDS = _env_int("READER_MANIFEST_CACHE_TTL_SECONDS", 1800)
READER_BOOK_CACHE_TTL_SECONDS = _env_int("READER_BOOK_CACHE_TTL_SECONDS", 900)
READER_CHAPTER_CACHE_TTL_SECONDS = _env_int("READER_CHAPTER_CACHE_TTL_SECONDS", 3600)
READER_RUM_ENABLED = _env_bool("READER_RUM_ENABLED", True)
READER_RUM_SAMPLE_RATE = min(1.0, _env_float("READER_RUM_SAMPLE_RATE", 0.15, minimum=0.0))
READER_RUM_SLOW_MS = _env_int("READER_RUM_SLOW_MS", 1800)
READER_RUM_AGGREGATE_TTL_SECONDS = _env_int("READER_RUM_AGGREGATE_TTL_SECONDS", 7 * 24 * 3600)
STARTUP_MAINTENANCE_VERSION = os.environ.get("STARTUP_MAINTENANCE_VERSION", "2026-06-03-autoscale-v2").strip()
STARTUP_MAINTENANCE_LOCK_SECONDS = _env_int("STARTUP_MAINTENANCE_LOCK_SECONDS", 180)
STARTUP_MAINTENANCE_WAIT_SECONDS = _env_int("STARTUP_MAINTENANCE_WAIT_SECONDS", 45)
STARTUP_MAINTENANCE_DONE_TTL_SECONDS = _env_int("STARTUP_MAINTENANCE_DONE_TTL_SECONDS", 600)
STARTUP_DB_MAINTENANCE_ATTEMPTS = _env_int("STARTUP_DB_MAINTENANCE_ATTEMPTS", 5)
_redis_client: Any = None
_redis_available = False
_cache_stats: Dict[str, int] = defaultdict(int)
_redis_config_status: Dict[str, Any] = {}


# ---------- Razorpay test-mode config ----------
RAZORPAY_KEY_ID = os.environ.get("RAZORPAY_KEY_ID", "").strip()
RAZORPAY_KEY_SECRET = os.environ.get("RAZORPAY_KEY_SECRET", "").strip()
RAZORPAY_WEBHOOK_SECRET = os.environ.get("RAZORPAY_WEBHOOK_SECRET", "").strip()
RAZORPAY_MODE = os.environ.get("RAZORPAY_MODE", "test").strip().lower()
TOPUP_INTENT_TTL_SECONDS = _env_int("TOPUP_INTENT_TTL_SECONDS", 24 * 60 * 60, minimum=60)

B2_S3_ENDPOINT = (os.environ.get("B2_S3_ENDPOINT") or os.environ.get("B2_ENDPOINT") or "").strip().rstrip("/")
B2_REGION = os.environ.get("B2_REGION", "").strip()
B2_BUCKET = (os.environ.get("B2_BUCKET") or os.environ.get("B2_BUCKET_NAME") or "").strip()
B2_ACCESS_KEY_ID = (os.environ.get("B2_ACCESS_KEY_ID") or os.environ.get("B2_KEY_ID") or "").strip()
B2_SECRET_ACCESS_KEY = (os.environ.get("B2_SECRET_ACCESS_KEY") or os.environ.get("B2_APP_KEY") or "").strip()

# Dependency-free per-process rate limits. For multi-instance scaling, move this
# counter to Redis or an edge/WAF layer so limits are shared across replicas.
RATE_LIMIT_ENABLED = _env_bool("RATE_LIMIT_ENABLED", True)
RATE_LIMIT_WINDOW_SECONDS = _env_int("RATE_LIMIT_WINDOW_SECONDS", 60)
RATE_LIMIT_DEFAULT_PER_MINUTE = _env_int("RATE_LIMIT_DEFAULT_PER_MINUTE", 1200)
RATE_LIMIT_PUBLIC_PER_MINUTE = _env_int("RATE_LIMIT_PUBLIC_PER_MINUTE", 30000)
RATE_LIMIT_READER_PER_MINUTE = _env_int("RATE_LIMIT_READER_PER_MINUTE", 15000)
RATE_LIMIT_AUTH_PER_MINUTE = _env_int("RATE_LIMIT_AUTH_PER_MINUTE", 120)
RATE_LIMIT_PAYMENT_PER_MINUTE = _env_int("RATE_LIMIT_PAYMENT_PER_MINUTE", 300)
RATE_LIMIT_WEBHOOK_PER_MINUTE = _env_int("RATE_LIMIT_WEBHOOK_PER_MINUTE", 600)
RATE_LIMIT_ANALYTICS_PER_MINUTE = _env_int("RATE_LIMIT_ANALYTICS_PER_MINUTE", 1800)
RATE_LIMIT_UPLOAD_PER_MINUTE = _env_int("RATE_LIMIT_UPLOAD_PER_MINUTE", 60)
_rate_limit_hits: Dict[str, Deque[float]] = defaultdict(deque)
_rate_limit_next_sweep = 0.0

# Small per-process cache for anonymous public catalogue reads. It intentionally
# excludes authenticated/user/admin routes. In multi-replica mode it is backed by
# Redis so public cache reads and invalidation are shared across Railway replicas.
PUBLIC_CACHE_ENABLED = _env_bool("PUBLIC_CACHE_ENABLED", True)
PUBLIC_CACHE_TTL_SECONDS = _env_int("PUBLIC_CACHE_TTL_SECONDS", 300)
PUBLIC_CACHE_MAX_ENTRIES = _env_int("PUBLIC_CACHE_MAX_ENTRIES", 96 if COST_CONTROL_MODE else 256)
HOME_BOOK_LIMIT = _env_int("HOME_BOOK_LIMIT", 120 if COST_CONTROL_MODE else 500)
HOME_BOOK_PAGE_DEFAULT_LIMIT = _env_int("HOME_BOOK_PAGE_DEFAULT_LIMIT", 8)
HOME_BOOK_PAGE_MAX_LIMIT = _env_int("HOME_BOOK_PAGE_MAX_LIMIT", 24)
HEALTH_CACHE_TTL_SECONDS = _env_int("HEALTH_CACHE_TTL_SECONDS", 5)
_public_cache: "OrderedDict[str, Tuple[float, Any]]" = OrderedDict()
_health_cache: dict = {"expires_at": 0.0, "payload": None}
_public_cache_generation = 0
_shutdown_state: dict = {"draining": False, "inflight": 0}


def _cloudinary_config_detected() -> bool:
    return bool(
        os.environ.get("CLOUDINARY_CLOUD_NAME")
        and os.environ.get("CLOUDINARY_API_KEY")
        and os.environ.get("CLOUDINARY_API_SECRET")
    )


def _cost_control_flags() -> Dict[str, Any]:
    return {
        "cost_control_mode": COST_CONTROL_MODE,
        "background_workers_enabled": ENABLE_BACKGROUND_WORKERS,
        "audiobook_pipeline_enabled": ENABLE_AUDIOBOOK_PIPELINE,
        "book_rendering_jobs_enabled": ENABLE_BOOK_RENDERING_JOBS,
        "cover_generation_enabled": ENABLE_COVER_GENERATION,
        "scheduled_jobs_enabled": ENABLE_SCHEDULED_JOBS,
        "queue_consumer_enabled": ENABLE_QUEUE_CONSUMER,
        "admin_media_uploads_enabled": ENABLE_ADMIN_MEDIA_UPLOADS,
        "startup_db_maintenance_enabled": ENABLE_STARTUP_DB_MAINTENANCE,
        "max_concurrent_jobs": MAX_CONCURRENT_JOBS,
        "request_body_limit_bytes": REQUEST_BODY_LIMIT_BYTES,
        "docx_upload_max_bytes": DOCX_UPLOAD_MAX_BYTES,
        "chapter_upload_max_bytes": CHAPTER_UPLOAD_MAX_BYTES,
        "admin_media_upload_max_bytes": ADMIN_MEDIA_UPLOAD_MAX_BYTES,
        "mongodb_max_pool_size": MONGODB_MAX_POOL_SIZE,
        "public_cache_max_entries": PUBLIC_CACHE_MAX_ENTRIES,
        "home_book_limit": HOME_BOOK_LIMIT,
    }


def _log_cost_control_startup() -> None:
    logger.info(
        "Railway cost-control startup: %s",
        _json.dumps(
            {
                **_cost_control_flags(),
                "openai_key_detected": bool(os.environ.get("OPENAI_API_KEY")),
                "cloudinary_config_detected": _cloudinary_config_detected(),
                "b2_storage_config_detected": _b2_is_configured() if "_b2_is_configured" in globals() else False,
            },
            sort_keys=True,
            ensure_ascii=True,
        ),
    )


def _request_body_limit_for_path(path: str) -> int:
    if path in {"/api/upload_docx", "/api/admin/books/import-template"}:
        return DOCX_UPLOAD_MAX_BYTES
    if path.startswith("/api/admin/books/") and path.endswith("/cover"):
        return ADMIN_MEDIA_UPLOAD_MAX_BYTES
    if path.startswith("/api/admin/books/") and "/chapters/" in path and path.endswith("/upload"):
        return CHAPTER_UPLOAD_MAX_BYTES
    if path == "/api/admin/upload/image":
        return ADMIN_MEDIA_UPLOAD_MAX_BYTES
    return REQUEST_BODY_LIMIT_BYTES


def _require_expensive_job_enabled(job_type: str, *, enabled: bool, confirm_expensive_job: bool) -> None:
    if not enabled:
        _expensive_job_state["blocked"][job_type] += 1
        raise HTTPException(
            status_code=503,
            detail=(
                f"{job_type} is disabled on the production web service by cost-control mode. "
                "Run the explicit manual worker/pipeline command with the matching ENABLE_* flag instead."
            ),
        )
    if not confirm_expensive_job:
        _expensive_job_state["blocked"][job_type] += 1
        raise HTTPException(
            status_code=400,
            detail=f"Set confirm_expensive_job=true to run {job_type}.",
        )


def _expensive_job_lock_for_loop() -> asyncio.Lock:
    global _expensive_job_lock
    if _expensive_job_lock is None:
        _expensive_job_lock = asyncio.Lock()
    return _expensive_job_lock


@asynccontextmanager
async def _expensive_job_slot(job_type: str):
    lock = _expensive_job_lock_for_loop()
    async with lock:
        active = int(_expensive_job_state.get("active", 0))
        if active >= MAX_CONCURRENT_JOBS:
            _expensive_job_state["blocked"][job_type] += 1
            raise HTTPException(
                status_code=429,
                detail=f"Maximum expensive job concurrency reached ({MAX_CONCURRENT_JOBS}).",
            )
        _expensive_job_state["active"] = active + 1
        _expensive_job_state["started"][job_type] += 1
    try:
        logger.warning("Starting explicitly confirmed expensive job: %s", job_type)
        yield
    finally:
        async with lock:
            _expensive_job_state["active"] = max(0, int(_expensive_job_state.get("active", 1)) - 1)

# Controlled-launch public truth gate. The production database can contain
# older published records, but the public launch surface must expose only the
# rights-approved Tier A core reading candidate until the next approval packet
# is intentionally merged.
CONTROLLED_PUBLICATION_TRUTH_GATE_VERSION = "preview-parity-v8"
CONTROLLED_LIVE_BOOK_SLUGS = CATALOG_TRUTH_LIVE_BOOK_SLUGS
CONTROLLED_PIPELINE_SLUGS = tuple(sorted(CATALOG_TRUTH_PIPELINE_SLUGS))
CONTROLLED_AUDIO_ENABLED_SLUGS = tuple(sorted(CATALOG_TRUTH_AUDIO_ENABLED_SLUGS))

# Server-owned pack catalogue. Frontend cannot influence amount/minutes.
# amount is in PAISE (Razorpay's smallest INR unit); minutes is integer minutes.
PACKS: List[dict] = [
    {
        "id": "30m",
        "label": "The First Chapter",
        "minutes": 30,
        "amount_paise": 4900,
        "price_inr": 49,
        "note": "Continue after the free preview, one careful sitting at a time.",
    },
    {
        "id": "1h",
        "label": "The Quiet Hour",
        "minutes": 60,
        "amount_paise": 8900,
        "price_inr": 89,
        "note": "Best first choice — enough time to settle into Dracula.",
    },
    {
        "id": "3h",
        "label": "The Deep Reading Pass",
        "minutes": 180,
        "amount_paise": 23900,
        "price_inr": 239,
        "note": "A longer weekend return to the castle and the count.",
    },
    {
        "id": "10h",
        "label": "The Reader’s Reserve",
        "minutes": 600,
        "amount_paise": 49900,
        "price_inr": 499,
        "note": "Ten quiet hours kept for Dracula and the classics coming next.",
    },
]
PACKS_BY_ID = {p["id"]: p for p in PACKS}
READER_STREAK_REWARD_KEY = "reader_reserve_streak_3"
READER_STREAK_REQUIRED_DAYS = 3
READER_STREAK_REWARD_SECONDS = 10 * 60


def razorpay_keys_configured() -> bool:
    return bool(RAZORPAY_KEY_ID and RAZORPAY_KEY_SECRET)


def get_razorpay_client():
    """Lazy-import the SDK so the app boots even if it's not installed."""
    if not razorpay_keys_configured():
        return None
    try:
        import razorpay  # type: ignore
    except ImportError:
        return None
    return razorpay.Client(auth=(RAZORPAY_KEY_ID, RAZORPAY_KEY_SECRET))


def _hmac_sha256_hex(secret: str, body: bytes) -> str:
    return hmac.new(secret.encode("utf-8"), body, hashlib.sha256).hexdigest()


# ---------- Auth helpers ----------
def hash_password(p: str) -> str:
    return bcrypt.hashpw(p.encode(), bcrypt.gensalt()).decode()

def verify_password(p: str, h: Optional[str]) -> bool:
    if not h or not isinstance(h, str):
        return False
    try:
        return bcrypt.checkpw(p.encode(), h.encode())
    except (TypeError, ValueError):
        return False


def _password_login_unavailable_detail(user: dict) -> str:
    provider = str(user.get("auth_provider") or "this sign-in method").replace("_", " ")
    if provider == "email":
        provider = "this account"
    return (
        f"This account was created with {provider} and does not have a password credential yet. "
        "Please use the original sign-in method or contact sales@reoenterprise.org to add a password."
    )


def _has_password_credential(user: Optional[dict]) -> bool:
    return bool(user and isinstance(user.get("password_hash"), str) and user.get("password_hash"))

def create_token(sub: str, email: str) -> str:
    payload = {
        "sub": sub,
        "email": email,
        "role": "admin",
        "exp": datetime.now(timezone.utc) + timedelta(minutes=JWT_EXPIRE_MINUTES),
        "type": "access",
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALG)

def create_user_token(sub: str, email: str, session_id: str, device_fingerprint: str) -> str:
    payload = {
        "sub": sub,
        "email": email,
        "role": "user",
        "sid": session_id,
        "fp": device_fingerprint,
        "exp": datetime.now(timezone.utc) + timedelta(minutes=USER_ACCESS_TOKEN_EXPIRE_MINUTES),
        "type": "access",
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALG)


def _as_utc_dt(value) -> Optional[datetime]:
    if isinstance(value, datetime):
        return value.replace(tzinfo=timezone.utc) if value.tzinfo is None else value.astimezone(timezone.utc)
    if isinstance(value, str) and value:
        try:
            parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
            return parsed.replace(tzinfo=timezone.utc) if parsed.tzinfo is None else parsed.astimezone(timezone.utc)
        except ValueError:
            return None
    return None


def _hash_secret(value: str) -> str:
    return hmac.new(JWT_SECRET.encode("utf-8"), value.encode("utf-8"), hashlib.sha256).hexdigest()


def _device_fingerprint(request: Request) -> str:
    ua = request.headers.get("user-agent", "")[:240]
    lang = request.headers.get("accept-language", "")[:80]
    browser_hint = request.headers.get("x-client-fingerprint", "")[:120]
    raw = f"{_client_ip(request)}|{ua}|{lang}|{browser_hint}"
    return _hash_secret(raw)


def _set_user_refresh_cookie(response: Response, refresh_token: str) -> None:
    response.set_cookie(
        key=USER_REFRESH_COOKIE,
        value=refresh_token,
        max_age=USER_REFRESH_TOTAL_HOURS * 3600,
        httponly=True,
        secure=COOKIE_SECURE,
        samesite=COOKIE_SAMESITE,
        path="/",
    )


def _clear_user_refresh_cookie(response: Response) -> None:
    response.delete_cookie(
        key=USER_REFRESH_COOKIE,
        httponly=True,
        secure=COOKIE_SECURE,
        samesite=COOKIE_SAMESITE,
        path="/",
    )


async def _create_user_session(user: dict, request: Request, response: Response) -> str:
    """Create a new reader session and revoke older active sessions.

    TRUSTED_DEVICE_MAX_ACTIVE_SESSIONS defaults to 1, which enforces one active
    login. Raising it later allows a small trusted-device whitelist without
    changing token semantics.
    """
    now = datetime.now(timezone.utc)
    session_id = str(uuid.uuid4())
    refresh_token = secrets.token_urlsafe(48)
    device_fp = _device_fingerprint(request)

    await db.user_sessions.insert_one({
        "id": session_id,
        "user_id": user["id"],
        "email": user["email"],
        "device_fingerprint": device_fp,
        "refresh_token_hash": _hash_secret(refresh_token),
        "status": "active",
        "created_at": now,
        "last_seen_at": now,
        "idle_expires_at": now + timedelta(minutes=USER_REFRESH_IDLE_MINUTES),
        "absolute_expires_at": now + timedelta(hours=USER_REFRESH_TOTAL_HOURS),
        "ip_hash": _hash_secret(_client_ip(request)),
        "user_agent": request.headers.get("user-agent", "")[:240],
    })

    active = await db.user_sessions.find(
        {"user_id": user["id"], "status": "active"},
        {"_id": 0, "id": 1, "created_at": 1},
    ).sort("created_at", -1).to_list(50)
    keep = {row["id"] for row in active[:TRUSTED_DEVICE_MAX_ACTIVE_SESSIONS]}
    revoked = [row["id"] for row in active if row["id"] not in keep]
    if revoked:
        await db.user_sessions.update_many(
            {"id": {"$in": revoked}},
            {"$set": {"status": "revoked", "revoked_at": now, "revoked_reason": "new_login"}},
        )

    await db.users.update_one(
        {"id": user["id"]},
        {"$set": {"active_user_session_id": session_id, "last_login_at": now}},
    )
    await _invalidate_user_cache(user["id"], session_ids=revoked)
    await _cache_user_session({
        "id": session_id,
        "user_id": user["id"],
        "email": user["email"],
        "device_fingerprint": device_fp,
        "status": "active",
        "created_at": now,
        "last_seen_at": now,
        "idle_expires_at": now + timedelta(minutes=USER_REFRESH_IDLE_MINUTES),
        "absolute_expires_at": now + timedelta(hours=USER_REFRESH_TOTAL_HOURS),
    })
    _set_user_refresh_cookie(response, refresh_token)
    return create_user_token(user["id"], user["email"], session_id, device_fp)


async def _refresh_user_session(refresh_token: str, request: Request, response: Response) -> Optional[dict]:
    now = datetime.now(timezone.utc)
    token_hash = _hash_secret(refresh_token)
    session = await db.user_sessions.find_one({"refresh_token_hash": token_hash, "status": "active"}, {"_id": 0})
    if not session:
        return None
    idle_expires = _as_utc_dt(session.get("idle_expires_at"))
    absolute_expires = _as_utc_dt(session.get("absolute_expires_at"))
    if not idle_expires or not absolute_expires or idle_expires <= now or absolute_expires <= now:
        await db.user_sessions.update_one(
            {"id": session["id"]},
            {"$set": {"status": "expired", "expired_at": now}},
        )
        await _invalidate_user_cache(session.get("user_id", ""), session_ids=[session["id"]])
        _clear_user_refresh_cookie(response)
        return None

    user = await _cached_user_doc(session["user_id"])
    if user and user.get("role") != "user":
        user = None
    if not user or user.get("status") == "blocked":
        return None
    if TRUSTED_DEVICE_MAX_ACTIVE_SESSIONS <= 1 and user.get("active_user_session_id") != session["id"]:
        await db.user_sessions.update_one(
            {"id": session["id"]},
            {"$set": {"status": "revoked", "revoked_at": now, "revoked_reason": "new_login"}},
        )
        await _invalidate_user_cache(user["id"], session_ids=[session["id"]])
        _clear_user_refresh_cookie(response)
        return None

    device_fp = _device_fingerprint(request)
    if session.get("device_fingerprint") != device_fp:
        await db.user_sessions.update_one(
            {"id": session["id"]},
            {"$set": {"status": "revoked", "revoked_at": now, "revoked_reason": "device_fingerprint_changed"}},
        )
        await _invalidate_user_cache(user["id"], session_ids=[session["id"]])
        _clear_user_refresh_cookie(response)
        return None

    idle_expires_at = now + timedelta(minutes=USER_REFRESH_IDLE_MINUTES)
    await db.user_sessions.update_one(
        {"id": session["id"]},
        {"$set": {"last_seen_at": now, "idle_expires_at": idle_expires_at}},
    )
    session["last_seen_at"] = now
    session["idle_expires_at"] = idle_expires_at
    await _cache_user_session(session)
    token = create_user_token(user["id"], user["email"], session["id"], device_fp)
    return {"token": token, "user": UserOut(**_user_public(user))}

bearer = HTTPBearer(auto_error=False)

async def require_admin(
    creds: Optional[HTTPAuthorizationCredentials] = Depends(bearer),
    ear_session: Optional[str] = Cookie(default=None, alias=SESSION_COOKIE),
) -> dict:
    # Accept either an Authorization Bearer header OR the httpOnly session cookie.
    # Cookies are preferred for browsers; Bearer is kept for server-side tests / curl.
    token: Optional[str] = None
    if creds and creds.scheme.lower() == "bearer":
        token = creds.credentials
    elif ear_session:
        token = ear_session
    if not token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALG])
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")
    if payload.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    return payload


async def require_user(
    request: Request,
    creds: Optional[HTTPAuthorizationCredentials] = Depends(bearer),
) -> dict:
    """Authenticate a reader user via Bearer token; reject admin tokens."""
    if not creds or creds.scheme.lower() != "bearer":
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        payload = jwt.decode(creds.credentials, JWT_SECRET, algorithms=[JWT_ALG])
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")
    if payload.get("role") != "user":
        raise HTTPException(status_code=403, detail="User access required")
    session_id = payload.get("sid")
    token_fp = payload.get("fp")
    if not session_id or not token_fp:
        raise HTTPException(status_code=401, detail="Session expired, please login again.")
    if token_fp != _device_fingerprint(request):
        raise HTTPException(status_code=401, detail="Session security check failed. Please login again.")
    user = await _cached_user_doc(payload["sub"])
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    if user.get("role") != "user":
        raise HTTPException(status_code=403, detail="User access required")
    if user.get("status") == "blocked":
        raise HTTPException(status_code=403, detail="Account is blocked")
    if TRUSTED_DEVICE_MAX_ACTIVE_SESSIONS <= 1 and user.get("active_user_session_id") != session_id:
        raise HTTPException(status_code=401, detail="You've been logged out: new login detected.")
    session = await _cached_user_session(session_id, user["id"])
    if not session or session.get("status") != "active":
        raise HTTPException(status_code=401, detail="You've been logged out: new login detected.")
    now = datetime.now(timezone.utc)
    idle_expires = _as_utc_dt(session.get("idle_expires_at"))
    absolute_expires = _as_utc_dt(session.get("absolute_expires_at"))
    if not idle_expires or not absolute_expires or idle_expires <= now or absolute_expires <= now:
        await db.user_sessions.update_one(
            {"id": session_id},
            {"$set": {"status": "expired", "expired_at": now}},
        )
        await _invalidate_user_cache(user["id"], session_ids=[session_id])
        raise HTTPException(status_code=401, detail="Session expired, please login again.")
    last_seen_at = _as_utc_dt(session.get("last_seen_at"))
    should_touch_session = (
        SESSION_TOUCH_INTERVAL_SECONDS <= 0
        or not last_seen_at
        or (now - last_seen_at).total_seconds() >= SESSION_TOUCH_INTERVAL_SECONDS
    )
    if should_touch_session:
        idle_expires_at = now + timedelta(minutes=USER_REFRESH_IDLE_MINUTES)
        await db.user_sessions.update_one(
            {"id": session_id},
            {"$set": {"last_seen_at": now, "idle_expires_at": idle_expires_at}},
        )
        session["last_seen_at"] = now
        session["idle_expires_at"] = idle_expires_at
        await _cache_user_session(session)
    user["session_id"] = session_id
    user["device_fingerprint"] = token_fp
    return user


async def optional_principal(
    request: Request,
    creds: Optional[HTTPAuthorizationCredentials] = Depends(bearer),
) -> Optional[dict]:
    """Return the caller as a dict tagged with role, or None for guests.

    Never raises. Used by endpoints that return different shapes for
    guests / readers / admins (e.g. content-gated chapter fetch).
    """
    if not creds or creds.scheme.lower() != "bearer":
        return None
    try:
        payload = jwt.decode(creds.credentials, JWT_SECRET, algorithms=[JWT_ALG])
    except (jwt.ExpiredSignatureError, jwt.InvalidTokenError):
        return None
    role = payload.get("role")
    if role == "admin":
        return {"role": "admin", "id": payload.get("sub"), "email": payload.get("email")}
    if role == "user":
        u = await _cached_user_doc(payload.get("sub"))
        if not u:
            return None
        session_id = payload.get("sid")
        if not session_id or payload.get("fp") != _device_fingerprint(request):
            return None
        if TRUSTED_DEVICE_MAX_ACTIVE_SESSIONS <= 1 and u.get("active_user_session_id") != session_id:
            return None
        session = await _cached_user_session(session_id, u["id"])
        if not session or session.get("status") != "active":
            return None
        now = datetime.now(timezone.utc)
        idle_expires = _as_utc_dt(session.get("idle_expires_at"))
        absolute_expires = _as_utc_dt(session.get("absolute_expires_at"))
        if not idle_expires or not absolute_expires or idle_expires <= now or absolute_expires <= now:
            return None
        u["role"] = "user"
        u["session_id"] = session_id
        return u
    return None


def _set_session_cookie(response: Response, token: str) -> None:
    response.set_cookie(
        key=SESSION_COOKIE,
        value=token,
        max_age=SESSION_TTL_SECONDS,
        httponly=True,
        secure=COOKIE_SECURE,
        samesite=COOKIE_SAMESITE,
        path="/",
    )


def _clear_session_cookie(response: Response) -> None:
    response.delete_cookie(key=SESSION_COOKIE, path="/")


# ---------- Utils ----------
def normalize_text(text: str) -> str:
    return unicodedata.normalize("NFC", text or "")


def slugify(text: str, fallback: Optional[str] = None) -> str:
    text = normalize_text(text)
    text = re.sub(r"[^a-zA-Z0-9\s-]", "", text).strip().lower()
    slug = re.sub(r"[\s_-]+", "-", text).strip("-")
    return slug or fallback or str(uuid.uuid4())[:8]

DEFAULT_CATEGORY_SLUG = "literary-fiction"
CANONICAL_CATEGORY_SLUGS = {
    "bengali-classics",
    "literary-fiction",
    "young-readers",
    "business",
    "technology",
    "history-strategy",
    "adventure",
    "science-fiction",
    "gothic-fiction",
}
LEGACY_CATEGORY_SLUG_MAP = {
    "classic-literature": "literary-fiction",
    "literature": "literary-fiction",
    "children-classics": "young-readers",
    "children": "young-readers",
    "business-entrepreneurship": "business",
    "technology-ai": "technology",
    "history-politics": "history-strategy",
    "bengali": "bengali-classics",
    "bengali-reading": "bengali-classics",
}


def _category_value_to_slug(value: str) -> str:
    text = normalize_text(value or "")
    text = re.sub(r"[^a-zA-Z0-9\s-]", "", text).strip().lower()
    return re.sub(r"[\s_-]+", "-", text).strip("-")


def normalize_category_slug(value: str) -> str:
    slug = _category_value_to_slug(value)
    return LEGACY_CATEGORY_SLUG_MAP.get(slug, slug)


def canonical_category_slug(value: str, default: str = DEFAULT_CATEGORY_SLUG) -> str:
    slug = normalize_category_slug(value)
    if slug in CANONICAL_CATEGORY_SLUGS:
        return slug
    return default

def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def topup_intent_expires_at(created_at: Optional[str] = None) -> str:
    created = _as_utc_dt(created_at) or datetime.now(timezone.utc)
    return (created + timedelta(seconds=TOPUP_INTENT_TTL_SECONDS)).isoformat()


def _topup_intent_is_expired(intent: dict) -> bool:
    if intent.get("status") == "credited":
        return False
    expires = _as_utc_dt(intent.get("expires_at"))
    if not expires:
        created = _as_utc_dt(intent.get("created_at"))
        if not created:
            return True
        expires = created + timedelta(seconds=TOPUP_INTENT_TTL_SECONDS)
    return expires <= datetime.now(timezone.utc)


BOOK_METADATA_PROJECTION = {
    "_id": 0,
    "chapters.content": 0,
    "rights_metadata": 0,
}
READER_ACCESS_PROJECTION = {
    "_id": 0,
    "slug": 1,
    "title": 1,
    "is_published": 1,
    "publication_status": 1,
    "approved_to_publish": 1,
    "rights_metadata": 1,
    "source_url": 1,
    "source_name": 1,
    "source_license": 1,
    "source_hash": 1,
    "content_hash": 1,
    "provenance_hash": 1,
    "qa_status": 1,
    "chapters.id": 1,
    "chapters.title": 1,
    "chapters.order": 1,
    "chapters.is_preview": 1,
}
CHAPTER_CONTENT_PROJECTION = {
    "_id": 0,
    "chapters.$": 1,
}
BOOK_LIST_PROJECTION = {
    "_id": 0,
    "rights_metadata": 0,
    "chapters": 0,
}
BOOK_SUMMARY_PROJECTION = {
    "_id": 0,
    "id": 1,
    "slug": 1,
    "title": 1,
    "subtitle": 1,
    "author": 1,
    "category_slug": 1,
    "short_description": 1,
    "cover_url": 1,
    "cover_image_url": 1,
    "thumbnail_url": 1,
    "blur_placeholder": 1,
    "dominant_color": 1,
    "estimated_reading_time": 1,
    "is_published": 1,
    "created_at": 1,
    "updated_at": 1,
    "publication_status": 1,
    "approved_to_publish": 1,
    "rights_metadata": 1,
    "source_url": 1,
    "source_name": 1,
    "source_license": 1,
    "source_hash": 1,
    "content_hash": 1,
    "provenance_hash": 1,
    "qa_status": 1,
    "chapters.id": 1,
    "chapters.is_preview": 1,
}
PUBLIC_CACHE_PATHS = {
    "/api/home",
    "/api/home/books",
    "/api/categories",
    "/api/books",
    "/api/blog",
    "/api/featured",
    "/api/settings/social",
    "/api/settings/brand",
    "/api/settings/public",
    "/api/payments/packs",
    "/api/payments/config",
    "/api/reading/packs",
}
PUBLIC_CACHE_PREFIXES = (
    "/api/books/",
    "/api/blog/",
)


async def initialize_replica_state_backends() -> None:
    global _redis_client, _redis_available
    if not REDIS_CACHE_ENABLED and not MULTI_REPLICA_ENABLED:
        logger.info("Redis cache/state disabled; using per-process local cache and rate-limit state.")
        return
    if not REDIS_URL:
        message = "REDIS_URL is required for shared Redis cache/state."
        if REDIS_CACHE_FAIL_FAST or MULTI_REPLICA_ENABLED:
            raise RuntimeError(message)
        logger.warning("%s Continuing without Redis.", message)
        return
    try:
        import redis.asyncio as redis  # type: ignore
    except Exception as exc:
        message = "Redis cache/state requires the redis Python package."
        if REDIS_CACHE_FAIL_FAST or MULTI_REPLICA_ENABLED:
            raise RuntimeError(message) from exc
        logger.warning("%s Continuing without Redis.", message)
        return
    _redis_client = redis.from_url(
        REDIS_URL,
        socket_connect_timeout=REDIS_SOCKET_CONNECT_TIMEOUT_SECONDS,
        socket_timeout=REDIS_SOCKET_TIMEOUT_SECONDS,
        retry_on_timeout=True,
    )
    try:
        await _redis_client.ping()
    except Exception as exc:
        _redis_client = None
        message = f"Redis cache/state ping failed: {exc}"
        if REDIS_CACHE_FAIL_FAST or MULTI_REPLICA_ENABLED:
            raise RuntimeError(message) from exc
        logger.warning("%s Continuing without Redis.", message)
        return
    _redis_available = True
    await _configure_redis_cache_policy()
    logger.info("Redis-backed cache/state is enabled.")


async def close_replica_state_backends() -> None:
    global _redis_available
    if _redis_client is not None:
        await _redis_client.aclose()
    _redis_available = False


def _redis_key(*parts: str) -> str:
    cleaned = [re.sub(r"[^a-zA-Z0-9_.:-]", "_", str(part)) for part in parts if str(part)]
    return ":".join([REDIS_KEY_PREFIX, *cleaned])


def _redis_state_enabled() -> bool:
    return bool(_redis_available and _redis_client is not None)


def _redis_text(value: Any) -> str:
    if isinstance(value, bytes):
        return value.decode("utf-8", errors="replace")
    return "" if value is None else str(value)


async def _configure_redis_cache_policy() -> None:
    global _redis_config_status
    _redis_config_status = {"attempted": False, "applied": {}, "errors": {}}
    if not REDIS_CONFIGURE_ON_STARTUP or _redis_client is None:
        return
    _redis_config_status["attempted"] = True
    config_pairs: Dict[str, str] = {}
    if REDIS_MAXMEMORY:
        config_pairs["maxmemory"] = REDIS_MAXMEMORY
    if REDIS_MAXMEMORY_POLICY:
        config_pairs["maxmemory-policy"] = REDIS_MAXMEMORY_POLICY
    for name, value in config_pairs.items():
        try:
            await _redis_client.config_set(name, value)
            _redis_config_status["applied"][name] = value
        except Exception as exc:
            _redis_config_status["errors"][name] = str(exc)[:240]
            logger.warning("Redis CONFIG SET %s failed; continuing with provider defaults.", name, exc_info=True)


def _ttl_with_jitter(ttl_seconds: int) -> int:
    ttl = int(ttl_seconds or 0)
    if ttl <= 0:
        return ttl
    jitter = min(max(0, REDIS_CACHE_TTL_JITTER_SECONDS), max(0, ttl // 5))
    if jitter <= 0:
        return ttl
    return ttl + secrets.randbelow(jitter + 1)


REDIS_CACHE_ALLOWED_PAYLOADS = (
    "metadata",
    "reader_manifests",
    "chapter_text",
    "short_lived_user_state",
    "session_state",
    "payment_state",
    "rate_limit_state",
    "reader_rum_aggregates",
)
REDIS_CACHE_EXCLUDED_PAYLOADS = (
    "book_cover_image_binaries",
    "audiobook_binaries",
    "video_binaries",
    "file_upload_streams",
    "response_objects",
    "inline_media_data_uris",
)
MEDIA_DATA_URI_RE = re.compile(
    r"data:(?:image|audio|video|application/octet-stream|application/pdf)/",
    re.IGNORECASE,
)


def _redis_cache_payload_is_media(value: Any, *, _seen: int = 0) -> bool:
    if _seen > 800:
        return False
    if isinstance(value, (bytes, bytearray, memoryview, io.IOBase, Response, UploadFile)):
        return True
    if isinstance(value, str):
        return bool(MEDIA_DATA_URI_RE.search(value))
    if isinstance(value, dict):
        for key, nested in value.items():
            key_text = str(key).lower()
            if isinstance(nested, (bytes, bytearray, memoryview, io.IOBase)):
                return True
            if (
                isinstance(nested, str)
                and key_text in {"body", "blob", "bytes", "binary", "file", "stream", "content"}
                and MEDIA_DATA_URI_RE.search(nested)
            ):
                return True
            if _redis_cache_payload_is_media(nested, _seen=_seen + 1):
                return True
        return False
    if isinstance(value, (list, tuple, set)):
        return any(_redis_cache_payload_is_media(item, _seen=_seen + 1) for item in value)
    return False


def _cache_payload_encode(value: Any) -> bytes:
    blob = pickle.dumps(value, protocol=pickle.HIGHEST_PROTOCOL)
    if len(blob) >= REDIS_CACHE_COMPRESS_MIN_BYTES:
        return b"z:" + zlib.compress(blob, level=6)
    return b"p:" + blob


def _cache_payload_encode_for_redis(namespace: str, value: Any) -> Optional[bytes]:
    if _redis_cache_payload_is_media(value):
        _cache_stats[f"{namespace}_media_skip"] += 1
        logger.info("Redis cache skipped media/binary payload for namespace=%s", namespace)
        return None
    return _cache_payload_encode(value)


def _cache_payload_decode(blob: bytes) -> Any:
    if blob.startswith(b"z:"):
        return pickle.loads(zlib.decompress(blob[2:]))
    if blob.startswith(b"p:"):
        return pickle.loads(blob[2:])
    # Backward-compatible reader for entries written before payload markers.
    return pickle.loads(blob)


def _cache_digest_key(namespace: str, key: str) -> str:
    digest = hashlib.sha256(key.encode("utf-8")).hexdigest()
    return _redis_key("cache", namespace, digest)


async def _redis_cache_get(namespace: str, key: str) -> Any:
    if not _redis_state_enabled():
        _cache_stats[f"{namespace}_miss"] += 1
        return None
    redis_key = _cache_digest_key(namespace, key)
    try:
        blob = await _redis_client.get(redis_key)
    except Exception:
        logger.warning("Redis cache get failed for namespace=%s", namespace, exc_info=True)
        _cache_stats[f"{namespace}_error"] += 1
        return None
    if not blob:
        _cache_stats[f"{namespace}_miss"] += 1
        return None
    try:
        value = _cache_payload_decode(blob)
    except Exception:
        logger.warning("Failed to decode Redis cache entry namespace=%s key=%s", namespace, key)
        _cache_stats[f"{namespace}_error"] += 1
        return None
    _cache_stats[f"{namespace}_hit"] += 1
    return value


async def _redis_cache_set(namespace: str, key: str, value: Any, ttl_seconds: int) -> None:
    if ttl_seconds <= 0 or not _redis_state_enabled():
        return
    redis_key = _cache_digest_key(namespace, key)
    try:
        payload = _cache_payload_encode_for_redis(namespace, value)
        if payload is None:
            return
        await _redis_client.setex(redis_key, _ttl_with_jitter(ttl_seconds), payload)
    except Exception:
        logger.warning("Redis cache set failed for namespace=%s", namespace, exc_info=True)
        _cache_stats[f"{namespace}_error"] += 1


async def _redis_cache_delete(namespace: str, key: str) -> None:
    if not _redis_state_enabled():
        return
    try:
        await _redis_client.delete(_cache_digest_key(namespace, key))
    except Exception:
        logger.warning("Redis cache delete failed for namespace=%s", namespace, exc_info=True)
        _cache_stats[f"{namespace}_error"] += 1


async def _redis_cache_delete_keys(*keys: str) -> None:
    if not keys or not _redis_state_enabled():
        return
    try:
        await _redis_client.delete(*keys)
    except Exception:
        logger.warning("Redis cache key delete failed", exc_info=True)
        _cache_stats["redis_delete_error"] += 1


def _public_cache_storage_key(generation: int, key: str) -> str:
    digest = hashlib.sha256(key.encode("utf-8")).hexdigest()
    return _redis_key("public-cache", str(generation), digest)


def _public_cache_key(scope: str, **params) -> str:
    payload = {
        "truth_gate": CONTROLLED_PUBLICATION_TRUTH_GATE_VERSION,
        **params,
    }
    return f"{scope}:{_json.dumps(payload, sort_keys=True, ensure_ascii=True, default=str)}"


def _controlled_public_book_query(extra: Optional[dict] = None) -> dict:
    return live_approved_mongo_query(extra)


def _is_controlled_public_slug(slug: str) -> bool:
    return str(slug or "").strip().lower() in CONTROLLED_LIVE_BOOK_SLUGS


def _public_projection_is_live(projected: Optional[dict]) -> bool:
    preview_enabled = projected.get("preview_enabled") if projected else None
    return bool(
        projected
        and projected.get("slug") in CONTROLLED_LIVE_BOOK_SLUGS
        and projected.get("publication_status") == "LIVE_APPROVED"
        and projected.get("reader_enabled") is True
        and isinstance(preview_enabled, bool)
        and bool(projected.get("preview_url")) is preview_enabled
        and projected.get("audio_enabled") is False
        and projected.get("audiobook_enabled") is False
    )


def _safe_live_public_projection(book: Optional[dict]) -> Optional[dict]:
    projected = public_book_projection(_strip_all_chapter_content(book)) if book else None
    return projected if _public_projection_is_live(projected) else None


def _slug_is_dracula(slug: str) -> bool:
    return str(slug or "").strip().lower() == "dracula"


def _controlled_artifact_doc(slug: str, *, include_content: bool = False) -> Optional[dict]:
    normalized_slug = str(slug or "").strip().lower()
    if normalized_slug not in CONTROLLED_LIVE_BOOK_SLUGS:
        return None
    doc = load_controlled_artifact_book(normalized_slug, include_content=include_content)
    if not doc:
        return None
    projected = public_book_projection(_strip_all_chapter_content(doc)) or {}
    if not _public_projection_is_live(projected):
        return None
    return doc


def _reader_audio_truth_doc(book: Optional[dict], slug: str) -> Optional[dict]:
    normalized_slug = str(slug or "").strip().lower()
    artifact = _controlled_artifact_doc(normalized_slug, include_content=False)
    if artifact and can_expose_audio({**artifact, "slug": normalized_slug}):
        return artifact
    if normalized_slug in CATALOG_TRUTH_AUDIO_ENABLED_SLUGS:
        return None
    return book


def _dracula_artifact_doc(*, include_content: bool = False) -> Optional[dict]:
    return _controlled_artifact_doc("dracula", include_content=include_content)


def _matches_public_filters(book: dict, *, category_filter: Optional[str] = None, q: str = "") -> bool:
    if category_filter and book.get("category_slug") != category_filter:
        return False
    q_norm = normalize_text(q).strip().lower()
    if not q_norm:
        return True
    searchable = " ".join(
        normalize_text(book.get(key))
        for key in ("title", "subtitle", "author", "short_description", "description", "category_slug")
    ).lower()
    chapter_titles = " ".join(normalize_text(chapter.get("title")) for chapter in book.get("chapters") or []).lower()
    return q_norm in searchable or q_norm in chapter_titles


def _append_controlled_artifact_projections(
    books: list[dict],
    *,
    category_filter: Optional[str] = None,
    q: str = "",
) -> list[dict]:
    existing_slugs = {book.get("slug") for book in books}
    appended: list[dict] = []
    for slug in CONTROLLED_LIVE_BOOK_SLUGS:
        if slug in existing_slugs:
            continue
        artifact = _controlled_artifact_doc(slug, include_content=False)
        if not artifact or not _matches_public_filters(artifact, category_filter=category_filter, q=q):
            continue
        projected = _safe_live_public_projection(artifact)
        if projected:
            appended.append(projected)
    return [*appended, *books]


def _append_dracula_artifact_projection(
    books: list[dict],
    *,
    category_filter: Optional[str] = None,
    q: str = "",
) -> list[dict]:
    return _append_controlled_artifact_projections(books, category_filter=category_filter, q=q)


async def _find_public_book_candidate(
    slug: str,
    projection: dict,
    *,
    include_artifact_content: bool = False,
) -> tuple[Optional[dict], str]:
    if not _is_controlled_public_slug(slug):
        return None, "missing"
    doc = await db.books.find_one(_controlled_public_book_query({"slug": slug}), projection)
    if doc and _safe_live_public_projection(doc):
        return doc, "db"
    artifact = _controlled_artifact_doc(slug, include_content=include_artifact_content)
    if artifact:
        return artifact, "artifact"
    return None, "missing"


async def _public_cache_generation_value() -> int:
    if not _redis_state_enabled():
        return _public_cache_generation
    raw = await _redis_client.get(_redis_key("public-cache", "generation"))
    if raw is None:
        return 0
    try:
        return int(raw)
    except (TypeError, ValueError):
        return 0


async def _public_cache_get(key: str):
    if not PUBLIC_CACHE_ENABLED:
        return None
    if _redis_state_enabled():
        redis_key = _public_cache_storage_key(await _public_cache_generation_value(), key)
        blob = await _redis_client.get(redis_key)
        if not blob:
            return None
        try:
            return _cache_payload_decode(blob)
        except Exception:
            logger.warning("Failed to decode Redis public cache entry: %s", key)
            return None
    item = _public_cache.get(key)
    if not item:
        return None
    expires_at, value = item
    if expires_at <= time.monotonic():
        _public_cache.pop(key, None)
        return None
    _public_cache.move_to_end(key)
    return value


async def _public_cache_set(key: str, value) -> None:
    if not PUBLIC_CACHE_ENABLED:
        return
    if _redis_state_enabled():
        redis_key = _public_cache_storage_key(await _public_cache_generation_value(), key)
        payload = _cache_payload_encode_for_redis("public-cache", value)
        if payload is None:
            return
        await _redis_client.setex(redis_key, _ttl_with_jitter(PUBLIC_CACHE_TTL_SECONDS), payload)
        return
    if _redis_cache_payload_is_media(value):
        _cache_stats["public-cache_media_skip"] += 1
        return
    _public_cache[key] = (time.monotonic() + PUBLIC_CACHE_TTL_SECONDS, value)
    _public_cache.move_to_end(key)
    while len(_public_cache) > PUBLIC_CACHE_MAX_ENTRIES:
        _public_cache.popitem(last=False)


async def _public_cache_clear() -> None:
    global _public_cache_generation
    if _redis_state_enabled():
        await _redis_client.incr(_redis_key("public-cache", "generation"))
        await _redis_client.incr(_redis_key("reader-content-cache", "generation"))
        return
    _public_cache_generation += 1
    _public_cache.clear()


def _client_etag_matches(request: Request, etag: str) -> bool:
    if not etag:
        return False
    raw = request.headers.get("if-none-match", "")
    if not raw:
        return False
    candidates = {item.strip() for item in raw.split(",") if item.strip()}
    return "*" in candidates or etag in candidates


def _is_public_cache_path(path: str) -> bool:
    return path in PUBLIC_CACHE_PATHS or any(path.startswith(prefix) for prefix in PUBLIC_CACHE_PREFIXES)


async def _reader_content_cache_generation_value() -> int:
    if not _redis_state_enabled():
        return 0
    raw = await _redis_client.get(_redis_key("reader-content-cache", "generation"))
    if raw is None:
        return 0
    try:
        return int(raw)
    except (TypeError, ValueError):
        return 0


def _user_cache_key(user_id: str) -> str:
    return _redis_key("user", user_id)


def _user_wallet_cache_key(user_id: str) -> str:
    return _redis_key("user-wallet", user_id)


def _user_session_cache_key(session_id: str) -> str:
    return _redis_key("user-session", session_id)


def _user_transactions_cache_id(user_id: str) -> str:
    return f"user:{user_id}:wallet-transactions"


def _user_payment_intents_cache_id(user_id: str) -> str:
    return f"user:{user_id}:payment-intents"


async def _cache_user_doc(user: Optional[dict]) -> None:
    if not user or not _redis_state_enabled():
        return
    doc = {k: v for k, v in dict(user).items() if k not in {"_id", "password_hash"}}
    try:
        payload = _cache_payload_encode_for_redis("user_doc", doc)
        if payload is None:
            return
        await _redis_client.setex(_user_cache_key(doc["id"]), _ttl_with_jitter(USER_AUTH_CACHE_TTL_SECONDS), payload)
    except Exception:
        logger.warning("Redis user cache set failed", exc_info=True)
        _cache_stats["user_doc_error"] += 1


async def _cached_user_doc(user_id: str) -> Optional[dict]:
    if not user_id:
        return None
    if _redis_state_enabled():
        try:
            blob = await _redis_client.get(_user_cache_key(user_id))
            if blob:
                _cache_stats["user_doc_hit"] += 1
                return _cache_payload_decode(blob)
        except Exception:
            logger.warning("Redis user cache get failed", exc_info=True)
            _cache_stats["user_doc_error"] += 1
    _cache_stats["user_doc_miss"] += 1
    user = await db.users.find_one({"id": user_id}, {"_id": 0, "password_hash": 0})
    await _cache_user_doc(user)
    return user


async def _cache_user_session(session: Optional[dict]) -> None:
    if not session or not _redis_state_enabled():
        return
    doc = {k: v for k, v in dict(session).items() if k != "_id"}
    try:
        payload = _cache_payload_encode_for_redis("user_session", doc)
        if payload is None:
            return
        await _redis_client.setex(_user_session_cache_key(doc["id"]), _ttl_with_jitter(USER_SESSION_CACHE_TTL_SECONDS), payload)
    except Exception:
        logger.warning("Redis user session cache set failed", exc_info=True)
        _cache_stats["user_session_error"] += 1


async def _cached_user_session(session_id: str, user_id: Optional[str] = None) -> Optional[dict]:
    if not session_id:
        return None
    if _redis_state_enabled():
        try:
            blob = await _redis_client.get(_user_session_cache_key(session_id))
            if blob:
                session = _cache_payload_decode(blob)
                if not user_id or session.get("user_id") == user_id:
                    _cache_stats["user_session_hit"] += 1
                    return session
        except Exception:
            logger.warning("Redis user session cache get failed", exc_info=True)
            _cache_stats["user_session_error"] += 1
    _cache_stats["user_session_miss"] += 1
    query = {"id": session_id}
    if user_id:
        query["user_id"] = user_id
    session = await db.user_sessions.find_one(query, {"_id": 0})
    await _cache_user_session(session)
    return session


async def _set_user_wallet_cache(user_id: str, wallet_seconds: int) -> None:
    if not user_id or not _redis_state_enabled():
        return
    try:
        await _redis_client.setex(_user_wallet_cache_key(user_id), _ttl_with_jitter(USER_WALLET_CACHE_TTL_SECONDS), int(wallet_seconds))
    except Exception:
        logger.warning("Redis wallet cache set failed", exc_info=True)
        _cache_stats["user_wallet_error"] += 1


async def _cached_user_wallet_seconds(user_id: str) -> int:
    if _redis_state_enabled():
        try:
            raw = await _redis_client.get(_user_wallet_cache_key(user_id))
            if raw is not None:
                _cache_stats["user_wallet_hit"] += 1
                return int(_redis_text(raw) or 0)
        except Exception:
            logger.warning("Redis wallet cache get failed", exc_info=True)
            _cache_stats["user_wallet_error"] += 1
    _cache_stats["user_wallet_miss"] += 1
    fresh = await db.users.find_one({"id": user_id}, {"_id": 0, "reading_seconds_balance": 1, "wallet_seconds": 1}) or {}
    wallet = int(fresh.get("reading_seconds_balance", fresh.get("wallet_seconds", 0)) or 0)
    await _set_user_wallet_cache(user_id, wallet)
    return wallet


async def _invalidate_user_cache(user_id: str, *, session_ids: Optional[List[str]] = None) -> None:
    if not user_id or not _redis_state_enabled():
        return
    keys = [_user_cache_key(user_id), _user_wallet_cache_key(user_id)]
    for session_id in session_ids or []:
        if session_id:
            keys.append(_user_session_cache_key(session_id))
    await _redis_cache_delete("user-private", _user_transactions_cache_id(user_id))
    await _redis_cache_delete("user-private", _user_payment_intents_cache_id(user_id))
    await _redis_cache_delete_keys(*keys)


async def _reader_book_access_doc(slug: str, *, admin_preview: bool = False) -> Optional[dict]:
    if not admin_preview and not _is_controlled_public_slug(slug):
        return None
    generation = await _reader_content_cache_generation_value()
    cache_key = f"book-access:{CONTROLLED_PUBLICATION_TRUTH_GATE_VERSION}:{generation}:{'admin' if admin_preview else 'public'}:{slug}"
    cached = await _redis_cache_get("reader-content", cache_key)
    if cached is not None:
        return cached
    query = {"slug": slug} if admin_preview else _controlled_public_book_query({"slug": slug})
    doc = await db.books.find_one(
        query,
        READER_ACCESS_PROJECTION,
    )
    if doc and not admin_preview and not can_expose_reader(doc):
        doc = None
    if not doc and not admin_preview:
        doc = _controlled_artifact_doc(slug, include_content=False)
    if doc:
        await _redis_cache_set("reader-content", cache_key, doc, READER_BOOK_CACHE_TTL_SECONDS)
    return doc


async def _reader_chapter_content(slug: str, chapter_id: str, *, admin_preview: bool = False) -> str:
    if not admin_preview and not _is_controlled_public_slug(slug):
        return ""
    generation = await _reader_content_cache_generation_value()
    cache_key = f"chapter-content:{CONTROLLED_PUBLICATION_TRUTH_GATE_VERSION}:{generation}:{'admin' if admin_preview else 'public'}:{slug}:{chapter_id}"
    cached = await _redis_cache_get("reader-content", cache_key)
    if cached is not None:
        return str(cached or "")
    if not admin_preview:
        book = await _reader_book_access_doc(slug)
        if not book:
            return ""
    query = {"slug": slug} if admin_preview else _controlled_public_book_query({"slug": slug})
    content_doc = await db.books.find_one(
        {**query, "chapters.id": chapter_id},
        {"_id": 0, "chapters.$": 1},
    )
    target = ((content_doc or {}).get("chapters") or [{}])[0]
    content = target.get("content", "")
    if not content and not admin_preview:
        artifact = _controlled_artifact_doc(slug, include_content=True) or {}
        target = next((chapter for chapter in artifact.get("chapters") or [] if chapter.get("id") == chapter_id), {})
        content = target.get("content", "")
    await _redis_cache_set("reader-content", cache_key, content, READER_CHAPTER_CACHE_TTL_SECONDS)
    return content


def _stable_digest(value: Any, length: int = 16) -> str:
    payload = _json.dumps(value, sort_keys=True, ensure_ascii=True, default=str, separators=(",", ":"))
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()[:length]


def _reader_chapter_content_version(chapter: dict) -> str:
    content = str(chapter.get("content") or "")
    return _stable_digest({
        "id": chapter.get("id", ""),
        "title": chapter.get("title", ""),
        "order": chapter.get("order", 0),
        "is_preview": bool(chapter.get("is_preview")),
        "updated_at": chapter.get("updated_at", ""),
        "uploaded_at": chapter.get("uploaded_at", ""),
        "word_count": chapter.get("word_count", 0),
        "content_length": len(content),
        "content_hash": hashlib.sha256(content.encode("utf-8")).hexdigest(),
    }, length=20)


def _reader_word_count(text: str) -> int:
    return len(re.findall(r"\b[\w\u0980-\u09FF'-]+\b", text or ""))


def _book_audiobook_doc(book: dict) -> dict:
    value = book.get("audiobook")
    return value if isinstance(value, dict) else {}


def _book_audiobook_provider(book: dict) -> str:
    nested = _book_audiobook_doc(book)
    return str(nested.get("provider") or book.get("audiobook_provider") or "").strip().lower()


def _book_audiobook_url(book: dict) -> str:
    nested = _book_audiobook_doc(book)
    assets = book.get("audiobook_assets") if isinstance(book.get("audiobook_assets"), dict) else {}
    return str(nested.get("url") or assets.get("mp3") or "").strip()


def _book_audiobook_asset_url(book: dict, key: str) -> str:
    nested = _book_audiobook_doc(book)
    nested_assets = nested.get("assets") if isinstance(nested.get("assets"), dict) else {}
    assets = book.get("audiobook_assets") if isinstance(book.get("audiobook_assets"), dict) else {}
    if key == "mp3":
        return str(nested.get("url") or nested_assets.get(key) or assets.get(key) or "").strip()
    return str(nested_assets.get(key) or assets.get(key) or "").strip()


def _audio_asset_looks_like_b2(url: str) -> bool:
    if not url:
        return False
    parsed = urlparse(url)
    endpoint_host = urlparse(B2_S3_ENDPOINT or "").netloc
    path_parts = [unquote(part) for part in parsed.path.split("/") if part]
    return bool(
        parsed.scheme in {"http", "https"}
        and (
            (endpoint_host and parsed.netloc == endpoint_host and path_parts and path_parts[0] == B2_BUCKET)
            or (B2_BUCKET and parsed.netloc.startswith(f"{B2_BUCKET}."))
            or "backblazeb2.com" in parsed.netloc
        )
    )


def _reader_audio_asset_url(book: dict, slug: str, key: str, url: str) -> str:
    if _audio_asset_looks_like_b2(url):
        if key == "mp3":
            return f"/api/reader/book/{slug}/audiobook"
        return f"/api/reader/book/{slug}/audiobook/{key}"
    return url


def _reader_manifest_audio(book: dict, slug: str) -> dict:
    if not can_expose_audio({**book, "slug": slug}):
        version = _stable_digest({
            "enabled": False,
            "slug": slug,
            "truth_gate": CONTROLLED_PUBLICATION_TRUTH_GATE_VERSION,
        })
        return {
            "enabled": False,
            "asset_slug": "",
            "provider": "",
            "voice": "",
            "assets": {},
            "url": "",
            "size": 0,
            "duration_ms": 0,
            "release_gate": "",
            "qa_status": "",
            "sync_mode": "",
            "highlight_sync_enabled": False,
            "version": version,
            "updated_at": "",
        }

    raw_assets = book.get("audiobook_assets") if isinstance(book.get("audiobook_assets"), dict) else {}
    assets = {
        str(key): _reader_audio_asset_url(book, slug, str(key), str(value))
        for key, value in raw_assets.items()
        if value
    }
    nested = _book_audiobook_doc(book)
    provider = _book_audiobook_provider(book)
    audio_slug = slugify(book.get("audio_asset_slug") or slug, fallback=slug)
    release_gate = "APPROVED"
    qa_status = audio_release_qa_status({**book, "slug": slug})
    sync_mode = book.get("sync_mode", nested.get("sync_mode", ""))
    highlight_sync_enabled = bool(
        book.get("highlight_sync_enabled", nested.get("highlight_sync_enabled", False))
    )
    enabled = bool(
        book.get("audiobook_enabled")
        or book.get("generate_audiobook")
        or assets
        or _book_audiobook_url(book)
    )
    version = _stable_digest({
        "enabled": enabled,
        "audio_slug": audio_slug,
        "provider": provider,
        "voice": book.get("audiobook_voice", ""),
        "assets": assets,
        "audiobook": nested,
        "release_gate": release_gate,
        "qa_status": qa_status,
        "sync_mode": sync_mode,
        "highlight_sync_enabled": highlight_sync_enabled,
        "truth_gate": CONTROLLED_PUBLICATION_TRUTH_GATE_VERSION,
        "updated_at": book.get("audiobook_assets_updated_at", ""),
    })
    return {
        "enabled": enabled,
        "asset_slug": audio_slug,
        "provider": provider,
        "voice": book.get("audiobook_voice", ""),
        "assets": assets,
        "url": _reader_audio_asset_url(book, slug, "mp3", _book_audiobook_url(book)),
        "size": int(nested.get("size", 0) or 0),
        "duration_ms": int(nested.get("duration_ms", nested.get("duration", 0)) or 0),
        "release_gate": release_gate,
        "qa_status": qa_status,
        "sync_mode": sync_mode,
        "highlight_sync_enabled": highlight_sync_enabled,
        "version": version,
        "updated_at": book.get("audiobook_assets_updated_at", ""),
    }


async def _reader_book_manifest_doc(slug: str, *, admin_preview: bool = False) -> Optional[dict]:
    if not admin_preview and not _is_controlled_public_slug(slug):
        return None
    generation = await _reader_content_cache_generation_value()
    cache_key = f"book-manifest:{CONTROLLED_PUBLICATION_TRUTH_GATE_VERSION}:{generation}:{'admin' if admin_preview else 'public'}:{slug}"
    cached = await _redis_cache_get("reader-manifest", cache_key)
    if cached is not None:
        return cached

    query = {"slug": slug} if admin_preview else _controlled_public_book_query({"slug": slug})
    doc = await db.books.find_one(query, {"_id": 0, "rights_metadata": 0})
    if doc and not admin_preview and not can_expose_reader(doc):
        doc = None
    if not doc and not admin_preview:
        doc = _controlled_artifact_doc(slug, include_content=True)
    if not doc:
        return None

    chapters = []
    preview_ids = _free_preview_chapter_ids(doc)
    for chapter in sorted((doc.get("chapters") or []), key=lambda c: c.get("order", 0)):
        version = _reader_chapter_content_version(chapter)
        chapter_id = chapter.get("id") or ""
        chapters.append({
            "id": chapter_id,
            "title": chapter.get("title", ""),
            "order": chapter.get("order", 0),
            "is_preview": chapter_id in preview_ids,
            "content_version": version,
            "word_count": int(chapter.get("word_count", 0) or _reader_word_count(chapter.get("content", ""))),
            "reading_minutes": int(chapter.get("reading_minutes", 0) or 0),
            "processing_status": chapter.get("processing_status", "ready"),
            "has_images": bool(chapter.get("has_images", False)),
            "content_url": f"/api/reader/chapter/{slug}/{chapter_id}?v={version}" if chapter_id else "",
        })

    audio_source = doc if admin_preview else (_reader_audio_truth_doc(doc, slug) or {})
    audio = _reader_manifest_audio(audio_source, slug)
    book_public = public_book_projection(_strip_all_chapter_content(doc)) or {}
    if not admin_preview and not _public_projection_is_live(book_public):
        return None
    book_public["chapters"] = chapters
    manifest_version = _stable_digest({
        "slug": slug,
        "admin_preview": admin_preview,
        "book": {
            "title": doc.get("title", ""),
            "updated_at": doc.get("updated_at", ""),
            "created_at": doc.get("created_at", ""),
            "is_published": doc.get("is_published", False),
        },
        "chapters": [{c["id"]: c["content_version"]} for c in chapters],
        "audio": audio.get("version", ""),
    }, length=20)
    result = {
        "book": book_public,
        "chapters": chapters,
        "audio": audio,
        "version": manifest_version,
        "content_generation": generation,
        "generated_at": now_iso(),
    }
    await _redis_cache_set("reader-manifest", cache_key, result, READER_MANIFEST_CACHE_TTL_SECONDS)
    return result


HTML_TAG_RE = re.compile(r"</?[a-z][\s\S]*>", re.IGNORECASE)


def _manual_content_to_render_html(content: str) -> tuple[str, List[str]]:
    text = normalize_text(content or "")
    if not text.strip():
        return "", []
    if not HTML_TAG_RE.search(text):
        paragraphs = re.split(r"\n{2,}", text.strip())
        html = "".join(
            f"<p>{_html.escape(p.strip()).replace(chr(10), '<br>')}</p>"
            for p in paragraphs
            if p.strip()
        )
        return html, []
    try:
        from utils.content_processor import sanitize_chapter_html_fragment  # type: ignore
    except Exception:
        return _html.escape(text), ["Content sanitizer unavailable; HTML was escaped."]
    return sanitize_chapter_html_fragment(text)


def _publish_blockers(book: dict) -> List[str]:
    if not normalize_text(book.get("title", "")).strip():
        return ["Title is required."]
    if not book.get("is_published"):
        return []

    blockers: List[str] = []
    slug = str(book.get("slug") or "").strip().lower()
    if slug not in CONTROLLED_LIVE_BOOK_SLUGS:
        blockers.append(
            "Publication safety allowlist blocks live publication for non-approved books."
        )
    if not (book.get("cover_image_url") or book.get("cover_url")):
        blockers.append("Front cover is required before publishing.")

    for issue in rights_publish_blockers(book):
        blockers.append(f"Rights verification: {issue}")

    for chapter in book.get("chapters") or []:
        status = chapter.get("processing_status") or "ready"
        if status != "ready":
            title = normalize_text(chapter.get("title", "")).strip() or "Untitled chapter"
            blockers.append(f"Chapter '{title}' is {status}.")
    return blockers


def _assert_publishable(book: dict) -> None:
    blockers = _publish_blockers(book)
    if blockers:
        raise HTTPException(status_code=400, detail={"message": "Book is not ready to publish.", "issues": blockers})


def _assert_public_rights_approved(book: dict, asset_label: str) -> None:
    if not book.get("is_published"):
        return
    blockers = rights_publish_blockers(book)
    if blockers:
        raise HTTPException(
            status_code=400,
            detail={
                "message": f"{asset_label} cannot be published without approved rights metadata.",
                "issues": [f"Rights verification: {issue}" for issue in blockers],
            },
        )


BOOK_TEMPLATE_LABELS = {
    "title": "title",
    "booktitle": "title",
    "subtitle": "subtitle",
    "author": "author",
    "category": "category_slug",
    "categoryslug": "category_slug",
    "shortdescription": "short_description",
    "description": "description",
    "estimatedreadingtime": "estimated_reading_time",
    "readingtime": "estimated_reading_time",
    "formats": "formats",
    "format": "formats",
    "benefits": "benefits",
    "whothisisfor": "who_for",
    "whofor": "who_for",
    "whatyouwilllearn": "learnings",
    "learnings": "learnings",
    "aboutauthor": "about_author",
    "abouttheauthor": "about_author",
    "aboutpublisher": "about_author",
    "abouttheauthorpublisher": "about_author",
    "buyurl": "buy_url",
    "paperbackprice": "price_paperback",
    "ebookprice": "price_ebook",
}
BOOK_TEMPLATE_LIST_FIELDS = {"formats", "benefits", "who_for", "learnings"}
BOOK_TEMPLATE_MULTILINE_FIELDS = {"short_description", "description", "about_author"}


def _book_template_label_key(label: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", normalize_text(label).strip().lower())


def _append_book_template_value(target: dict, field: str, value: str) -> None:
    value = normalize_text(value).strip()
    if not value:
        return
    target.setdefault(field, []).append(value)


def _book_template_list(values: List[str]) -> List[str]:
    out: List[str] = []
    for value in values:
        for part in re.split(r"[\n,;]", value):
            item = re.sub(r"^\s*[-*•–—\d.)]+\s*", "", normalize_text(part)).strip()
            if item:
                out.append(item)
    return out


def _parse_book_template_docx(file_bytes: bytes) -> Tuple[dict, List[str], str]:
    try:
        import mammoth  # type: ignore
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"DOCX parser unavailable: {e}")

    result = mammoth.extract_raw_text(io.BytesIO(file_bytes))
    raw_text = normalize_text(result.value or "")
    warnings = [str(m.message) for m in result.messages if getattr(m, "message", None)]
    collected: dict = {}
    current_field: Optional[str] = None

    for raw_line in raw_text.splitlines():
        line = normalize_text(raw_line).strip()
        if not line:
            continue
        match = re.match(r"^([^:：]{2,80})[:：]\s*(.*)$", line)
        if match:
            field = BOOK_TEMPLATE_LABELS.get(_book_template_label_key(match.group(1)))
            if field:
                current_field = field
                _append_book_template_value(collected, field, match.group(2))
                continue
            current_field = None
        elif current_field:
            _append_book_template_value(collected, current_field, line)

    def scalar(field: str, default: str = "") -> str:
        values = collected.get(field) or []
        if field in BOOK_TEMPLATE_MULTILINE_FIELDS:
            return "\n".join(values).strip() or default
        return (values[0].strip() if values else "") or default

    category_value = scalar("category_slug", "business")
    imported = {
        "title": scalar("title"),
        "subtitle": scalar("subtitle"),
        "author": scalar("author", "The Earnalism"),
        "category_slug": canonical_category_slug(category_value, default="business"),
        "short_description": scalar("short_description"),
        "description": scalar("description"),
        "estimated_reading_time": scalar("estimated_reading_time"),
        "price_paperback": scalar("price_paperback"),
        "price_ebook": scalar("price_ebook"),
        "buy_url": scalar("buy_url"),
        "formats": _book_template_list(collected.get("formats") or []) or ["Ebook"],
        "benefits": _book_template_list(collected.get("benefits") or []),
        "who_for": _book_template_list(collected.get("who_for") or []),
        "learnings": _book_template_list(collected.get("learnings") or []),
        "about_author": scalar("about_author"),
        "is_published": False,
    }
    if not imported["title"]:
        warnings.append("No Title field was found in the DOCX template.")
    return imported, warnings, raw_text


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
SECURE_STORAGE_DIR = ROOT_DIR / "secure_storage"
AI_MARKER_PATTERNS = [
    r"\bas an ai language model\b",
    r"\bi cannot (?:provide|assist|help)\b",
    r"\bit is important to note\b",
    r"\bin conclusion\b",
    r"\bdelve into\b",
]
COPYRIGHT_SURFACE_PATTERNS = [
    r"\bit was the best of times\b",
    r"\bcall me ishmael\b",
    r"\ball happy families are alike\b",
]
OFFENSIVE_SURFACE_PATTERNS = [
    r"\bkill all\b",
    r"\bexterminate\b",
    r"\bgenocide\b",
]
BENGALI_TEXT_RE = re.compile(r"[\u0980-\u09FF]")
KEYWORD_STOPWORDS = {
    "about", "after", "again", "also", "because", "before", "being", "between", "chapter",
    "could", "every", "first", "from", "have", "into", "more", "other", "their", "there",
    "these", "this", "through", "under", "were", "when", "where", "which", "with", "would",
    "your", "book", "author", "story", "page", "part", "will",
}


def contains_bengali_text(value: str) -> bool:
    return bool(BENGALI_TEXT_RE.search(value or ""))


def _sanitize_docx_plain(value: str, limit: int = 5000) -> str:
    value = normalize_text(value or "")
    value = re.sub(r"<[^>]+>", "", value)
    value = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", value)
    return value.strip()[:limit]


def _safe_storage_filename(filename: str) -> str:
    stem = Path(filename or "manuscript.docx").stem
    safe_stem = re.sub(r"[^a-zA-Z0-9._-]+", "-", stem).strip("-")[:80] or "manuscript"
    return f"{safe_stem}.docx"


def _paragraph_text(paragraph) -> str:
    return _sanitize_docx_plain(paragraph.text or "", 10000)


def _paragraph_is_italic(paragraph) -> bool:
    runs = [run for run in paragraph.runs if (run.text or "").strip()]
    return bool(runs) and all(bool(run.italic) for run in runs)


def _derive_docx_metadata(document, raw_text: str, template_data: Optional[dict] = None) -> dict:
    paragraphs = list(document.paragraphs)
    title = ""
    subtitle = ""
    author = ""
    chapter_count = 0

    for index, paragraph in enumerate(paragraphs):
        style = (paragraph.style.name if paragraph.style else "").lower()
        text = _paragraph_text(paragraph)
        if not text:
            continue
        if not title and "heading 1" in style:
            title = text
            for next_para in paragraphs[index + 1:index + 5]:
                next_text = _paragraph_text(next_para)
                if next_text and _paragraph_is_italic(next_para):
                    subtitle = next_text
                    break
        if "heading 2" in style:
            chapter_count += 1
        if not author:
            match = re.match(r"^\s*by\s+(.+)$", text, flags=re.IGNORECASE)
            if match:
                author = _sanitize_docx_plain(match.group(1), 180)

    lines = [_sanitize_docx_plain(line, 300) for line in raw_text.splitlines() if _sanitize_docx_plain(line, 300)]
    if not title and lines:
        title = lines[0]
    if not subtitle:
        for line in lines[1:5]:
            if not re.match(r"^\s*by\s+", line, flags=re.IGNORECASE):
                subtitle = line
                break
    if not author:
        for line in lines[:20]:
            match = re.match(r"^\s*by\s+(.+)$", line, flags=re.IGNORECASE)
            if match:
                author = _sanitize_docx_plain(match.group(1), 180)
                break

    props = getattr(document, "core_properties", None)
    keyword_source = _sanitize_docx_plain(getattr(props, "keywords", "") if props else "", 500)
    if keyword_source:
        keywords = [item.strip() for item in re.split(r"[,;]", keyword_source) if item.strip()][:5]
    else:
        words = re.findall(r"[A-Za-z][A-Za-z'-]{3,}", raw_text.lower())
        counts = defaultdict(int)
        for word in words:
            base = word.strip("'-")
            if base and base not in KEYWORD_STOPWORDS:
                counts[base] += 1
        keywords = [item[0] for item in sorted(counts.items(), key=lambda pair: (-pair[1], pair[0]))[:5]]

    word_count = len(re.findall(r"\b[\w\u0980-\u09FF'-]+\b", raw_text or ""))
    template_data = template_data or {}
    return {
        "title": template_data.get("title") or title,
        "subtitle": template_data.get("subtitle") or subtitle,
        "author": template_data.get("author") or author or "The Earnalism",
        "keywords": keywords,
        "chapter_count": chapter_count,
        "word_count": word_count,
    }


def _validate_docx_document(document, raw_text: str) -> Tuple[List[dict], List[str]]:
    checks: List[dict] = []
    warnings: List[str] = []

    def add_check(name: str, passed: bool, severity: str, detail: str, findings: Optional[list] = None):
        checks.append({
            "name": name,
            "status": "PASS" if passed else "FAIL",
            "severity": severity,
            "detail": detail,
            "findings": findings or [],
        })

    text = raw_text or ""
    ai_findings = []
    for idx, paragraph in enumerate(document.paragraphs):
        paragraph_text = _paragraph_text(paragraph)
        for pattern in AI_MARKER_PATTERNS:
            if re.search(pattern, paragraph_text, flags=re.IGNORECASE):
                ai_findings.append({"paragraph_index": idx, "excerpt": paragraph_text[:220]})
    add_check(
        "AI-content marker scan",
        not ai_findings,
        "WARNING",
        "No common AI-authorship markers found." if not ai_findings else "Common AI-authorship markers were found for admin review.",
        ai_findings,
    )
    if ai_findings:
        warnings.append("Potential AI-content markers detected. Review manually before publishing.")

    copyright_findings = []
    for pattern in COPYRIGHT_SURFACE_PATTERNS:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            copyright_findings.append({"excerpt": text[max(0, match.start() - 80):match.end() + 80]})
    add_check(
        "Plagiarism surface scan",
        not copyright_findings,
        "WARNING",
        "No known common passage matches found." if not copyright_findings else "Known common passage surface matches found.",
        copyright_findings,
    )

    offensive_findings = []
    for pattern in OFFENSIVE_SURFACE_PATTERNS:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            offensive_findings.append({"excerpt": text[max(0, match.start() - 80):match.end() + 80]})
    add_check(
        "Offensive text surface scan",
        not offensive_findings,
        "WARNING",
        "No configured offensive text markers found." if not offensive_findings else "Configured offensive text markers were found.",
        offensive_findings,
    )

    heading_names = [(p.style.name if p.style else "") for p in document.paragraphs if _paragraph_text(p)]
    has_heading = any("Heading" in name for name in heading_names)
    if not has_heading:
        warnings.append("No DOCX heading styles were detected. Chapter detection may be incomplete.")
    add_check(
        "Heading hierarchy",
        has_heading,
        "INFO" if has_heading else "WARNING",
        "DOCX heading styles detected." if has_heading else "No heading styles detected.",
    )

    inline_images = len(getattr(document, "inline_shapes", []))
    floating_images = document.element.xml.count("<wp:anchor")
    add_check(
        "Inline image rendering",
        floating_images == 0,
        "INFO" if floating_images == 0 else "WARNING",
        f"{inline_images} inline image(s), {floating_images} floating image layout risk(s).",
        [{"inline_images": inline_images, "floating_images": floating_images}],
    )
    if floating_images:
        warnings.append("Floating images can shift unpredictably in reader preview. Prefer inline images.")

    readable = len(text.strip()) > 0
    add_check(
        "Structural integrity",
        readable,
        "ERROR" if not readable else "INFO",
        "DOCX text was readable." if readable else "No readable text could be extracted.",
    )
    return checks, warnings


def _format_docx_document(document) -> None:
    try:
        from docx.shared import Pt  # type: ignore
    except Exception:
        Pt = None
    for paragraph in document.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        clean_text = _sanitize_docx_plain(paragraph.text, 20000)
        if paragraph.text != clean_text:
            for run in paragraph.runs:
                run.text = ""
            paragraph.add_run(clean_text)
        if Pt:
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.line_spacing = 1.15
            for run in paragraph.runs:
                run.font.name = "Noto Serif Bengali" if contains_bengali_text(run.text) else "Georgia"
                if "Heading 1" in style_name:
                    run.font.size = Pt(22)
                    run.bold = True
                elif "Heading 2" in style_name:
                    run.font.size = Pt(18)
                    run.bold = True
                elif "Heading 3" in style_name:
                    run.font.size = Pt(15)
                    run.bold = True
                else:
                    run.font.size = Pt(12)


def _docx_credit_rows(admin_user_id: str, session_id: str, file_name: str, upload_id: str, tasks: List[dict]) -> List[dict]:
    now = now_iso()
    rows = []
    for task in tasks:
        units = float(task.get("units", 0) or 0)
        credits = round(float(task.get("credits_used", 0) or 0), 4)
        rows.append({
            "id": str(uuid.uuid4()),
            "user_id": admin_user_id,
            "session_id": session_id,
            "file_name": file_name,
            "upload_id": upload_id,
            "operation_type": "docx_validation_upload",
            "task": task.get("task", "unknown"),
            "units": units,
            "credits_used": credits,
            "timestamp": now,
        })
    return rows


async def _store_admin_upload_artifact(upload_id: str, safe_filename: str, kind: str, body: bytes, content_type: str) -> str:
    if MULTI_REPLICA_ENABLED:
        file_name = f"{upload_id}/{kind}/{safe_filename}"
        file_id = await admin_upload_files.upload_from_stream(
            file_name,
            body,
            metadata={
                "upload_id": upload_id,
                "kind": kind,
                "filename": safe_filename,
                "content_type": content_type,
                "created_at": now_iso(),
            },
        )
        return f"gridfs://admin_upload_files/{file_id}"

    storage_dir = SECURE_STORAGE_DIR / upload_id
    target_dir = storage_dir / "validated" if kind == "validated" else storage_dir
    target_dir.mkdir(parents=True, exist_ok=True)
    target_path = target_dir / safe_filename
    target_path.write_bytes(body)
    prefix = "secure_storage/{upload_id}/validated" if kind == "validated" else "secure_storage/{upload_id}"
    return f"{prefix.format(upload_id=upload_id)}/{safe_filename}"


async def _process_docx_upload(
    *,
    docx_file: UploadFile,
    request: Request,
    admin: dict,
    front_cover: Optional[UploadFile] = None,
    back_cover: Optional[UploadFile] = None,
) -> dict:
    start = time.perf_counter()
    filename = docx_file.filename or "book-upload.docx"
    content_type = (docx_file.content_type or "").split(";", 1)[0].strip().lower()
    if not filename.lower().endswith(".docx") or content_type != DOCX_MIME:
        raise HTTPException(status_code=400, detail="Upload a valid .DOCX file with the correct Office Open XML MIME type.")
    body = await docx_file.read()
    if len(body) > DOCX_UPLOAD_MAX_BYTES:
        raise HTTPException(status_code=400, detail=f"DOCX upload must be under {DOCX_UPLOAD_MAX_BYTES} bytes.")

    upload_id = str(uuid.uuid4())
    safe_filename = _safe_storage_filename(filename)

    try:
        await _store_admin_upload_artifact(upload_id, safe_filename, "original", body, DOCX_MIME)
        from docx import Document  # type: ignore
        document = Document(io.BytesIO(body))
    except Exception as e:
        await db.admin_upload_audit.insert_one({
            "id": str(uuid.uuid4()),
            "upload_id": upload_id,
            "admin_user_id": admin.get("sub", ""),
            "admin_email": admin.get("email", ""),
            "file_name": filename,
            "status": "failed",
            "error": f"Unreadable DOCX: {e}",
            "created_at": now_iso(),
        })
        raise HTTPException(status_code=400, detail="DOCX could not be read. Please retry with a valid Office Open XML file.")

    raw_text_result = _parse_book_template_docx(body)
    template_data, template_warnings, raw_text = raw_text_result
    metadata = _derive_docx_metadata(document, raw_text, template_data)
    checks, warnings = _validate_docx_document(document, raw_text)
    warnings.extend(template_warnings)
    _format_docx_document(document)
    formatted_buffer = io.BytesIO()
    document.save(formatted_buffer)
    storage_ref = await _store_admin_upload_artifact(upload_id, safe_filename, "validated", formatted_buffer.getvalue(), DOCX_MIME)

    book_data = {
        **template_data,
        "title": metadata["title"],
        "subtitle": metadata["subtitle"],
        "author": metadata["author"],
        "estimated_reading_time": template_data.get("estimated_reading_time") or f"{max(1, round(metadata['word_count'] / 238))} min",
        "is_published": False,
        "audiobook_enabled": False,
        "generate_audiobook": False,
    }
    if metadata["keywords"]:
        book_data["keywords"] = metadata["keywords"]

    import_id = f"docx-{upload_id[:12]}"

    if front_cover or back_cover:
        _require_expensive_job_enabled(
            "cover_generation",
            enabled=ENABLE_COVER_GENERATION,
            confirm_expensive_job=True,
        )

    async def process_import_cover(file: Optional[UploadFile], kind: str) -> Optional[dict]:
        if not file:
            return None
        if file.content_type not in _ALLOWED_COVER_TYPES:
            raise HTTPException(status_code=400, detail=f"Unsupported {kind} cover type. Use JPG, PNG, WebP, or GIF.")
        cover_body = await file.read()
        if len(cover_body) > ADMIN_MEDIA_UPLOAD_MAX_BYTES:
            raise HTTPException(status_code=400, detail=f"{kind.title()} cover must be under {ADMIN_MEDIA_UPLOAD_MAX_BYTES} bytes")
        _ensure_cloudinary()
        try:
            from utils.content_processor import process_book_cover  # type: ignore
            return process_book_cover(cover_body, import_id, kind=kind)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"{kind.title()} cover processing failed: {e}")

    front = await process_import_cover(front_cover, "front")
    back = await process_import_cover(back_cover, "back")
    if front:
        book_data.update({
            "cover_image_url": front["cover_url"],
            "cover_url": front["cover_url"],
            "thumbnail_url": front["thumbnail_url"],
        })
    if back:
        book_data.update({
            "back_cover_image_url": back["cover_url"],
            "back_cover_url": back["cover_url"],
            "back_cover_thumbnail_url": back["thumbnail_url"],
        })

    elapsed_ms = round((time.perf_counter() - start) * 1000, 2)
    ai_tokens = sum(len(finding.get("excerpt", "").split()) for check in checks for finding in check.get("findings", []))
    credit_usage = [
        {"task": "upload_io", "units": round(len(body) / 1024, 2), "unit_label": "KB", "credits_used": round(len(body) / (1024 * 1024) * 0.05, 4)},
        {"task": "docx_parse", "units": metadata["word_count"], "unit_label": "words", "credits_used": round(metadata["word_count"] / 1000 * 0.1, 4)},
        {"task": "validation_rules", "units": len(checks), "unit_label": "checks", "credits_used": round(len(checks) * 0.02, 4)},
        {"task": "formatting", "units": len(document.paragraphs), "unit_label": "paragraphs", "credits_used": round(len(document.paragraphs) * 0.002, 4)},
        {"task": "ai_marker_review", "units": ai_tokens, "unit_label": "tokens", "credits_used": round(ai_tokens / 1000 * 0.2, 4)},
        {"task": "compute_time", "units": elapsed_ms / 1000, "unit_label": "seconds", "credits_used": round(elapsed_ms / 1000 * 0.01, 4)},
    ]
    total_credits = round(sum(item["credits_used"] for item in credit_usage), 4)
    admin_user_id = admin.get("sub", admin.get("email", "admin"))
    session_id = getattr(request.state, "request_id", upload_id)
    credit_rows = _docx_credit_rows(admin_user_id, session_id, filename, upload_id, credit_usage)
    if credit_rows:
        await db.credit_log.insert_many(credit_rows)

    status = "failed" if any(check["severity"] == "ERROR" and check["status"] == "FAIL" for check in checks) else "passed"
    audit = {
        "id": str(uuid.uuid4()),
        "upload_id": upload_id,
        "admin_user_id": admin_user_id,
        "admin_email": admin.get("email", ""),
        "file_name": filename,
        "operation_type": "docx_validation_upload",
        "status": status,
        "word_count": metadata["word_count"],
        "chapter_count": metadata["chapter_count"],
        "credits_used": total_credits,
        "storage_ref": storage_ref,
        "created_at": now_iso(),
    }
    await db.admin_upload_audit.insert_one(audit)

    return {
        "success": status == "passed",
        "upload_id": upload_id,
        "book": book_data,
        "metadata": metadata,
        "validation_summary": {
            "title": "Validation Result – Earnalism Compliance v1.0",
            "status": status.upper(),
            "checks": checks,
            "warnings": sorted(set(warnings)),
            "formatted_file": {
                "file_name": safe_filename,
                "storage_ref": storage_ref,
            },
            "footer_note": "Validated content complies with anti-AI & copyright policies.",
        },
        "warnings": sorted(set(warnings)),
        "credit_usage": credit_usage,
        "credits_used": total_credits,
        "credit_report": {
            "user_id": admin_user_id,
            "session_id": session_id,
            "upload_id": upload_id,
        },
    }


# ---------- Models ----------
class LoginIn(BaseModel):
    email: EmailStr
    password: str

class ChangePasswordIn(BaseModel):
    current_password: str
    new_password: str

class TokenOut(BaseModel):
    token: str
    email: str
    role: str

class Category(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    slug: str
    name: str
    description: str = ""
    image_url: str = ""
    order: int = 0

class CategoryIn(BaseModel):
    name: str
    description: str = ""
    image_url: str = ""
    order: int = 0
    slug: Optional[str] = None

class Chapter(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    content: str = ""
    order: int = 0
    is_preview: bool = False
    has_images: bool = False
    image_count: int = 0
    word_count: int = 0
    reading_minutes: int = 0
    language_hint: str = ""
    processing_status: str = "ready"
    processing_error: str = ""
    processing_warnings: List[str] = Field(default_factory=list)
    source_filename: str = ""
    uploaded_at: str = ""
    updated_at: str = ""

class ChapterIn(BaseModel):
    title: str
    content: str = ""
    is_preview: bool = False

class ChapterReorderIn(BaseModel):
    ids: List[str]

class Book(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    slug: str
    title: str
    subtitle: str = ""
    author: str = "The Earnalism"
    category_slug: str
    short_description: str = ""
    description: str = ""
    cover_url: str = ""
    cover_image_url: str = ""
    thumbnail_url: str = ""
    blur_placeholder: str = ""
    dominant_color: str = ""
    back_cover_url: str = ""
    back_cover_image_url: str = ""
    back_cover_thumbnail_url: str = ""
    back_cover_blur_placeholder: str = ""
    back_cover_dominant_color: str = ""
    cover_processing_status: str = ""
    cover_processing_error: str = ""
    back_cover_processing_status: str = ""
    back_cover_processing_error: str = ""
    estimated_reading_time: str = ""
    price_paperback: str = ""
    price_ebook: str = ""
    buy_url: str = ""
    formats: List[str] = Field(default_factory=lambda: ["Paperback", "Ebook"])
    benefits: List[str] = Field(default_factory=list)
    who_for: List[str] = Field(default_factory=list)
    learnings: List[str] = Field(default_factory=list)
    about_author: str = ""
    chapters: List[Chapter] = Field(default_factory=list)
    audiobook_enabled: bool = False
    generate_audiobook: bool = False
    audiobook_provider: str = ""
    audiobook_voice: str = ""
    audiobook_assets_updated_at: str = ""
    audio_asset_slug: str = ""
    audiobook_assets: Dict[str, str] = Field(default_factory=dict)
    audiobook: Dict[str, Any] = Field(default_factory=dict)
    rights_metadata: Dict[str, Any] = Field(default_factory=dict)
    readerStatus: str = "ready_for_editorial_review"
    publicationStatus: str = "draft"
    isPublic: bool = False
    isLive: bool = False
    showInPublicLibrary: bool = False
    showInHomepage: bool = False
    allowPublicReading: bool = False
    allowCheckout: bool = False
    allowPayment: bool = False
    is_published: bool = False
    created_at: str = Field(default_factory=now_iso)


class PublicChapterOut(BaseModel):
    model_config = ConfigDict(extra="ignore")

    id: str = ""
    title: str = ""
    order: int = 0
    is_preview: bool = False
    has_images: bool = False
    image_count: int = 0
    word_count: int = 0
    reading_minutes: int = 0
    language_hint: str = ""
    processing_status: str = ""
    processing_warnings: List[str] = Field(default_factory=list)
    source_filename: str = ""
    uploaded_at: str = ""
    updated_at: str = ""


class PublicBookOut(BaseModel):
    """Safe public book shape for controlled-launch catalog/detail routes."""

    model_config = ConfigDict(extra="ignore")

    id: str = ""
    slug: str = ""
    title: str = ""
    subtitle: str = ""
    author: str = ""
    category_slug: str = ""
    short_description: str = ""
    description: str = ""
    cover_url: str = ""
    cover_image_url: str = ""
    thumbnail_url: str = ""
    blur_placeholder: str = ""
    dominant_color: str = ""
    back_cover_url: str = ""
    back_cover_image_url: str = ""
    back_cover_thumbnail_url: str = ""
    back_cover_blur_placeholder: str = ""
    back_cover_dominant_color: str = ""
    estimated_reading_time: str = ""
    formats: List[str] = Field(default_factory=list)
    benefits: List[str] = Field(default_factory=list)
    who_for: List[str] = Field(default_factory=list)
    learnings: List[str] = Field(default_factory=list)
    about_author: str = ""
    chapters: List[PublicChapterOut] = Field(default_factory=list)
    is_published: bool = False
    created_at: str = ""
    updated_at: str = ""
    publication_status: str = ""
    launch_status: str = ""
    reader_enabled: bool = False
    preview_enabled: bool = False
    audio_enabled: bool = False
    audiobook_enabled: bool = False
    public_route: str = ""
    reader_url: str = ""
    preview_url: str = ""
    audio_url: str = ""
    audio_status: str = ""
    cta_label: str = ""
    secondary_cta_label: str = ""
    public_json_ld_enabled: bool = False
    source_note: str = ""
    rights_note: str = ""


class BookIn(BaseModel):
    title: str
    subtitle: str = ""
    author: str = "The Earnalism"
    category_slug: str
    short_description: str = ""
    description: str = ""
    cover_image_url: str = ""
    back_cover_image_url: str = ""
    estimated_reading_time: str = ""
    price_paperback: str = ""
    price_ebook: str = ""
    buy_url: str = ""
    formats: List[str] = Field(default_factory=lambda: ["Paperback", "Ebook"])
    benefits: List[str] = Field(default_factory=list)
    who_for: List[str] = Field(default_factory=list)
    learnings: List[str] = Field(default_factory=list)
    about_author: str = ""
    rights_metadata: Dict[str, Any] = Field(default_factory=dict)
    source_url: str = ""
    source_name: str = ""
    source_license: str = ""
    source_hash: str = ""
    content_hash: str = ""
    provenance_hash: str = ""
    rights_basis: str = ""
    rights_tier: str = ""
    verification_status: str = ""
    qa_status: str = ""
    approved_to_publish: bool = False
    publication_status: str = ""
    audiobook_enabled: bool = False
    generate_audiobook: bool = False
    readerStatus: str = "ready_for_editorial_review"
    publicationStatus: str = "draft"
    isPublic: bool = False
    isLive: bool = False
    showInPublicLibrary: bool = False
    showInHomepage: bool = False
    allowPublicReading: bool = False
    allowCheckout: bool = False
    allowPayment: bool = False
    is_published: bool = False
    slug: Optional[str] = None


class BookAudiobookIn(BaseModel):
    audiobook_enabled: bool = True
    generate_audiobook: bool = True
    audiobook_provider: str = ""
    audiobook_voice: str = ""
    audio_asset_slug: str = ""
    audiobook_assets: Dict[str, str] = Field(default_factory=dict)
    audiobook_size: int = 0
    audiobook_duration_ms: int = 0


ALLOWED_AUDIO_ASSET_KEYS = {"mp3", "timestamps", "vtt", "chapters", "meta", "manifest"}


def _safe_audiobook_assets(value: Optional[Dict[str, Any]]) -> Dict[str, str]:
    assets: Dict[str, str] = {}
    for key, raw_url in (value or {}).items():
        normalized_key = str(key or "").strip().lower()
        if normalized_key not in ALLOWED_AUDIO_ASSET_KEYS:
            continue
        url = str(raw_url or "").strip()
        if not url or len(url) > 2000:
            continue
        if not (url.startswith("https://") or url.startswith("/audio/")):
            continue
        assets[normalized_key] = url
    return assets

class BlogPost(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    slug: str
    title: str
    excerpt: str = ""
    content: str = ""
    category: str = "Reflections"
    cover_image_url: str = ""
    author: str = "The Earnalism"
    pull_quote: str = ""
    is_published: bool = True
    created_at: str = Field(default_factory=now_iso)

class BlogPostIn(BaseModel):
    title: str
    excerpt: str = ""
    content: str = ""
    category: str = "Reflections"
    cover_image_url: str = ""
    author: str = "The Earnalism"
    pull_quote: str = ""
    is_published: bool = True
    slug: Optional[str] = None

class NewsletterIn(BaseModel):
    name: str
    email: EmailStr

class ContactIn(BaseModel):
    name: str
    email: EmailStr
    subject: str = ""
    message: str

class SocialIn(BaseModel):
    instagram: str = ""
    facebook: str = ""
    youtube: str = ""
    linkedin: str = ""
    twitter: str = ""

class BrandIn(BaseModel):
    logo_url: str = ""
    og_image_url: str = ""

class FeaturedIn(BaseModel):
    book_slug: str

class ContactStatusIn(BaseModel):
    status: str  # one of: new, read, responded

VALID_CONTACT_STATUSES = {"new", "read", "responded"}


# ---------- Reader User / Wallet / Session models ----------
class UserSignupIn(BaseModel):
    name: str
    email: EmailStr
    password: str

class UserLoginIn(BaseModel):
    email: EmailStr
    password: str

class UserOut(BaseModel):
    id: str
    name: str
    email: EmailStr
    role: str = "user"
    reading_seconds_balance: int = 0
    status: str = "active"
    auth_provider: str = "email"
    created_at: str

class UserAuthOut(BaseModel):
    token: str
    user: UserOut

class GoogleAuthIn(BaseModel):
    credential: str

class OTPRequestIn(BaseModel):
    mobile: str

class OTPVerifyIn(BaseModel):
    mobile: str
    otp: str

class WalletAdjustIn(BaseModel):
    minutes: int  # may be negative; converted to seconds server-side
    reason: str = ""

class WalletRefundApproveIn(BaseModel):
    candidate_ids: List[str] = Field(default_factory=list)
    note: str = ""

class WalletTransactionOut(BaseModel):
    id: str
    user_id: str
    type: str  # "credit" | "debit" | "consume"
    seconds: int
    reason: str
    created_at: str
    actor: str = "system"  # "admin" | "system" | "user"
    session_id: str = ""

class ReaderSessionStartIn(BaseModel):
    session_id: Optional[str] = None
    book_id: Optional[str] = None
    book_slug: Optional[str] = None
    chapter_id: Optional[str] = None

class ReaderHeartbeatIn(BaseModel):
    session_id: str
    visible: bool = True
    idle: bool = False
    chapter_id: Optional[str] = None

class ReaderSessionEndIn(BaseModel):
    session_id: str

class ReadingPulseIn(BaseModel):
    session_id: str
    visible: bool = True
    idle: bool = False

class ReaderCompletionIn(BaseModel):
    book_slug: str
    chapter_id: str
    chapter_title: str = ""
    progress: int = 100

class ReaderMetricIn(BaseModel):
    event: str = "reader_metric"
    session_id: str = ""
    book_slug: str = ""
    chapter_id: str = ""
    route: str = ""
    timings: Dict[str, Any] = Field(default_factory=dict)
    metrics: Dict[str, Any] = Field(default_factory=dict)
    tags: Dict[str, Any] = Field(default_factory=dict)

class AnalyticsEventIn(BaseModel):
    event: str = ""
    event_name: str = ""
    route: str = ""
    book_slug: str = ""
    anonymous_session_id: str = ""
    metadata: Dict[str, Any] = Field(default_factory=dict)

class SecureReaderEventIn(BaseModel):
    session_id: str
    event_type: str
    book_slug: str = ""
    chapter_id: str = ""
    access_token_fingerprint: str = ""
    counts: Dict[str, int] = Field(default_factory=dict)
    metadata: Dict[str, Any] = Field(default_factory=dict)

class UserStatusIn(BaseModel):
    status: str  # "active" | "blocked"


# ---------- Payments / Razorpay top-up models ----------
class PackOut(BaseModel):
    id: str
    label: str
    minutes: int
    price_inr: int
    amount_paise: int
    note: str


class TopUpCreateIn(BaseModel):
    pack_id: str


class TopUpCreateOut(BaseModel):
    intent_id: str
    razorpay_order_id: str
    key_id: str
    amount: int  # in paise
    currency: str = "INR"
    name: str
    description: str
    pack: PackOut
    prefill: dict


class PaymentVerifyIn(BaseModel):
    razorpay_order_id: str
    razorpay_payment_id: str
    razorpay_signature: str


class PaymentReconcileIn(BaseModel):
    note: str = ""


def _user_public(doc: dict) -> dict:
    return {
        "id": doc["id"],
        "name": doc.get("name", ""),
        "email": doc["email"],
        "role": "user",
        "reading_seconds_balance": int(doc.get("reading_seconds_balance", 0)),
        "status": doc.get("status", "active"),
        "auth_provider": doc.get("auth_provider", "email"),
        "created_at": doc.get("created_at", now_iso()),
    }


async def _ledger_balance(user_id: str) -> int:
    rows = await db.wallet_ledger.aggregate([
        {"$match": {"user_id": user_id}},
        {"$group": {"_id": "$user_id", "credit": {"$sum": "$credit"}, "debit": {"$sum": "$debit"}}},
    ]).to_list(1)
    if not rows:
        return 0
    return int(rows[0].get("credit", 0)) - int(rows[0].get("debit", 0))


async def _flag_wallet_divergence(user_id: str, stored_balance: int) -> None:
    computed = await _ledger_balance(user_id)
    if computed == int(stored_balance or 0):
        return
    await db.wallet_integrity_alerts.insert_one({
        "id": str(uuid.uuid4()),
        "user_id": user_id,
        "stored_balance": int(stored_balance or 0),
        "ledger_balance": computed,
        "delta": int(stored_balance or 0) - computed,
        "created_at": now_iso(),
        "status": "open",
    })


async def _record_wallet_ledger(
    *,
    user_id: str,
    session_id: str = "",
    action: str,
    seconds_delta: int,
    reason: str,
    actor: str = "system",
    balance_after: Optional[int] = None,
    source_transaction_id: Optional[str] = None,
    extra: Optional[dict] = None,
) -> dict:
    seconds = int(seconds_delta)
    tx_type = "credit" if seconds >= 0 else ("debit" if action == "admin_adjustment" else "consume")
    tx = {
        "id": source_transaction_id or str(uuid.uuid4()),
        "user_id": user_id,
        "type": tx_type,
        "seconds": seconds,
        "reason": reason,
        "created_at": now_iso(),
        "actor": actor,
        "session_id": session_id,
    }
    if extra:
        tx.update(extra)
    if not source_transaction_id:
        await db.wallet_transactions.insert_one(tx)

    if balance_after is None:
        fresh = await db.users.find_one({"id": user_id}, {"_id": 0, "reading_seconds_balance": 1, "wallet_seconds": 1})
        balance_after = int((fresh or {}).get("reading_seconds_balance", (fresh or {}).get("wallet_seconds", 0)) or 0)

    ledger = {
        "id": str(uuid.uuid4()),
        "user_id": user_id,
        "session_id": session_id,
        "action": action,
        "debit": max(0, -seconds),
        "credit": max(0, seconds),
        "timestamp": now_iso(),
        "reason": reason,
        "actor": actor,
        "balance_after": int(balance_after or 0),
        "source_transaction_id": tx["id"],
    }
    if extra:
        ledger["metadata"] = {k: str(v)[:240] for k, v in extra.items()}
    await db.wallet_ledger.update_one(
        {"source_transaction_id": tx["id"]},
        {"$setOnInsert": ledger},
        upsert=True,
    )
    await _flag_wallet_divergence(user_id, int(balance_after or 0))
    await _invalidate_user_cache(user_id)
    await _set_user_wallet_cache(user_id, int(balance_after or 0))
    return tx


async def _migrate_wallet_transactions_to_ledger() -> None:
    rows = await db.wallet_transactions.find({}, {"_id": 0}).sort("created_at", 1).to_list(10000)
    running: Dict[str, int] = {}
    for tx in rows:
        tx_id = tx.get("id")
        if not tx_id:
            continue
        user_id = tx.get("user_id")
        if not user_id:
            continue
        seconds = int(tx.get("seconds", 0) or 0)
        running[user_id] = running.get(user_id, 0) + seconds
        await db.wallet_ledger.update_one(
            {"source_transaction_id": tx_id},
            {
                "$setOnInsert": {
                    "id": str(uuid.uuid4()),
                    "user_id": user_id,
                    "session_id": tx.get("session_id", ""),
                    "action": tx.get("type", "wallet_adjustment"),
                    "debit": max(0, -seconds),
                    "credit": max(0, seconds),
                    "timestamp": tx.get("created_at", now_iso()),
                    "reason": tx.get("reason", ""),
                    "actor": tx.get("actor", "system"),
                    "balance_after": running[user_id],
                    "source_transaction_id": tx_id,
                    "migrated": True,
                }
            },
            upsert=True,
        )


REFUND_DUPLICATE_PULSE_SECONDS = 10


def _wallet_refund_candidate_id(user_id: str, tx_id: str, issue: str, refundable_seconds: int) -> str:
    raw = f"{user_id}:{tx_id}:{issue}:{int(refundable_seconds)}"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:24]


def _wallet_refund_candidates(transactions: List[dict], refunded_candidate_ids: Optional[set[str]] = None) -> List[dict]:
    """Find high-confidence reader billing candidates that are safe to refund.

    The rules intentionally stay conservative: a single reader pulse should
    never consume more than HEARTBEAT_TICK_SECONDS, and duplicate pulses from
    the same session within a few seconds are treated as retry/concurrency
    artifacts. Ambiguous activity is left out for manual admin adjustment.
    """
    refunded = refunded_candidate_ids or set()
    rows = sorted(
        (tx for tx in transactions if tx.get("type") == "consume" and int(tx.get("seconds", 0) or 0) < 0),
        key=lambda tx: _as_utc_dt(tx.get("created_at")) or datetime.min.replace(tzinfo=timezone.utc),
    )
    candidates: List[dict] = []
    previous_by_session: Dict[str, dict] = {}
    for tx in rows:
        tx_id = tx.get("id") or ""
        user_id = tx.get("user_id") or ""
        session_id = tx.get("session_id") or ""
        charged = abs(int(tx.get("seconds", 0) or 0))
        if not tx_id or not user_id or charged <= 0:
            continue

        issue = ""
        refundable = 0
        evidence = ""
        confidence = "high"

        if charged > HEARTBEAT_TICK_SECONDS:
            issue = "stale_gap_overcharge"
            refundable = charged - HEARTBEAT_TICK_SECONDS
            evidence = (
                f"Single debit consumed {charged}s. The protected reader policy caps a visible, active pulse "
                f"at {HEARTBEAT_TICK_SECONDS}s, so only one pulse is retained."
            )
        elif session_id:
            current_at = _as_utc_dt(tx.get("created_at"))
            previous = previous_by_session.get(session_id)
            previous_at = _as_utc_dt((previous or {}).get("created_at"))
            if current_at and previous_at:
                gap = max(0, int((current_at - previous_at).total_seconds()))
                if gap <= REFUND_DUPLICATE_PULSE_SECONDS:
                    issue = "duplicate_pulse"
                    refundable = min(charged, HEARTBEAT_TICK_SECONDS)
                    evidence = (
                        f"Two debits in the same reading session landed {gap}s apart. Reader pulses should "
                        f"not settle that close together."
                    )

        if session_id:
            previous_by_session[session_id] = tx

        if refundable <= 0 or not issue:
            continue

        candidate_id = _wallet_refund_candidate_id(user_id, tx_id, issue, refundable)
        if candidate_id in refunded:
            continue
        candidates.append({
            "candidate_id": candidate_id,
            "source_transaction_id": tx_id,
            "session_id": session_id,
            "created_at": tx.get("created_at", ""),
            "charged_seconds": charged,
            "refundable_seconds": refundable,
            "issue": issue,
            "reason": tx.get("reason", ""),
            "evidence": evidence,
            "confidence": confidence,
        })
    return candidates


def _free_preview_chapter_ids(book: dict) -> set[str]:
    """Return only chapters with explicit preview approval evidence."""
    return set(explicit_preview_chapter_ids(book))


def _is_free_preview_chapter(book: dict, chapter_id: Optional[str]) -> bool:
    if not chapter_id:
        return False
    return chapter_id in _free_preview_chapter_ids(book)


def _strip_paid_chapter_content(book: dict) -> dict:
    """Return a shallow copy of `book` with non-preview chapter `content`
    blanked. Explicit preview chapters keep content; every other chapter is
    metadata-only. Used for PUBLIC `/api/books/*` endpoints so chapter
    bodies are never leaked to guests via the catalog API.
    """
    if not book:
        return book
    chapters = book.get("chapters") or []
    if not chapters:
        out = dict(book)
        out.pop("rights_metadata", None)
        return out
    preview_ids = _free_preview_chapter_ids(book)
    masked = []
    for c in chapters:
        c2 = dict(c)
        if c.get("id") not in preview_ids:
            c2["content"] = ""
        masked.append(c2)
    out = dict(book)
    out["chapters"] = masked
    out.pop("rights_metadata", None)
    return out


def _strip_all_chapter_content(book: dict) -> dict:
    if not book:
        return book
    chapters = book.get("chapters") or []
    if not chapters:
        out = dict(book)
        out.pop("rights_metadata", None)
        return out
    out = dict(book)
    out["chapters"] = [{**c, "content": ""} for c in chapters]
    out.pop("rights_metadata", None)
    return out


# ---------- App ----------
async def initialize_database_indexes() -> None:
    await db.users.create_index("email", unique=True)
    await db.users.create_index("id", sparse=True)
    await db.users.create_index("mobile", sparse=True)
    await db.users.create_index([("role", 1), ("status", 1), ("created_at", -1)])

    await db.books.create_index("slug", unique=True)
    await db.books.create_index([("slug", 1), ("is_published", 1)])
    await db.books.create_index([("is_published", 1), ("created_at", -1)])
    await db.books.create_index([("category_slug", 1), ("is_published", 1), ("created_at", -1)])
    await db.books.create_index([("rights_metadata.verification_status", 1), ("rights_metadata.rights_tier", 1)])
    await db.books.create_index([("rights_metadata.publication_region", 1), ("is_published", 1)])

    await db.categories.create_index("slug", unique=True)
    await db.categories.create_index([("order", 1), ("slug", 1)])

    await db.blog_posts.create_index("slug", unique=True)
    await db.blog_posts.create_index([("is_published", 1), ("created_at", -1)])

    await db.newsletter.create_index("email", unique=True)
    await db.newsletter.create_index([("created_at", -1)])

    await db.contacts.create_index([("status", 1), ("created_at", -1)])
    await db.contacts.create_index([("created_at", -1)])

    await db.settings.create_index("key", unique=True)

    await db.wallet_transactions.create_index([("user_id", 1), ("created_at", -1)])

    await db.wallet_ledger.create_index([("user_id", 1), ("timestamp", -1)])
    await db.wallet_ledger.create_index([("session_id", 1), ("timestamp", -1)])
    await db.wallet_ledger.create_index(
        "source_transaction_id",
        unique=True,
        partialFilterExpression={"source_transaction_id": {"$type": "string"}},
    )
    await db.wallet_integrity_alerts.create_index([("user_id", 1), ("created_at", -1)])
    await db.wallet_refunds.create_index("candidate_id", unique=True)
    await db.wallet_refunds.create_index([("user_id", 1), ("created_at", -1)])
    await db.wallet_refunds.create_index([("status", 1), ("created_at", -1)])

    await db.credit_log.create_index([("user_id", 1), ("timestamp", -1)])
    await db.credit_log.create_index([("upload_id", 1), ("timestamp", -1)])
    await db.admin_upload_audit.create_index([("admin_user_id", 1), ("created_at", -1)])

    await db.user_sessions.create_index("id", unique=True)
    await db.user_sessions.create_index([("user_id", 1), ("status", 1), ("created_at", -1)])
    await db.user_sessions.create_index(
        "refresh_token_hash",
        unique=True,
        partialFilterExpression={"refresh_token_hash": {"$type": "string"}},
    )

    await db.reading_sessions.create_index("id", sparse=True)
    await db.reading_sessions.create_index([("user_id", 1), ("started_at", -1)])
    await db.reading_sessions.create_index([("user_id", 1), ("status", 1), ("started_at", -1)])
    await db.reading_sessions.create_index("status")

    await db.topup_intents.create_index("id", sparse=True)
    await db.topup_intents.create_index(
        "razorpay_order_id",
        unique=True,
        partialFilterExpression={"razorpay_order_id": {"$type": "string"}},
    )
    await db.topup_intents.create_index(
        "razorpay_payment_id",
        unique=True,
        partialFilterExpression={"razorpay_payment_id": {"$type": "string"}},
    )
    await db.topup_intents.create_index([("user_id", 1), ("created_at", -1)])
    await db.topup_intents.create_index([("status", 1), ("created_at", -1)])
    await db.topup_intents.create_index("status")

    await db.payment_webhook_events.create_index("event_id", unique=True, sparse=True)
    await db.payment_webhook_events.create_index([("created_at", -1)])

    await db.analytics_events.create_index([("event", 1), ("created_at", -1)])
    await db.analytics_events.create_index([("route", 1), ("created_at", -1)])
    await db.analytics_events.create_index([("anonymous_session_id", 1), ("created_at", -1)])

    await db.reader_security_events.create_index([("event_type", 1), ("created_at", -1)])
    await db.reader_security_events.create_index([("session_id", 1), ("created_at", -1)])
    await db.reader_security_events.create_index([("user_id", 1), ("created_at", -1)])
    await db.reader_experience_events.create_index([("event", 1), ("created_at", -1)])
    await db.reader_experience_events.create_index([("book_slug", 1), ("chapter_id", 1), ("created_at", -1)])

    await db.reader_completions.create_index([("user_id", 1), ("completed_on", -1)])
    await db.reader_completions.create_index(
        [("user_id", 1), ("book_slug", 1), ("chapter_id", 1), ("completed_on", 1)],
        unique=True,
    )

    await db.reward_claims.create_index([("user_id", 1), ("reward_key", 1)], unique=True)


def _mark_shutdown_draining() -> None:
    _shutdown_state["draining"] = True
    logger.warning(
        "SIGTERM received; marking replica as draining with %s in-flight request(s).",
        _shutdown_state.get("inflight", 0),
    )


def _install_sigterm_drain_marker() -> None:
    previous_handler = signal.getsignal(signal.SIGTERM)

    def _handle_sigterm(signum, frame):
        _mark_shutdown_draining()
        if callable(previous_handler):
            previous_handler(signum, frame)
        elif previous_handler == signal.SIG_DFL:
            raise KeyboardInterrupt

    signal.signal(signal.SIGTERM, _handle_sigterm)


async def _run_startup_database_maintenance_once() -> None:
    await initialize_database_indexes()
    await _migrate_wallet_transactions_to_ledger()

    # admin (seed only if absent; do NOT overwrite password so admin can change it via UI)
    existing = await db.users.find_one({"email": ADMIN_EMAIL.lower()})
    if not existing:
        if not ADMIN_PASSWORD:
            raise RuntimeError("ADMIN_PASSWORD is required to seed the initial admin user")
        await db.users.insert_one({
            "id": str(uuid.uuid4()),
            "email": ADMIN_EMAIL.lower(),
            "password_hash": hash_password(ADMIN_PASSWORD),
            "name": "Admin",
            "role": "admin",
            "created_at": now_iso(),
        })
        logger.info("Seeded admin user")

    # Optional non-production seed reader (idempotent; disabled unless explicitly enabled).
    test_user_email = SEED_TEST_READER_EMAIL.lower()
    if SEED_TEST_READER and not await db.users.find_one({"email": test_user_email}):
        if not SEED_TEST_READER_PASSWORD:
            raise RuntimeError("SEED_TEST_READER_PASSWORD is required when SEED_TEST_READER=true")
        await db.users.insert_one({
            "id": str(uuid.uuid4()),
            "email": test_user_email,
            "password_hash": hash_password(SEED_TEST_READER_PASSWORD),
            "name": "Sample Reader",
            "role": "user",
            "auth_provider": "email",
            "reading_seconds_balance": 0,
            "status": "active",
            "created_at": now_iso(),
        })
        logger.info("Seeded sample reader user")

    # Backfill new fields on legacy user documents (role=user)
    await db.users.update_many(
        {"role": "user", "reading_seconds_balance": {"$exists": False}},
        {"$set": {"reading_seconds_balance": 0}},
    )
    await db.users.update_many(
        {"role": "user", "status": {"$exists": False}},
        {"$set": {"status": "active"}},
    )
    await db.users.update_many(
        {"role": "user", "auth_provider": {"$exists": False}},
        {"$set": {"auth_provider": "email"}},
    )

    # Canonical shelf taxonomy. Keep older imported/admin book records attached
    # to their closest current shelf before refreshing category documents.
    for old_slug, new_slug in LEGACY_CATEGORY_SLUG_MAP.items():
        await db.books.update_many(
            {"category_slug": old_slug},
            {"$set": {"category_slug": new_slug}},
        )
    await db.books.update_many(
        {"$or": [
            {"category_slug": {"$exists": False}},
            {"category_slug": ""},
            {"category_slug": {"$nin": list(CANONICAL_CATEGORY_SLUGS)}},
        ]},
        {"$set": {"category_slug": DEFAULT_CATEGORY_SLUG}},
    )
    await db.categories.delete_many({"slug": {"$nin": list(CANONICAL_CATEGORY_SLUGS)}})
    for c in SEED_CATEGORIES:
        await db.categories.update_one(
            {"slug": c["slug"]},
            {"$set": c, "$setOnInsert": {"id": str(uuid.uuid4())}},
            upsert=True,
        )

    # Retired sample books used to be seeded on startup, which made them
    # reappear after admins deleted them. Keep them removed permanently.
    retired_result = await db.books.delete_many({"slug": {"$in": list(RETIRED_SEED_BOOK_SLUGS)}})
    if retired_result.deleted_count:
        logger.info("Removed %s retired seed book(s)", retired_result.deleted_count)
    await db.settings.update_one(
        {"key": "featured_book", "book_slug": {"$in": list(RETIRED_SEED_BOOK_SLUGS)}},
        {"$unset": {"book_slug": ""}},
    )

    # featured setting
    await db.settings.update_one(
        {"key": "featured_book"},
        {"$setOnInsert": {"key": "featured_book"}},
        upsert=True,
    )

    # blog posts
    for p in SEED_POSTS:
        if not await db.blog_posts.find_one({"slug": p["slug"]}):
            post = BlogPost(**p)
            await db.blog_posts.insert_one(post.model_dump())

    # social settings (seed empty config so GET /api/settings/social always works)
    await db.settings.update_one(
        {"key": "social"},
        {"$setOnInsert": {"key": "social", "instagram": "", "facebook": "", "youtube": "", "linkedin": "", "twitter": ""}},
        upsert=True,
    )

    # brand settings (logo + OG image; seeded empty so the public endpoint always works)
    await db.settings.update_one(
        {"key": "brand"},
        {"$setOnInsert": {"key": "brand", "logo_url": "", "og_image_url": ""}},
        upsert=True,
    )

    # backfill contact.status for older entries that didn't have the field
    await db.contacts.update_many({"status": {"$exists": False}}, {"$set": {"status": "new"}})

    # backfill digital-library fields on any pre-existing books
    await db.books.update_many(
        {"author": {"$exists": False}},
        {"$set": {"author": "The Earnalism"}},
    )
    await db.books.update_many(
        {"estimated_reading_time": {"$exists": False}},
        {"$set": {"estimated_reading_time": ""}},
    )
    await db.books.update_many(
        {"chapters": {"$exists": False}},
        {"$set": {"chapters": []}},
    )
    # Draft-safe by default: only explicit controlled-launch slugs inherit
    # legacy live status when older rows are missing is_published.
    await db.books.update_many(
        {
            "is_published": {"$exists": False},
            "slug": {"$in": list(CONTROLLED_LIVE_BOOK_SLUGS)},
        },
        {"$set": {"is_published": True}},
    )
    await db.books.update_many(
        {
            "is_published": {"$exists": False},
            "slug": {"$nin": list(CONTROLLED_LIVE_BOOK_SLUGS)},
        },
        {"$set": {"is_published": False}},
    )

    logger.info("Startup seeding complete")


async def _run_startup_database_maintenance_with_retry() -> None:
    delay = 1.0
    for attempt in range(1, STARTUP_DB_MAINTENANCE_ATTEMPTS + 1):
        try:
            await _run_startup_database_maintenance_once()
            return
        except (AutoReconnect, ServerSelectionTimeoutError) as exc:
            if attempt >= STARTUP_DB_MAINTENANCE_ATTEMPTS:
                logger.exception("Startup database maintenance failed after %s attempt(s).", attempt)
                raise
            logger.warning(
                "Startup database maintenance attempt %s/%s hit a transient Mongo error: %s. Retrying in %.1fs.",
                attempt,
                STARTUP_DB_MAINTENANCE_ATTEMPTS,
                exc,
                delay,
            )
            await asyncio.sleep(delay)
            delay = min(delay * 2, 8.0)


async def _release_startup_maintenance_lock(lock_key: str, token: str) -> None:
    try:
        current = _redis_text(await _redis_client.get(lock_key))
        if current == token:
            await _redis_client.delete(lock_key)
    except Exception:
        logger.warning("Failed to release startup maintenance lock cleanly.", exc_info=True)


async def _run_startup_database_maintenance() -> None:
    if not _redis_state_enabled():
        await _run_startup_database_maintenance_with_retry()
        return

    lock_key = _redis_key("startup-maintenance", STARTUP_MAINTENANCE_VERSION, "lock")
    done_key = _redis_key("startup-maintenance", STARTUP_MAINTENANCE_VERSION, "done")
    if await _redis_client.get(done_key):
        logger.info("Startup database maintenance already completed for %s.", STARTUP_MAINTENANCE_VERSION)
        return

    token = str(uuid.uuid4())
    deadline = time.monotonic() + STARTUP_MAINTENANCE_WAIT_SECONDS
    while True:
        acquired = await _redis_client.set(
            lock_key,
            token,
            ex=STARTUP_MAINTENANCE_LOCK_SECONDS,
            nx=True,
        )
        if acquired:
            logger.info("Acquired startup database maintenance lock.")
            try:
                await _run_startup_database_maintenance_with_retry()
                await _redis_client.setex(done_key, STARTUP_MAINTENANCE_DONE_TTL_SECONDS, now_iso())
            finally:
                await _release_startup_maintenance_lock(lock_key, token)
            return

        if await _redis_client.get(done_key):
            logger.info("Startup database maintenance completed by another replica.")
            return
        if time.monotonic() >= deadline:
            logger.warning(
                "Startup database maintenance lock is still held after %ss; starting replica without rerunning maintenance.",
                STARTUP_MAINTENANCE_WAIT_SECONDS,
            )
            return
        await asyncio.sleep(1)


@asynccontextmanager
async def lifespan(_app: FastAPI):
    # ----- startup -----
    _install_sigterm_drain_marker()
    _log_cost_control_startup()
    await initialize_replica_state_backends()
    if ENABLE_STARTUP_DB_MAINTENANCE:
        await _run_startup_database_maintenance()
    else:
        logger.info("Startup database maintenance is disabled by cost-control mode.")

    yield

    # ----- shutdown -----
    _mark_shutdown_draining()
    drain_deadline = time.monotonic() + 15
    while int(_shutdown_state.get("inflight", 0)) > 0 and time.monotonic() < drain_deadline:
        await asyncio.sleep(0.1)
    await close_replica_state_backends()
    client.close()


app = FastAPI(
    title="The Earnalism API",
    lifespan=lifespan,
    default_response_class=UTF8JSONResponse,
    docs_url=None if ENVIRONMENT == "production" else "/docs",
    redoc_url=None if ENVIRONMENT == "production" else "/redoc",
)
if os.environ.get("JUDOSCALE_URL", "").strip():
    try:
        from judoscale.asgi.middleware import FastAPIRequestQueueTimeMiddleware  # type: ignore
    except Exception as exc:
        raise RuntimeError("JUDOSCALE_URL is set but Judoscale ASGI middleware is unavailable.") from exc
    app.add_middleware(FastAPIRequestQueueTimeMiddleware)
    logger.info("Judoscale FastAPI request queue middleware enabled.")
api = APIRouter(prefix="/api")

APPROVED_LAUNCH_ANALYTICS_EVENTS = {
    "homepage_view",
    "first_time_site_tour_shown",
    "first_time_site_tour_completed",
    "first_time_site_tour_skipped",
    "hero_read_chapter_free_click",
    "dracula_book_page_view",
    "start_dracula_click",
    "reader_opened",
    "reader_locked_state",
    "reader_low_balance_state",
    "pricing_page_view",
    "reading_pack_selected",
    "checkout_started",
    "payment_success_return",
    "payment_failed_or_cancelled",
    "wallet_credited_visible",
    "continue_reading_click",
    "return_resume_reading_click",
    "core_web_vital",
}

LAUNCH_MONITOR_FUNNEL_EVENTS = [
    "homepage_view",
    "first_time_site_tour_shown",
    "first_time_site_tour_completed",
    "first_time_site_tour_skipped",
    "hero_read_chapter_free_click",
    "dracula_book_page_view",
    "start_dracula_click",
    "reader_opened",
    "reader_locked_state",
    "reader_low_balance_state",
    "pricing_page_view",
    "reading_pack_selected",
    "checkout_started",
    "payment_success_return",
    "payment_failed_or_cancelled",
    "wallet_credited_visible",
    "continue_reading_click",
    "return_resume_reading_click",
]

ANALYTICS_BLOCKED_KEY_RE = re.compile(
    r"(email|phone|customer|payment|order|razorpay|signature|token|secret|password|authorization|invoice|billing|card|upi|bank|address|name)",
    re.IGNORECASE,
)
ANALYTICS_ALLOWED_SAFE_KEYS = {"payment_mode"}
ANALYTICS_UNSAFE_VALUE_PATTERNS = [
    re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.IGNORECASE),
    re.compile(r"\b(?:\+?91[-\s]?)?[6-9]\d{9}\b"),
    re.compile(r"\b(?:rzp|pay|order|cust)_[A-Za-z0-9_-]{8,}\b", re.IGNORECASE),
    re.compile(r"\bsk[-_][A-Za-z0-9_-]{12,}\b", re.IGNORECASE),
    re.compile(r"\bBearer\s+[A-Za-z0-9._-]{12,}\b", re.IGNORECASE),
    re.compile(r"\b(?:card|upi|bank|ifsc|account)\b", re.IGNORECASE),
]


def _safe_analytics_event_name(value: str) -> str:
    event = re.sub(r"[^a-zA-Z0-9_.:-]", "_", str(value or "").strip())[:80]
    return event if event in APPROVED_LAUNCH_ANALYTICS_EVENTS else ""


def _analytics_metadata_violations(metadata: Dict[str, Any]) -> list[str]:
    if not isinstance(metadata, dict):
        return ["metadata must be an object"]
    violations: list[str] = []
    for raw_key, raw_value in list(metadata.items())[:50]:
        key = str(raw_key or "")
        if not key or (ANALYTICS_BLOCKED_KEY_RE.search(key) and key.lower() not in ANALYTICS_ALLOWED_SAFE_KEYS):
            violations.append(f"blocked metadata field: {key or '<empty>'}")
            continue
        if isinstance(raw_value, (dict, list, tuple, set)):
            violations.append(f"nested metadata is not allowed: {key}")
            continue
        value = "" if raw_value is None else str(raw_value)
        if len(value) > 500:
            violations.append(f"metadata value too long: {key}")
            continue
        if any(pattern.search(value) for pattern in ANALYTICS_UNSAFE_VALUE_PATTERNS):
            violations.append(f"unsafe metadata value: {key}")
    return violations


def _safe_analytics_metadata(metadata: Dict[str, Any]) -> Dict[str, Any]:
    """Keep conversion analytics useful while avoiding secrets or large payloads."""
    if not isinstance(metadata, dict):
        return {}

    safe: Dict[str, Any] = {}
    for raw_key, raw_value in list(metadata.items())[:20]:
        key = re.sub(r"[^a-zA-Z0-9_.:-]", "_", str(raw_key))[:60]
        if not key:
            continue
        if isinstance(raw_value, bool) or raw_value is None:
            safe[key] = raw_value
        elif isinstance(raw_value, (int, float)):
            safe[key] = raw_value
        elif isinstance(raw_value, str):
            safe[key] = raw_value[:240]
        else:
            safe[key] = str(raw_value)[:240]
    return safe


def _safe_analytics_route(value: str) -> str:
    route = str(value or "").strip()
    if not route:
        return ""
    parsed = urlparse(route)
    if parsed.scheme or parsed.netloc:
        route = parsed.path or "/"
    return re.sub(r"[^a-zA-Z0-9_./?=&:-]", "_", route)[:180]


def _safe_analytics_session_id(value: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_.:-]", "_", str(value or "").strip())[:80]


def _safe_analytics_book_slug(value: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_.:-]", "_", str(value or "").strip())[:80]


def _analytics_event_document(payload: AnalyticsEventIn, request: Request, principal: Optional[dict]) -> dict:
    event = _safe_analytics_event_name(payload.event_name or payload.event)
    if not event:
        raise HTTPException(status_code=400, detail="Unknown launch analytics event")

    metadata = dict(payload.metadata or {})
    route = _safe_analytics_route(payload.route or metadata.get("route") or metadata.get("path") or "")
    if route:
        metadata["route"] = route
    book_slug = _safe_analytics_book_slug(payload.book_slug or metadata.get("book_slug") or metadata.get("book") or "")
    if book_slug:
        metadata["book_slug"] = book_slug
    session_id = _safe_analytics_session_id(payload.anonymous_session_id or metadata.get("anonymous_session_id") or "")

    violations = _analytics_metadata_violations(metadata)
    if violations:
        raise HTTPException(status_code=400, detail={"message": "Unsafe analytics metadata", "violations": violations[:5]})

    return {
        "id": str(uuid.uuid4()),
        "event": event,
        "metadata": _safe_analytics_metadata(metadata),
        "route": route,
        "book_slug": book_slug,
        "anonymous_session_id": session_id,
        "principal_role": principal.get("role") if principal else "guest",
        "principal_id": principal.get("id") if principal else "",
        "path": route or str(request.headers.get("referer", ""))[:240],
        "user_agent": str(request.headers.get("user-agent", ""))[:180],
        "created_at": now_iso(),
    }


def _launch_monitor_cutoff(hours: int) -> str:
    return (datetime.now(timezone.utc) - timedelta(hours=hours)).isoformat()


def _launch_monitor_today_cutoff() -> str:
    now = datetime.now(timezone.utc)
    return datetime(now.year, now.month, now.day, tzinfo=timezone.utc).isoformat()


async def _count_since(collection_name: str, time_field: str, cutoff: str, extra: Optional[dict] = None) -> int:
    query = dict(extra or {})
    query[time_field] = {"$gte": cutoff}
    try:
        return int(await getattr(db, collection_name).count_documents(query))
    except Exception:
        logger.warning("Launch monitor count failed for %s", collection_name, exc_info=True)
        return 0


async def _group_counts_since(
    collection_name: str,
    field: str,
    time_field: str,
    cutoff: str,
    extra: Optional[dict] = None,
    limit: int = 200,
) -> Dict[str, int]:
    match = dict(extra or {})
    match[time_field] = {"$gte": cutoff}
    try:
        rows = await getattr(db, collection_name).aggregate([
            {"$match": match},
            {"$group": {"_id": f"${field}", "count": {"$sum": 1}}},
        ]).to_list(limit)
    except Exception:
        logger.warning("Launch monitor group failed for %s.%s", collection_name, field, exc_info=True)
        return {}
    counts: Dict[str, int] = {}
    for row in rows:
        key = str(row.get("_id") or "unknown")
        counts[key] = int(row.get("count", 0) or 0)
    return counts


def _rate(numerator: int, denominator: int) -> float:
    if denominator <= 0:
        return 0.0
    return round((int(numerator) / int(denominator)) * 100, 2)


def _conversion_rates(counts: Dict[str, int]) -> dict:
    return {
        "homepage_to_dracula_cta_pct": _rate(counts.get("hero_read_chapter_free_click", 0), counts.get("homepage_view", 0)),
        "dracula_to_reader_pct": _rate(counts.get("reader_opened", 0), counts.get("start_dracula_click", 0)),
        "reader_locked_to_pricing_pct": _rate(
            counts.get("pricing_page_view", 0),
            counts.get("reader_locked_state", 0) + counts.get("reader_low_balance_state", 0),
        ),
        "pricing_to_checkout_pct": _rate(counts.get("checkout_started", 0), counts.get("pricing_page_view", 0)),
        "checkout_to_payment_success_pct": _rate(counts.get("payment_success_return", 0), counts.get("checkout_started", 0)),
        "payment_success_to_continue_reading_pct": _rate(counts.get("continue_reading_click", 0), counts.get("payment_success_return", 0)),
    }


async def _launch_monitor_funnel_window(cutoff: str) -> dict:
    counts = await _group_counts_since(
        "analytics_events",
        "event",
        "created_at",
        cutoff,
        {"event": {"$in": LAUNCH_MONITOR_FUNNEL_EVENTS}},
    )
    full_counts = {event: int(counts.get(event, 0)) for event in LAUNCH_MONITOR_FUNNEL_EVENTS}
    return {
        "counts": full_counts,
        "conversion_rates": _conversion_rates(full_counts),
    }


async def _launch_monitor_payment_window(cutoff: str) -> dict:
    status_counts = await _group_counts_since("topup_intents", "status", "created_at", cutoff)
    webhook_status_counts = await _group_counts_since("payment_webhook_events", "status", "created_at", cutoff)
    support_counts = await _group_counts_since("contacts", "status", "created_at", cutoff)
    refund_counts = await _group_counts_since("wallet_refunds", "status", "created_at", cutoff)
    return {
        "payment_intents_created": await _count_since("topup_intents", "created_at", cutoff),
        "payment_success_count": int(status_counts.get("credited", 0)),
        "payment_failed_count": int(status_counts.get("failed", 0)),
        "wallet_credit_count": await _count_since("wallet_ledger", "timestamp", cutoff, {"action": "topup_credit"}),
        "webhook_received_count": await _count_since("payment_webhook_events", "created_at", cutoff),
        "webhook_duplicate_replay_blocked_count": int(webhook_status_counts.get("duplicate_replay_blocked", 0)),
        "topup_status_counts": status_counts,
        "webhook_status_counts": webhook_status_counts,
        "support_queue": support_counts,
        "refund_queue": refund_counts,
    }


async def _launch_monitor_core_web_vitals(cutoff: str) -> dict:
    try:
        rows = await db.analytics_events.find(
            {"event": "core_web_vital", "created_at": {"$gte": cutoff}},
            {"_id": 0, "metadata": 1, "route": 1, "created_at": 1},
        ).sort("created_at", -1).to_list(500)
    except Exception:
        logger.warning("Launch monitor Core Web Vitals query failed", exc_info=True)
        rows = []

    metrics: Dict[str, dict] = {}
    for row in rows:
        metadata = row.get("metadata") if isinstance(row.get("metadata"), dict) else {}
        name = str(metadata.get("metric") or "unknown")
        rating = str(metadata.get("rating") or "info")
        try:
            value = float(metadata.get("value"))
        except (TypeError, ValueError):
            continue
        item = metrics.setdefault(name, {"count": 0, "sum": 0.0, "good": 0, "needs-improvement": 0, "poor": 0, "info": 0})
        item["count"] += 1
        item["sum"] += value
        item[rating if rating in item else "info"] += 1

    return {
        "status": "COLLECTING" if rows else "NO_CORE_WEB_VITALS_EVENTS_YET",
        "metrics": {
            name: {
                "count": data["count"],
                "average_value": round(data["sum"] / data["count"], 2) if data["count"] else 0,
                "ratings": {
                    "good": data.get("good", 0),
                    "needs_improvement": data.get("needs-improvement", 0),
                    "poor": data.get("poor", 0),
                    "info": data.get("info", 0),
                },
            }
            for name, data in metrics.items()
        },
    }


def _public_audio_leak_status() -> dict:
    audio_exts = {".mp3", ".wav", ".m4a", ".ogg", ".aac"}
    repo_root = ROOT_DIR.parent
    public_dirs = [repo_root / "frontend" / "public", repo_root / "frontend" / "build"]
    files: list[str] = []
    for public_dir in public_dirs:
        if not public_dir.exists():
            continue
        for path in public_dir.rglob("*"):
            if path.is_file() and path.suffix.lower() in audio_exts:
                files.append(str(path.relative_to(repo_root)))
    return {
        "status": "PASS_NO_PUBLIC_AUDIO_FILES" if not files else "BLOCKED_PUBLIC_AUDIO_FILES_DETECTED",
        "public_audio_release": "PUBLIC_AUDIO_RELEASE_BLOCKED",
        "audiobook_production_status": "PRODUCTION_BLOCKED",
        "files": files[:20],
    }


def _post_deploy_canary_status() -> dict:
    repo_root = ROOT_DIR.parent
    candidates = [
        repo_root / "POST_DEPLOY_READING_ONLY_CANARY_REPORT.md",
        repo_root / "output" / "launch" / "post_deploy_reading_canary.json",
    ]
    for path in candidates:
        if not path.exists():
            continue
        try:
            text = path.read_text(encoding="utf-8")
        except Exception:
            continue
        if "PASS" in text and "PUBLIC_AUDIO_RELEASE_BLOCKED" in text:
            return {"status": "PASS_RECORDED", "source": str(path.relative_to(repo_root))}
    return {"status": "NOT_RECORDED_IN_REPO", "source": ""}


async def build_launch_monitor_summary() -> dict:
    windows = {
        "today": _launch_monitor_today_cutoff(),
        "last_24h": _launch_monitor_cutoff(24),
        "last_48h": _launch_monitor_cutoff(48),
    }
    funnel = {}
    payment = {}
    for label, cutoff in windows.items():
        funnel[label] = await _launch_monitor_funnel_window(cutoff)
        payment[label] = await _launch_monitor_payment_window(cutoff)

    return {
        "launch_status": "LIVE_VERIFIED",
        "public_audio_status": "PUBLIC_AUDIO_RELEASE_BLOCKED",
        "audiobook_production_status": "PRODUCTION_BLOCKED",
        "dashboard_status": "OWNER_ADMIN_ONLY",
        "windows": list(windows.keys()),
        "funnel": funnel,
        "payment": payment,
        "core_web_vitals": await _launch_monitor_core_web_vitals(windows["last_48h"]),
        "ops_health": {
            "backend_errors": {"status": "NOT_PERSISTED", "count": 0},
            "post_deploy_canary": _post_deploy_canary_status(),
            "public_audio_leak_check": _public_audio_leak_status(),
        },
        "action_checklist": [
            "Monitor Razorpay dashboard for live checkout anomalies.",
            "Confirm wallet credits match successful payments.",
            "Check support and refund queues twice daily during the first 48 hours.",
            "Run post-deploy canary after every production change.",
            "Keep audiobook release blocked until separate sync, accessibility, and release gates pass.",
        ],
        "privacy": {
            "pii_collected": False,
            "payment_identifiers_collected": False,
            "third_party_pixels": False,
            "analytics_mode": "first_party_opt_in_minimal_events",
        },
    }


def _safe_counter_map(counts: Dict[str, int]) -> Dict[str, int]:
    if not isinstance(counts, dict):
        return {}
    safe: Dict[str, int] = {}
    for raw_key, raw_value in list(counts.items())[:20]:
        key = re.sub(r"[^a-zA-Z0-9_.:-]", "_", str(raw_key))[:60]
        if not key:
            continue
        try:
            safe[key] = max(0, min(10000, int(raw_value)))
        except (TypeError, ValueError):
            safe[key] = 0
    return safe


def _safe_metric_name(value: str, fallback: str = "reader_metric") -> str:
    name = re.sub(r"[^a-zA-Z0-9_.:-]", "_", str(value or "").strip())[:80]
    return name or fallback


def _safe_numeric_map(values: Dict[str, Any], *, limit: int = 40, maximum: float = 10_000_000) -> Dict[str, float]:
    if not isinstance(values, dict):
        return {}
    safe: Dict[str, float] = {}
    for raw_key, raw_value in list(values.items())[:limit]:
        key = _safe_metric_name(raw_key, fallback="")
        if not key:
            continue
        try:
            number = float(raw_value)
        except (TypeError, ValueError):
            continue
        if not (number == number) or number < 0:
            continue
        safe[key] = min(maximum, number)
    return safe


def _safe_tag_map(values: Dict[str, Any], *, limit: int = 30) -> Dict[str, str]:
    if not isinstance(values, dict):
        return {}
    safe: Dict[str, str] = {}
    for raw_key, raw_value in list(values.items())[:limit]:
        key = _safe_metric_name(raw_key, fallback="")
        if not key:
            continue
        safe[key] = normalize_text(str(raw_value or ""))[:180]
    return safe


async def _aggregate_reader_metric(event: str, timings: Dict[str, float], metrics: Dict[str, float]) -> None:
    if not _redis_state_enabled():
        return
    day = datetime.utcnow().date().isoformat()
    redis_key = _redis_key("reader-rum", day, event)
    try:
        pipe = _redis_client.pipeline()
        pipe.hincrby(redis_key, "count", 1)
        for key, value in timings.items():
            pipe.hincrbyfloat(redis_key, f"timing:{key}:sum", float(value))
            pipe.hincrby(redis_key, f"timing:{key}:count", 1)
        for key, value in metrics.items():
            pipe.hincrbyfloat(redis_key, f"metric:{key}:sum", float(value))
            pipe.hincrby(redis_key, f"metric:{key}:count", 1)
        pipe.expire(redis_key, READER_RUM_AGGREGATE_TTL_SECONDS)
        await pipe.execute()
    except Exception:
        logger.warning("Reader RUM aggregate failed", exc_info=True)


def _should_persist_reader_metric(timings: Dict[str, float]) -> bool:
    slowest = max(timings.values(), default=0)
    if slowest >= READER_RUM_SLOW_MS:
        return True
    if READER_RUM_SAMPLE_RATE <= 0:
        return False
    return secrets.randbelow(1_000_000) < int(READER_RUM_SAMPLE_RATE * 1_000_000)


SECURE_READER_RECORDED_EVENTS = {
    "right_click",
    "copy",
    "cut",
    "print",
    "print_screen",
    "drag",
    "drop",
    "blocked_shortcut",
}


def _utc_day() -> str:
    return datetime.utcnow().date().isoformat()


async def _reader_completion_streak(user_id: str) -> int:
    rows = await db.reader_completions.find(
        {"user_id": user_id},
        {"_id": 0, "completed_on": 1},
    ).sort("completed_on", -1).to_list(90)
    completed_days = {row.get("completed_on") for row in rows if row.get("completed_on")}
    today = datetime.utcnow().date()
    streak = 0
    while (today - timedelta(days=streak)).isoformat() in completed_days:
        streak += 1
    return streak


async def _reader_reward_state(user_id: str) -> dict:
    streak = await _reader_completion_streak(user_id)
    claim = await db.reward_claims.find_one(
        {"user_id": user_id, "reward_key": READER_STREAK_REWARD_KEY},
        {"_id": 0},
    )
    claimed = bool(claim)
    return {
        "streak_days": streak,
        "required_days": READER_STREAK_REQUIRED_DAYS,
        "eligible": streak >= READER_STREAK_REQUIRED_DAYS and not claimed,
        "claimed": claimed,
        "credit_seconds": READER_STREAK_REWARD_SECONDS,
        "credit_minutes": READER_STREAK_REWARD_SECONDS // 60,
        "target_pack_id": "10h",
        "target_pack_label": "The Reader’s Reserve",
    }


def _client_ip(request: Request) -> str:
    for header in ("cf-connecting-ip", "x-real-ip"):
        value = request.headers.get(header)
        if value:
            return value.split(",", 1)[0].strip()[:64]
    forwarded_for = request.headers.get("x-forwarded-for")
    if forwarded_for:
        return forwarded_for.split(",", 1)[0].strip()[:64]
    if request.client and request.client.host:
        return request.client.host[:64]
    return "unknown"


def _rate_limit_scope(path: str) -> Tuple[str, int]:
    if _is_public_cache_path(path):
        return "public", RATE_LIMIT_PUBLIC_PER_MINUTE
    if path.startswith(("/api/reader/chapter/", "/api/reader/book/", "/api/reader/metrics")):
        return "reader", RATE_LIMIT_READER_PER_MINUTE
    if path.startswith("/api/auth/"):
        return "auth", RATE_LIMIT_AUTH_PER_MINUTE
    if path == "/api/payments/webhook":
        return "webhook", RATE_LIMIT_WEBHOOK_PER_MINUTE
    if path.startswith("/api/payments/"):
        return "payments", RATE_LIMIT_PAYMENT_PER_MINUTE
    if path.startswith("/api/analytics/"):
        return "analytics", RATE_LIMIT_ANALYTICS_PER_MINUTE
    if "upload" in path:
        return "upload", RATE_LIMIT_UPLOAD_PER_MINUTE
    return "api", RATE_LIMIT_DEFAULT_PER_MINUTE


def _rate_limit_identity(request: Request) -> str:
    auth = request.headers.get("authorization", "")
    if auth.lower().startswith("bearer "):
        token = auth.split(" ", 1)[1].strip()
        if token:
            try:
                payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALG])
                subject = str(payload.get("sub") or "")[:80]
                role = str(payload.get("role") or "user")[:20]
                if subject:
                    return f"token:{role}:{subject}"
            except (jwt.ExpiredSignatureError, jwt.InvalidTokenError):
                pass
    return f"ip:{_client_ip(request)}"


def _sweep_rate_limit_buckets(now: float) -> None:
    global _rate_limit_next_sweep
    if now < _rate_limit_next_sweep:
        return
    _rate_limit_next_sweep = now + RATE_LIMIT_WINDOW_SECONDS
    cutoff = now - RATE_LIMIT_WINDOW_SECONDS
    for key, bucket in list(_rate_limit_hits.items()):
        while bucket and bucket[0] <= cutoff:
            bucket.popleft()
        if not bucket:
            _rate_limit_hits.pop(key, None)


async def _redis_rate_limit_retry_after(key: str, limit: int, now: float) -> Optional[int]:
    redis_key = _redis_key("rate-limit", hashlib.sha256(key.encode("utf-8")).hexdigest())
    cutoff = now - RATE_LIMIT_WINDOW_SECONDS
    await _redis_client.zremrangebyscore(redis_key, 0, cutoff)
    count = await _redis_client.zcard(redis_key)
    if count >= limit:
        oldest = await _redis_client.zrange(redis_key, 0, 0, withscores=True)
        if oldest:
            oldest_score = float(oldest[0][1])
            return max(1, int(RATE_LIMIT_WINDOW_SECONDS - (now - oldest_score)))
        return RATE_LIMIT_WINDOW_SECONDS
    member = f"{now:.6f}:{uuid.uuid4()}"
    await _redis_client.zadd(redis_key, {member: now})
    await _redis_client.expire(redis_key, RATE_LIMIT_WINDOW_SECONDS * 2)
    return None


async def _rate_limit_retry_after(request: Request) -> Optional[int]:
    if not RATE_LIMIT_ENABLED or request.method == "OPTIONS":
        return None
    path = request.url.path
    if path in {"/health", "/healthz", "/api/health", "/api/healthz", "/docs", "/redoc", "/openapi.json", "/favicon.ico"}:
        return None
    scope, limit = _rate_limit_scope(path)
    now = time.monotonic()
    key = f"{scope}:{_rate_limit_identity(request)}"
    if _redis_state_enabled():
        return await _redis_rate_limit_retry_after(key, limit, now)
    _sweep_rate_limit_buckets(now)
    bucket = _rate_limit_hits[key]
    cutoff = now - RATE_LIMIT_WINDOW_SECONDS
    while bucket and bucket[0] <= cutoff:
        bucket.popleft()
    if len(bucket) >= limit:
        return max(1, int(RATE_LIMIT_WINDOW_SECONDS - (now - bucket[0])))
    bucket.append(now)
    return None


def _error_payload(request: Request, status_code: int, message: str, detail=None) -> dict:
    request_id = getattr(request.state, "request_id", None)
    return {
        "ok": False,
        "detail": detail if detail is not None else message,
        "error": {
            "code": status_code,
            "message": message,
        },
        "request_id": request_id,
    }


def _security_headers(response: Response) -> None:
    response.headers.setdefault("X-Content-Type-Options", "nosniff")
    response.headers.setdefault("X-Frame-Options", "DENY")
    response.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
    response.headers.setdefault("Permissions-Policy", "camera=(), microphone=(), geolocation=()")
    response.headers.setdefault("Content-Security-Policy", "frame-ancestors 'none'; object-src 'none'; base-uri 'none'")
    if ENVIRONMENT == "production":
        response.headers.setdefault("Strict-Transport-Security", "max-age=63072000; includeSubDomains; preload")


def _json_error_response(
    request: Request,
    status_code: int,
    message: str,
    detail=None,
    headers: Optional[dict] = None,
) -> JSONResponse:
    response = UTF8JSONResponse(
        status_code=status_code,
        content=_error_payload(request, status_code, message, detail),
        headers=headers,
    )
    if request_id := getattr(request.state, "request_id", None):
        response.headers["X-Request-ID"] = request_id
    _security_headers(response)
    return response


def _sanitize_validation_errors(errors: list) -> list:
    sanitized = []
    for error in errors:
        sanitized.append({
            "loc": error.get("loc", ()),
            "msg": error.get("msg", "Invalid input"),
            "type": error.get("type", "value_error"),
        })
    return sanitized


@app.middleware("http")
async def production_hardening_middleware(request: Request, call_next):
    request_id = request.headers.get("x-request-id") or str(uuid.uuid4())
    request.state.request_id = request_id
    started = time.perf_counter()
    status_code = 500
    path = request.url.path

    if request.method in {"POST", "PUT", "PATCH"}:
        raw_content_length = request.headers.get("content-length")
        if raw_content_length:
            try:
                content_length = int(raw_content_length)
            except ValueError:
                content_length = 0
            body_limit = _request_body_limit_for_path(path)
            if content_length > body_limit:
                response = _json_error_response(
                    request,
                    status_code=413,
                    message="Request body too large for production cost-control limits",
                    detail={"limit_bytes": body_limit, "content_length": content_length},
                )
                response.headers["X-Request-ID"] = request_id
                return response

    if _shutdown_state.get("draining") and path not in {"/health", "/healthz", "/api/health", "/api/healthz"}:
        response = _json_error_response(
            request,
            status_code=503,
            message="Server is draining before shutdown",
            headers={"Connection": "close", "Retry-After": "15"},
        )
        response.headers["X-Request-ID"] = request_id
        return response

    _shutdown_state["inflight"] = int(_shutdown_state.get("inflight", 0)) + 1
    try:
        retry_after = await _rate_limit_retry_after(request)
        if retry_after is not None:
            status_code = 429
            response = _json_error_response(
                request,
                status_code=429,
                message="Too many requests",
                headers={"Retry-After": str(retry_after)},
            )
        else:
            response = await call_next(request)
            status_code = response.status_code
    except Exception:
        logger.exception("Unhandled request failure")
        raise
    finally:
        _shutdown_state["inflight"] = max(0, int(_shutdown_state.get("inflight", 1)) - 1)

    response.headers["X-Request-ID"] = request_id
    _security_headers(response)
    duration_ms = round((time.perf_counter() - started) * 1000, 2)
    response.headers["Server-Timing"] = f"app;dur={duration_ms}"
    response.headers["X-Response-Time-ms"] = str(duration_ms)
    if (
        request.method == "GET"
        and response.status_code == 200
        and not request.headers.get("authorization")
        and _is_public_cache_path(path)
    ):
        response.headers.setdefault(
            "Cache-Control",
            f"public, max-age={min(PUBLIC_CACHE_TTL_SECONDS, 60)}, stale-while-revalidate=120",
        )
    elif path in {"/health", "/healthz", "/api/health", "/api/healthz"}:
        response.headers.setdefault("Cache-Control", "no-store")
    elif path.startswith(("/api/admin", "/api/users", "/api/reader", "/api/reading", "/api/payments")):
        response.headers.setdefault("Cache-Control", "no-store")
    if request.method in {"POST", "PUT", "PATCH", "DELETE"} and response.status_code < 400 and path.startswith("/api/admin"):
        await _public_cache_clear()
    logger.info(_json.dumps({
        "event": "http_request",
        "request_id": request_id,
        "method": request.method,
        "path": path,
        "status": status_code,
        "duration_ms": duration_ms,
        "client_ip": _client_ip(request),
        "user_agent": request.headers.get("user-agent", "")[:160],
    }, ensure_ascii=True))
    return response


@app.exception_handler(HTTPException)
async def structured_http_exception_handler(request: Request, exc: HTTPException):
    message = exc.detail if isinstance(exc.detail, str) else "Request failed"
    return _json_error_response(
        request,
        status_code=exc.status_code,
        message=message,
        detail=exc.detail,
        headers=exc.headers,
    )


@app.exception_handler(RequestValidationError)
async def structured_validation_exception_handler(request: Request, exc: RequestValidationError):
    return _json_error_response(
        request,
        status_code=422,
        message="Invalid request",
        detail=_sanitize_validation_errors(exc.errors()),
    )


@app.exception_handler(Exception)
async def structured_unhandled_exception_handler(request: Request, exc: Exception):
    logger.exception("Unhandled application error")
    return _json_error_response(
        request,
        status_code=500,
        message="Internal server error",
    )


# ---------- Auth ----------
@api.post("/auth/login", response_model=TokenOut)
async def login(payload: LoginIn):
    email = payload.email.lower().strip()
    user = await db.users.find_one({"email": email}, {"_id": 0})
    if not user or not verify_password(payload.password, user.get("password_hash")):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    token = create_token(user["id"], user["email"])
    return TokenOut(token=token, email=user["email"], role=user.get("role", "admin"))

@api.get("/auth/me")
async def me(user=Depends(require_admin)):
    return {"email": user["email"], "role": user["role"]}


@api.post("/auth/change-password")
async def change_password(payload: ChangePasswordIn, user=Depends(require_admin)):
    if len(payload.new_password) < 8:
        raise HTTPException(status_code=400, detail="New password must be at least 8 characters")
    existing = await db.users.find_one({"email": user["email"]}, {"_id": 0})
    if not existing or not verify_password(payload.current_password, existing.get("password_hash")):
        raise HTTPException(status_code=400, detail="Current password is incorrect")
    await db.users.update_one(
        {"email": user["email"]},
        {"$set": {"password_hash": hash_password(payload.new_password)}},
    )
    return {"ok": True, "message": "Password updated. Please sign in again."}


# ---------- Public: Categories ----------
@api.get("/health")
async def health_check():
    """Liveness + DB ping for load balancers and Docker healthchecks."""
    now = time.monotonic()
    if _health_cache["payload"] and _health_cache["expires_at"] > now:
        return _health_cache["payload"]
    db_ok = True
    try:
        await db.command("ping")
    except Exception:
        db_ok = False
    payload = {
        "ok": db_ok,
        "service": "the-earnalism-api",
        "mode": RAZORPAY_MODE,
        "razorpay_configured": razorpay_keys_configured(),
        "time": now_iso(),
    }
    _health_cache["payload"] = payload
    _health_cache["expires_at"] = now + HEALTH_CACHE_TTL_SECONDS
    return payload


@app.get("/health")
async def root_health_check():
    return await health_check()


@app.get("/healthz")
async def root_healthz_check():
    return {
        "status": "ok",
        "replica": "single" if not MULTI_REPLICA_ENABLED else os.environ.get("RAILWAY_REPLICA_ID", "multi"),
    }


@api.get("/healthz")
async def api_healthz_check():
    return await root_healthz_check()


@api.get("/controlled-launch/status")
async def controlled_launch_status():
    dracula_doc, dracula_source = await _find_public_book_candidate(
        "dracula",
        BOOK_METADATA_PROJECTION,
        include_artifact_content=False,
    )
    manifest = await _reader_book_manifest_doc("dracula")
    dracula_book_available = bool(_safe_live_public_projection(dracula_doc))
    dracula_manifest_available = bool(manifest and len(manifest.get("chapters") or []) == 27)
    catalog_truth_status = "PASS" if dracula_book_available and dracula_manifest_available else "HOLD"
    return {
        "backend_healthy": True,
        "live_approved_slugs": list(CONTROLLED_LIVE_BOOK_SLUGS),
        "pipeline_slugs": list(CONTROLLED_PIPELINE_SLUGS),
        "audio_enabled_slugs": list(CONTROLLED_AUDIO_ENABLED_SLUGS),
        "dracula_book_available": dracula_book_available,
        "dracula_manifest_available": dracula_manifest_available,
        "dracula_source": dracula_source,
        "catalog_truth_status": catalog_truth_status,
    }


@api.get("/home")
async def get_home_payload(books_limit: Optional[int] = None, books_offset: int = 0):
    """Cached landing-page payload.

    The home page needs categories, live books and a featured book immediately.
    Returning them together cuts the first-load API fanout from three database
    reads to one cached request while keeping the older public endpoints intact.
    """
    normalized_limit = _home_book_page_limit(books_limit, default=HOME_BOOK_LIMIT, maximum=HOME_BOOK_LIMIT)
    normalized_offset = max(0, int(books_offset or 0))
    cache_key = _public_cache_key("home_payload", books_limit=normalized_limit, books_offset=normalized_offset)
    cached = await _public_cache_get(cache_key)
    if cached is not None:
        return cached

    categories = await db.categories.find({}, {"_id": 0}).sort("order", 1).to_list(200)
    books_page = await _home_books_page(normalized_limit, normalized_offset, maximum=HOME_BOOK_LIMIT)

    featured_book = None
    setting = await db.settings.find_one({"key": "featured_book"}, {"_id": 0})
    featured_slug = (setting or {}).get("book_slug")
    featured_candidate = featured_slug if featured_slug in CONTROLLED_LIVE_BOOK_SLUGS else CONTROLLED_LIVE_BOOK_SLUGS[0]
    doc = await db.books.find_one(
        _controlled_public_book_query({"slug": featured_candidate}),
        BOOK_METADATA_PROJECTION,
    )
    featured_book = _safe_live_public_projection(doc)
    if not featured_book:
        featured_book = _safe_live_public_projection(_controlled_artifact_doc(featured_candidate, include_content=False))

    result = {
        "categories": categories,
        "books": books_page["books"],
        "books_page": books_page["pagination"],
        "featured": {"book": featured_book},
    }
    await _public_cache_set(cache_key, result)
    return result


def _home_book_page_limit(value: Optional[int], *, default: int, maximum: int = HOME_BOOK_PAGE_MAX_LIMIT) -> int:
    try:
        parsed = int(value if value is not None else default)
    except (TypeError, ValueError):
        parsed = default
    return max(1, min(parsed, maximum))


async def _home_books_page(limit: int, offset: int = 0, *, maximum: int = HOME_BOOK_PAGE_MAX_LIMIT) -> dict:
    normalized_limit = _home_book_page_limit(limit, default=HOME_BOOK_PAGE_DEFAULT_LIMIT, maximum=maximum)
    normalized_offset = max(0, int(offset or 0))
    cache_key = _public_cache_key("home_books_page", limit=normalized_limit, offset=normalized_offset)
    cached = await _public_cache_get(cache_key)
    if cached is not None:
        return cached

    query = _controlled_public_book_query()
    cursor = (
        db.books.find(query, BOOK_SUMMARY_PROJECTION)
        .sort("created_at", -1)
        .skip(normalized_offset)
        .limit(normalized_limit)
    )
    docs = await cursor.to_list(normalized_limit)
    total = await db.books.count_documents(query)
    books = []
    for doc in docs:
        projected = _safe_live_public_projection(doc)
        if projected:
            books.append(projected)
    if normalized_offset == 0:
        books = _append_controlled_artifact_projections(books)
    total = max(total, len(books))
    next_offset = normalized_offset + len(books)
    result = {
        "books": books,
        "pagination": {
            "offset": normalized_offset,
            "limit": normalized_limit,
            "count": len(books),
            "total": total,
            "next_offset": next_offset if next_offset < total else None,
            "has_more": next_offset < total,
        },
    }
    await _public_cache_set(cache_key, result)
    return result


@api.get("/home/books")
async def get_home_books_page(limit: int = HOME_BOOK_PAGE_DEFAULT_LIMIT, offset: int = 0):
    return await _home_books_page(limit, offset)


@api.post("/analytics/event")
async def record_analytics_event(
    payload: AnalyticsEventIn,
    request: Request,
    principal: Optional[dict] = Depends(optional_principal),
):
    await db.analytics_events.insert_one(_analytics_event_document(payload, request, principal))
    return {"ok": True}


@api.post("/analytics/events")
async def record_analytics_events_alias(
    payload: AnalyticsEventIn,
    request: Request,
    principal: Optional[dict] = Depends(optional_principal),
):
    return await record_analytics_event(payload, request, principal)


@api.post("/secure-reader/events")
async def record_secure_reader_event(
    payload: SecureReaderEventIn,
    request: Request,
    principal: Optional[dict] = Depends(optional_principal),
):
    event_type = re.sub(r"[^a-zA-Z0-9_.:-]", "_", payload.event_type.strip())[:80]
    if not event_type:
        raise HTTPException(status_code=400, detail="Event type is required")
    if event_type not in SECURE_READER_RECORDED_EVENTS:
        return {"ok": True, "stored": False}

    # Never store raw auth tokens. The frontend may send only a short fingerprint.
    await db.reader_security_events.insert_one({
        "id": str(uuid.uuid4()),
        "session_id": payload.session_id.strip()[:120],
        "event_type": event_type,
        "book_slug": payload.book_slug.strip()[:120],
        "chapter_id": payload.chapter_id.strip()[:120],
        "access_token_fingerprint": payload.access_token_fingerprint.strip()[:32],
        "counts": _safe_counter_map(payload.counts),
        "metadata": _safe_analytics_metadata(payload.metadata),
        "principal_role": principal.get("role") if principal else "guest",
        "user_id": principal.get("id") if principal and principal.get("role") == "user" else "",
        "user_email": principal.get("email") if principal and principal.get("role") == "user" else "",
        "ip": _client_ip(request),
        "user_agent": str(request.headers.get("user-agent", ""))[:180],
        "created_at": now_iso(),
    })
    return {"ok": True, "stored": True}


@api.get("/categories")
async def list_categories():
    cache_key = _public_cache_key("categories")
    cached = await _public_cache_get(cache_key)
    if cached is not None:
        return cached
    docs = await db.categories.find({}, {"_id": 0}).sort("order", 1).to_list(200)
    await _public_cache_set(cache_key, docs)
    return docs


@api.get("/home/curated")
async def get_home_curated():
    """Return the deterministic, file-backed Sprint 1 homepage projection."""
    return build_home_curated_payload()


# ---------- Public: Books ----------
@api.get("/books")
async def list_books(category: Optional[str] = None, q: Optional[str] = None):
    category_filter = None
    if category and category != "all":
        category_filter = normalize_category_slug(category) or category
    cache_key = _public_cache_key("books", category=category_filter or "all", q=normalize_text(q).strip() if q else "")
    cached = await _public_cache_get(cache_key)
    if cached is not None:
        return cached
    extra_query: dict = {}
    if category_filter:
        extra_query["category_slug"] = category_filter
    q_norm = normalize_text(q).strip() if q else ""
    if q_norm:
        pattern = re.escape(q_norm)
        extra_query["$or"] = [
            {"title": {"$regex": pattern, "$options": "i"}},
            {"subtitle": {"$regex": pattern, "$options": "i"}},
            {"author": {"$regex": pattern, "$options": "i"}},
            {"short_description": {"$regex": pattern, "$options": "i"}},
            {"description": {"$regex": pattern, "$options": "i"}},
            {"category_slug": {"$regex": pattern, "$options": "i"}},
            {"chapters.title": {"$regex": pattern, "$options": "i"}},
        ]
    query = _controlled_public_book_query(extra_query)
    docs = await db.books.find(query, BOOK_SUMMARY_PROJECTION).sort("created_at", -1).to_list(500)
    # Public list is shelf metadata-only so library browsing never ships chapter bodies.
    result = []
    for doc in docs:
        projected = _safe_live_public_projection(doc)
        if projected:
            result.append(projected)
    result = _append_controlled_artifact_projections(result, category_filter=category_filter, q=q_norm)
    await _public_cache_set(cache_key, result)
    return result

@api.get("/books/{slug}", response_model=PublicBookOut)
async def get_book(slug: str):
    cache_key = _public_cache_key("book_detail", slug=slug)
    cached = await _public_cache_get(cache_key)
    if cached is not None:
        return cached
    doc, _source = await _find_public_book_candidate(slug, BOOK_METADATA_PROJECTION, include_artifact_content=False)
    if not doc:
        raise HTTPException(status_code=404, detail="Book not found")
    # Public detail returns metadata + ToC only. Reader content is fetched through
    # the gated chapter endpoint so detail pages stay light for large books.
    result = _safe_live_public_projection(doc)
    if not result:
        raise HTTPException(status_code=404, detail="Book not found")
    await _public_cache_set(cache_key, result)
    return result


@api.get("/books/{slug}/chapters")
async def get_book_chapters(slug: str):
    cache_key = _public_cache_key("book_chapters", slug=slug)
    cached = await _public_cache_get(cache_key)
    if cached is not None:
        return cached
    doc, _source = await _find_public_book_candidate(slug, BOOK_METADATA_PROJECTION, include_artifact_content=False)
    if not doc:
        raise HTTPException(status_code=404, detail="Book not found")
    if not can_expose_reader(doc):
        raise HTTPException(status_code=404, detail="Book not found")
    chapters = _strip_all_chapter_content(doc).get("chapters") or []
    if not chapters:
        return []
    result = sorted(chapters, key=lambda c: c.get("order", 0))
    await _public_cache_set(cache_key, result)
    return result


@api.get("/books/{slug}/chapters/{chapter_id}")
async def get_book_chapter(slug: str, chapter_id: str):
    cache_key = _public_cache_key("book_chapter", slug=slug, chapter_id=chapter_id)
    cached = await _public_cache_get(cache_key)
    if cached is not None:
        return cached
    if not _is_controlled_public_slug(slug):
        raise HTTPException(status_code=404, detail="Book not found")
    book_meta, _source = await _find_public_book_candidate(slug, READER_ACCESS_PROJECTION, include_artifact_content=False)
    if not book_meta:
        raise HTTPException(status_code=404, detail="Book not found")
    if not can_expose_reader(book_meta):
        raise HTTPException(status_code=404, detail="Book not found")
    chapters = sorted((book_meta.get("chapters") or []), key=lambda c: c.get("order", 0))
    target_meta = next((c for c in chapters if c.get("id") == chapter_id), None)
    if not target_meta:
        raise HTTPException(status_code=404, detail="Chapter not found")

    chapter = dict(target_meta)
    chapter["content"] = ""
    if _is_free_preview_chapter(book_meta, chapter_id):
        content_doc = await db.books.find_one(
            {**_controlled_public_book_query({"slug": slug}), "chapters.id": chapter_id},
            CHAPTER_CONTENT_PROJECTION,
        )
        target = ((content_doc or {}).get("chapters") or [{}])[0]
        chapter["content"] = target.get("content", "")
        if not chapter["content"]:
            artifact = _controlled_artifact_doc(slug, include_content=True) or {}
            target = next((c for c in artifact.get("chapters") or [] if c.get("id") == chapter_id), {})
            chapter["content"] = target.get("content", "")
    if not chapter:
        raise HTTPException(status_code=404, detail="Chapter not found")
    await _public_cache_set(cache_key, chapter)
    return chapter


# ---------- Public: Blog ----------
@api.get("/blog", response_model=List[BlogPost])
async def list_blog(category: Optional[str] = None):
    cache_key = _public_cache_key("blog", category=category or "all")
    cached = await _public_cache_get(cache_key)
    if cached is not None:
        return cached
    query: dict = {
        "is_published": True,
        "slug": {"$nin": sorted(RETIRED_PUBLIC_BLOG_SLUGS)},
    }
    if category and category != "all":
        query["category"] = category
    docs = await db.blog_posts.find(query, {"_id": 0}).sort("created_at", -1).to_list(200)
    await _public_cache_set(cache_key, docs)
    return docs

@api.get("/blog/{slug}", response_model=BlogPost)
async def get_blog(slug: str):
    if slug in RETIRED_PUBLIC_BLOG_SLUGS:
        raise HTTPException(
            status_code=410,
            detail="Article removed",
            headers={"X-Robots-Tag": "noindex, nofollow, noarchive"},
        )
    cache_key = _public_cache_key("blog_detail", slug=slug)
    cached = await _public_cache_get(cache_key)
    if cached is not None:
        return cached
    doc = await db.blog_posts.find_one({"slug": slug, "is_published": True}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Article not found")
    await _public_cache_set(cache_key, doc)
    return doc


# ---------- Public: Featured ----------
@api.get("/featured")
async def get_featured():
    cache_key = _public_cache_key("featured")
    cached = await _public_cache_get(cache_key)
    if cached is not None:
        return cached
    s = await db.settings.find_one({"key": "featured_book"}, {"_id": 0})
    featured_candidate = (
        s.get("book_slug")
        if s and s.get("book_slug") in CONTROLLED_LIVE_BOOK_SLUGS
        else CONTROLLED_LIVE_BOOK_SLUGS[0]
    )
    book = await db.books.find_one(_controlled_public_book_query({"slug": featured_candidate}), BOOK_METADATA_PROJECTION)
    featured_book = _safe_live_public_projection(book)
    if not featured_book:
        featured_book = _safe_live_public_projection(_controlled_artifact_doc(featured_candidate, include_content=False))
    result = {"book": featured_book}
    await _public_cache_set(cache_key, result)
    return result


# ---------- Public: Submissions ----------
@api.post("/newsletter")
async def subscribe(payload: NewsletterIn):
    existing = await db.newsletter.find_one({"email": payload.email.lower()}, {"_id": 0})
    if existing:
        return {"ok": True, "message": "You're already part of the Reading Circle."}
    await db.newsletter.insert_one({
        "id": str(uuid.uuid4()),
        "name": payload.name.strip(),
        "email": payload.email.lower().strip(),
        "created_at": now_iso(),
    })
    return {"ok": True, "message": "Welcome to the Earnalism Reading Circle."}

@api.post("/contact")
async def contact(payload: ContactIn):
    await db.contacts.insert_one({
        "id": str(uuid.uuid4()),
        "name": payload.name.strip(),
        "email": payload.email.lower().strip(),
        "subject": payload.subject.strip(),
        "message": payload.message.strip(),
        "status": "new",
        "created_at": now_iso(),
    })
    return {"ok": True, "message": "Thank you. We'll respond with care."}


# ---------- Public: Social settings ----------
@api.get("/settings/social")
async def get_social():
    cache_key = _public_cache_key("settings_social")
    cached = await _public_cache_get(cache_key)
    if cached is not None:
        return cached
    doc = await db.settings.find_one({"key": "social"}, {"_id": 0}) or {}
    result = {
        "instagram": doc.get("instagram", ""),
        "facebook": doc.get("facebook", ""),
        "youtube": doc.get("youtube", ""),
        "linkedin": doc.get("linkedin", ""),
        "twitter": doc.get("twitter", ""),
    }
    await _public_cache_set(cache_key, result)
    return result


@api.get("/settings/brand")
async def get_brand():
    """Brand identity: logo URL + social-share OG image. Both optional.
    Empty strings are returned if not configured so the frontend can fall back
    to the existing premium text logo and hero image."""
    cache_key = _public_cache_key("settings_brand")
    cached = await _public_cache_get(cache_key)
    if cached is not None:
        return cached
    doc = await db.settings.find_one({"key": "brand"}, {"_id": 0}) or {}
    result = {
        "logo_url": doc.get("logo_url", ""),
        "og_image_url": doc.get("og_image_url", ""),
    }
    await _public_cache_set(cache_key, result)
    return result


@api.get("/settings/public")
async def get_public_settings():
    cache_key = _public_cache_key("settings_public")
    cached = await _public_cache_get(cache_key)
    if cached is not None:
        return cached
    result = {
        "social": await get_social(),
        "brand": await get_brand(),
    }
    await _public_cache_set(cache_key, result)
    return result


# ---------- Admin: Books ----------
@api.post("/admin/books", response_model=Book)
async def admin_create_book(payload: BookIn, _=Depends(require_admin)):
    data = payload.model_dump()
    rights_metadata = data.pop("rights_metadata", {}) or {}
    book_id = str(uuid.uuid4())
    data["title"] = normalize_text(data.get("title", "")).strip()
    if not data["title"]:
        raise HTTPException(status_code=400, detail="Title is required")
    data["author"] = normalize_text(data.get("author", "")).strip() or "The Earnalism"
    data["category_slug"] = canonical_category_slug(data.get("category_slug", ""))
    slug = slugify(data.pop("slug", None) or data["title"], fallback=f"book-{book_id[:8]}")
    if await db.books.find_one({"slug": slug}):
        raise HTTPException(status_code=400, detail="Slug already exists")
    book = Book(id=book_id, slug=slug, **data)
    doc = book.model_dump()
    doc["rights_metadata"] = rights_metadata
    _assert_publishable(doc)
    await db.books.insert_one(doc)
    return doc

@api.put("/admin/books/{slug}", response_model=Book)
async def admin_update_book(slug: str, payload: BookIn, _=Depends(require_admin)):
    existing = await db.books.find_one({"slug": slug}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Book not found")
    update = payload.model_dump()
    rights_metadata = update.pop("rights_metadata", None)
    update["title"] = normalize_text(update.get("title", "")).strip()
    if not update["title"]:
        raise HTTPException(status_code=400, detail="Title is required")
    update["author"] = normalize_text(update.get("author", "")).strip() or "The Earnalism"
    update["category_slug"] = canonical_category_slug(update.get("category_slug", ""))
    requested_slug = update.get("slug")
    if requested_slug:
        new_slug = slugify(requested_slug, fallback=slug)
    elif update["title"] != normalize_text(existing.get("title", "")):
        new_slug = slugify(update["title"], fallback=slug)
    else:
        new_slug = slug
    if new_slug != slug and await db.books.find_one({"slug": new_slug}, {"_id": 0}):
        raise HTTPException(status_code=400, detail="Slug already exists")
    update["slug"] = new_slug
    candidate = {**existing, **update}
    if rights_metadata:
        candidate["rights_metadata"] = rights_metadata
    _assert_publishable(candidate)
    if rights_metadata:
        update["rights_metadata"] = rights_metadata
    await db.books.update_one({"slug": slug}, {"$set": update})
    refreshed = await db.books.find_one({"slug": new_slug}, {"_id": 0})
    return refreshed


@api.patch("/admin/books/{slug}/audiobook", response_model=Book)
async def admin_update_book_audiobook(slug: str, payload: BookAudiobookIn, _=Depends(require_admin)):
    existing = await db.books.find_one({"slug": slug}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Book not found")

    provider = normalize_text(payload.audiobook_provider).strip()[:80].lower()
    assets = _safe_audiobook_assets(payload.audiobook_assets)
    if payload.audiobook_enabled or payload.generate_audiobook or assets:
        _assert_public_rights_approved(existing, "Audiobook")
    audiobook_doc = {
        **(_book_audiobook_doc(existing)),
        "url": assets.get("mp3", ""),
        "provider": provider,
        "size": max(0, int(payload.audiobook_size or 0)),
        "duration_ms": max(0, int(payload.audiobook_duration_ms or 0)),
        "assets": assets,
        "updated_at": now_iso(),
    }
    update = {
        "audiobook_enabled": bool(payload.audiobook_enabled),
        "generate_audiobook": bool(payload.generate_audiobook),
        "audiobook_provider": provider,
        "audiobook_voice": normalize_text(payload.audiobook_voice).strip()[:120],
        "audio_asset_slug": slugify(payload.audio_asset_slug or slug, fallback=slug),
        "audiobook_assets": assets,
        "audiobook": audiobook_doc,
        "audiobook_assets_updated_at": now_iso(),
    }
    await db.books.update_one({"slug": slug}, {"$set": update})
    refreshed = await db.books.find_one({"slug": slug}, {"_id": 0})
    return refreshed


@api.delete("/admin/books/{slug}")
async def admin_delete_book(slug: str, _=Depends(require_admin)):
    res = await db.books.delete_one({"slug": slug})
    return {"deleted": res.deleted_count}

@api.get("/admin/books", response_model=List[Book])
async def admin_list_books(_=Depends(require_admin)):
    return await db.books.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)


@api.get("/admin/books/summary", response_model=List[Book])
async def admin_list_books_summary(_=Depends(require_admin)):
    return await db.books.find({}, {"_id": 0, "chapters.content": 0}).sort("created_at", -1).to_list(1000)


@api.get("/admin/books/{slug}", response_model=Book)
async def admin_get_book(slug: str, _=Depends(require_admin)):
    return await _load_book_or_404(slug)


@api.get("/admin/rights/reports/{filename}")
async def admin_rights_report(filename: str, _=Depends(require_admin)):
    report_kind = next((kind for kind, expected in RIGHTS_REPORT_FILENAMES.items() if expected == filename), "")
    if not report_kind:
        raise HTTPException(status_code=404, detail="Unknown rights report")
    books = await db.books.find({}, {"_id": 0, "chapters.content": 0}).sort("created_at", -1).to_list(5000)
    rows = rights_report_rows(books, report_kind)
    return Response(
        rights_report_csv(rows),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@api.post("/upload_docx")
async def admin_upload_docx_validator(
    request: Request,
    docx_file: UploadFile = File(...),
    front_cover: Optional[UploadFile] = File(default=None),
    back_cover: Optional[UploadFile] = File(default=None),
    confirm_expensive_job: bool = False,
    admin=Depends(require_admin),
):
    _require_expensive_job_enabled(
        "book_rendering_jobs",
        enabled=ENABLE_BOOK_RENDERING_JOBS,
        confirm_expensive_job=confirm_expensive_job,
    )
    if front_cover or back_cover:
        _require_expensive_job_enabled(
            "cover_generation",
            enabled=ENABLE_COVER_GENERATION,
            confirm_expensive_job=confirm_expensive_job,
        )
    async with _expensive_job_slot("book_rendering_jobs"):
        return await _process_docx_upload(
            docx_file=docx_file,
            request=request,
            admin=admin,
            front_cover=front_cover,
            back_cover=back_cover,
        )


@api.post("/admin/books/import-template")
async def admin_import_book_template(
    request: Request,
    docx_file: UploadFile = File(...),
    front_cover: Optional[UploadFile] = File(default=None),
    back_cover: Optional[UploadFile] = File(default=None),
    confirm_expensive_job: bool = False,
    admin=Depends(require_admin),
):
    _require_expensive_job_enabled(
        "book_rendering_jobs",
        enabled=ENABLE_BOOK_RENDERING_JOBS,
        confirm_expensive_job=confirm_expensive_job,
    )
    if front_cover or back_cover:
        _require_expensive_job_enabled(
            "cover_generation",
            enabled=ENABLE_COVER_GENERATION,
            confirm_expensive_job=confirm_expensive_job,
        )
    async with _expensive_job_slot("book_rendering_jobs"):
        return await _process_docx_upload(
            docx_file=docx_file,
            request=request,
            admin=admin,
            front_cover=front_cover,
            back_cover=back_cover,
        )


@api.get("/credits/report")
async def credits_report(user_id: str, format: str = "json", _=Depends(require_admin)):
    requested_user = _sanitize_docx_plain(user_id, 160)
    if not requested_user:
        raise HTTPException(status_code=400, detail="user_id is required")
    rows = await db.credit_log.find({"user_id": requested_user}, {"_id": 0}).sort("timestamp", -1).to_list(1000)
    total = round(sum(float(row.get("credits_used", 0) or 0) for row in rows), 4)
    if format.lower() == "csv":
        headers = ["timestamp", "user_id", "session_id", "file_name", "operation_type", "task", "units", "credits_used", "upload_id"]
        lines = [",".join(headers)]
        for row in rows:
            line = []
            for header in headers:
                value = str(row.get(header, "")).replace('"', '""')
                line.append(f'"{value}"')
            lines.append(",".join(line))
        return Response(
            "\n".join(lines),
            media_type="text/csv",
            headers={"Content-Disposition": f'attachment; filename="earnalism-credit-report-{requested_user}.csv"'},
        )
    return {
        "user_id": requested_user,
        "total_credits_used": total,
        "records": rows,
        "generated_at": now_iso(),
    }


# ---------- Admin: Chapters (manual paste only for Phase 1) ----------
async def _load_book_or_404(slug: str) -> dict:
    doc = await db.books.find_one({"slug": slug}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Book not found")
    return doc

@api.post("/admin/books/{slug}/chapters", response_model=Book)
async def admin_add_chapter(slug: str, payload: ChapterIn, _=Depends(require_admin)):
    book = await _load_book_or_404(slug)
    _assert_public_rights_approved(book, "Source text")
    existing = book.get("chapters", []) or []
    next_order = max([c.get("order", 0) for c in existing], default=-1) + 1
    title = normalize_text(payload.title).strip()
    if not title:
        raise HTTPException(status_code=400, detail="Chapter title is required")
    content, warnings = _manual_content_to_render_html(payload.content)
    chapter = Chapter(
        title=title,
        content=content,
        order=next_order,
        is_preview=payload.is_preview,
        processing_status="ready",
        processing_warnings=warnings,
        updated_at=now_iso(),
    ).model_dump()
    await db.books.update_one({"slug": slug}, {"$push": {"chapters": chapter}})
    return await _load_book_or_404(slug)

@api.put("/admin/books/{slug}/chapters/reorder", response_model=Book)
async def admin_reorder_chapters(slug: str, payload: ChapterReorderIn, _=Depends(require_admin)):
    book = await _load_book_or_404(slug)
    existing = {c["id"]: c for c in (book.get("chapters") or [])}
    if set(payload.ids) != set(existing.keys()):
        raise HTTPException(status_code=400, detail="Reorder ids must match existing chapter ids exactly")
    reordered = []
    for i, cid in enumerate(payload.ids):
        c = dict(existing[cid])
        c["order"] = i
        reordered.append(c)
    await db.books.update_one({"slug": slug}, {"$set": {"chapters": reordered}})
    return await _load_book_or_404(slug)

@api.put("/admin/books/{slug}/chapters/{cid}", response_model=Book)
async def admin_update_chapter(slug: str, cid: str, payload: ChapterIn, _=Depends(require_admin)):
    book = await _load_book_or_404(slug)
    _assert_public_rights_approved(book, "Source text")
    title = normalize_text(payload.title).strip()
    if not title:
        raise HTTPException(status_code=400, detail="Chapter title is required")
    content, warnings = _manual_content_to_render_html(payload.content)
    res = await db.books.update_one(
        {"slug": slug, "chapters.id": cid},
        {"$set": {
            "chapters.$.title": title,
            "chapters.$.content": content,
            "chapters.$.is_preview": payload.is_preview,
            "chapters.$.processing_status": "ready",
            "chapters.$.processing_error": "",
            "chapters.$.processing_warnings": warnings,
            "chapters.$.updated_at": now_iso(),
        }},
    )
    if res.matched_count == 0:
        raise HTTPException(status_code=404, detail="Chapter not found")
    return await _load_book_or_404(slug)

@api.delete("/admin/books/{slug}/chapters/{cid}", response_model=Book)
async def admin_delete_chapter(slug: str, cid: str, _=Depends(require_admin)):
    await _load_book_or_404(slug)
    res = await db.books.update_one({"slug": slug}, {"$pull": {"chapters": {"id": cid}}})
    if res.modified_count == 0:
        raise HTTPException(status_code=404, detail="Chapter not found")
    return await _load_book_or_404(slug)


# ---------- Admin: Cloudinary uploads ----------
# Cloudinary setup is lazy — failed imports are surfaced at request time
# rather than blocking app boot when the optional deps aren't installed yet.
_CLOUDINARY_INITIALIZED = False

def _ensure_cloudinary():
    global _CLOUDINARY_INITIALIZED
    if _CLOUDINARY_INITIALIZED:
        return
    try:
        from config.cloudinary import init_cloudinary  # type: ignore
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Image pipeline unavailable: {e}")
    init_cloudinary()
    _CLOUDINARY_INITIALIZED = True


_ALLOWED_COVER_TYPES = {"image/jpeg", "image/png", "image/webp", "image/gif"}
_ALLOWED_CHAPTER_EXTS = {"docx", "md", "markdown", "html", "txt"}


@api.post("/admin/books/{slug}/cover")
async def admin_upload_cover(
    slug: str,
    kind: str = "front",
    confirm_expensive_job: bool = False,
    file: UploadFile = File(...),
    _=Depends(require_admin),
):
    _require_expensive_job_enabled(
        "cover_generation",
        enabled=ENABLE_COVER_GENERATION,
        confirm_expensive_job=confirm_expensive_job,
    )
    async with _expensive_job_slot("cover_generation"):
        cover_kind = "back" if kind == "back" else "front"
        if file.content_type not in _ALLOWED_COVER_TYPES:
            raise HTTPException(status_code=400, detail="Unsupported image type. Use JPG, PNG, WebP, or GIF.")
        body = await file.read()
        if len(body) > ADMIN_MEDIA_UPLOAD_MAX_BYTES:
            raise HTTPException(status_code=400, detail=f"Cover must be under {ADMIN_MEDIA_UPLOAD_MAX_BYTES} bytes")
        book = await _load_book_or_404(slug)
        _assert_public_rights_approved(book, "Visual asset")
        _ensure_cloudinary()
        try:
            from utils.content_processor import process_book_cover  # type: ignore
        except Exception as e:
            raise HTTPException(status_code=503, detail=f"Image pipeline unavailable: {e}")
        status_field = "back_cover_processing_status" if cover_kind == "back" else "cover_processing_status"
        error_field = "back_cover_processing_error" if cover_kind == "back" else "cover_processing_error"
        await db.books.update_one({"slug": slug}, {"$set": {status_field: "processing", error_field: ""}})
        try:
            result = process_book_cover(body, book.get("id") or slug, kind=cover_kind)
        except Exception as e:
            await db.books.update_one({"slug": slug}, {"$set": {status_field: "failed", error_field: str(e)}})
            raise HTTPException(status_code=400, detail=f"Cover processing failed: {e}")

        if cover_kind == "back":
            fields = {
                "back_cover_url": result["cover_url"],
                "back_cover_image_url": result["cover_url"],
                "back_cover_thumbnail_url": result["thumbnail_url"],
                "back_cover_blur_placeholder": result["blur_placeholder"],
                "back_cover_dominant_color": result["dominant_color"],
                "back_cover_processing_status": "ready",
                "back_cover_processing_error": "",
            }
        else:
            fields = {
                "cover_url": result["cover_url"],
                "cover_image_url": result["cover_url"],
                "thumbnail_url": result["thumbnail_url"],
                "blur_placeholder": result["blur_placeholder"],
                "dominant_color": result["dominant_color"],
                "cover_processing_status": "ready",
                "cover_processing_error": "",
            }
        await db.books.update_one(
            {"slug": slug},
            {"$set": fields},
        )
        return {"success": True, "kind": cover_kind, **result}


@api.post("/admin/books/{slug}/chapters/{chapter_id}/upload")
async def admin_upload_chapter_file(
    slug: str,
    chapter_id: str,
    confirm_expensive_job: bool = False,
    file: UploadFile = File(...),
    _=Depends(require_admin),
):
    _require_expensive_job_enabled(
        "book_rendering_jobs",
        enabled=ENABLE_BOOK_RENDERING_JOBS,
        confirm_expensive_job=confirm_expensive_job,
    )
    async with _expensive_job_slot("book_rendering_jobs"):
        ext = (file.filename or "").rsplit(".", 1)[-1].lower()
        if ext not in _ALLOWED_CHAPTER_EXTS:
            raise HTTPException(status_code=400, detail="Unsupported chapter format")
        body = await file.read()
        if len(body) > CHAPTER_UPLOAD_MAX_BYTES:
            raise HTTPException(status_code=400, detail=f"Chapter file must be under {CHAPTER_UPLOAD_MAX_BYTES} bytes")
        book = await _load_book_or_404(slug)
        _assert_public_rights_approved(book, "Source text")
        chapters = book.get("chapters") or []
        target = next((c for c in chapters if c.get("id") == chapter_id), None)
        if not target:
            raise HTTPException(status_code=404, detail="Chapter not found")
        _ensure_cloudinary()
        try:
            from utils.content_processor import process_chapter_content  # type: ignore
        except Exception as e:
            raise HTTPException(status_code=503, detail=f"Content pipeline unavailable: {e}")
        started_at = now_iso()
        await db.books.update_one(
            {"slug": slug, "chapters.id": chapter_id},
            {"$set": {
                "chapters.$.processing_status": "processing",
                "chapters.$.processing_error": "",
                "chapters.$.processing_warnings": [],
                "chapters.$.source_filename": file.filename or "chapter",
                "chapters.$.uploaded_at": started_at,
                "chapters.$.updated_at": started_at,
            }},
        )
        try:
            result = process_chapter_content(body, file.filename or "chapter", book.get("id") or slug)
        except Exception as e:
            await db.books.update_one(
                {"slug": slug, "chapters.id": chapter_id},
                {"$set": {
                    "chapters.$.processing_status": "failed",
                    "chapters.$.processing_error": str(e),
                    "chapters.$.updated_at": now_iso(),
                }},
            )
            raise HTTPException(status_code=400, detail=f"Chapter processing failed: {e}")

        warnings = result.get("warnings", [])
        new_chapters = []
        for c in chapters:
            if c.get("id") == chapter_id:
                c = dict(c)
                c["content"] = result["content_html"]
                c["has_images"] = result["has_images"]
                c["image_count"] = result["image_count"]
                c["word_count"] = result["word_count"]
                c["reading_minutes"] = result["reading_minutes"]
                c["language_hint"] = result.get("language_hint", "")
                c["processing_status"] = "ready"
                c["processing_error"] = ""
                c["processing_warnings"] = warnings
                c["source_filename"] = file.filename or "chapter"
                c["uploaded_at"] = started_at
                c["updated_at"] = now_iso()
            new_chapters.append(c)
        await db.books.update_one({"slug": slug}, {"$set": {"chapters": new_chapters}})
        return {
            "success": True,
            "processing_status": "ready",
            "word_count": result["word_count"],
            "reading_minutes": result["reading_minutes"],
            "has_images": result["has_images"],
            "image_count": result["image_count"],
            "language_hint": result.get("language_hint", ""),
            "warnings": warnings,
            "preview_html": result["content_html"],
        }


# ---------- Admin: Image upload (generic — for blog editor) ----------
@api.post("/admin/upload/image")
async def admin_upload_image(
    file: UploadFile = File(...),
    confirm_expensive_job: bool = False,
    _=Depends(require_admin),
):
    _require_expensive_job_enabled(
        "admin_media_uploads",
        enabled=ENABLE_ADMIN_MEDIA_UPLOADS,
        confirm_expensive_job=confirm_expensive_job,
    )
    async with _expensive_job_slot("admin_media_uploads"):
        if file.content_type not in _ALLOWED_COVER_TYPES:
            raise HTTPException(status_code=400, detail="Unsupported image type")
        body = await file.read()
        if len(body) > ADMIN_MEDIA_UPLOAD_MAX_BYTES:
            raise HTTPException(status_code=400, detail=f"Image must be under {ADMIN_MEDIA_UPLOAD_MAX_BYTES} bytes")
        _ensure_cloudinary()
        try:
            from config.cloudinary import upload_image  # type: ignore
        except Exception as e:
            raise HTTPException(status_code=503, detail=f"Image pipeline unavailable: {e}")
        result = upload_image(body, folder="earnalism/journal")
    return {"url": result["url"], "width": result.get("width"), "height": result.get("height")}


# ---------- Admin: Categories ----------
@api.post("/admin/categories", response_model=Category)
async def admin_create_cat(payload: CategoryIn, _=Depends(require_admin)):
    slug = slugify(payload.slug or payload.name)
    if await db.categories.find_one({"slug": slug}):
        raise HTTPException(status_code=400, detail="Slug already exists")
    cat = Category(slug=slug, **{k: v for k, v in payload.model_dump().items() if k != "slug"})
    await db.categories.insert_one(cat.model_dump())
    return cat

@api.put("/admin/categories/{slug}", response_model=Category)
async def admin_update_cat(slug: str, payload: CategoryIn, _=Depends(require_admin)):
    existing = await db.categories.find_one({"slug": slug}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Category not found")
    update = payload.model_dump()
    update["slug"] = slug  # keep slug stable
    await db.categories.update_one({"slug": slug}, {"$set": update})
    return await db.categories.find_one({"slug": slug}, {"_id": 0})

@api.delete("/admin/categories/{slug}")
async def admin_delete_cat(slug: str, _=Depends(require_admin)):
    res = await db.categories.delete_one({"slug": slug})
    return {"deleted": res.deleted_count}


# ---------- Admin: Blog ----------
@api.post("/admin/blog", response_model=BlogPost)
async def admin_create_post(payload: BlogPostIn, _=Depends(require_admin)):
    slug = slugify(payload.slug or payload.title)
    if await db.blog_posts.find_one({"slug": slug}):
        raise HTTPException(status_code=400, detail="Slug already exists")
    post = BlogPost(slug=slug, **{k: v for k, v in payload.model_dump().items() if k != "slug"})
    await db.blog_posts.insert_one(post.model_dump())
    return post

@api.put("/admin/blog/{slug}", response_model=BlogPost)
async def admin_update_post(slug: str, payload: BlogPostIn, _=Depends(require_admin)):
    existing = await db.blog_posts.find_one({"slug": slug}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Post not found")
    update = payload.model_dump()
    update["slug"] = slug
    await db.blog_posts.update_one({"slug": slug}, {"$set": update})
    return await db.blog_posts.find_one({"slug": slug}, {"_id": 0})

@api.delete("/admin/blog/{slug}")
async def admin_delete_post(slug: str, _=Depends(require_admin)):
    res = await db.blog_posts.delete_one({"slug": slug})
    return {"deleted": res.deleted_count}

@api.get("/admin/blog", response_model=List[BlogPost])
async def admin_list_posts(_=Depends(require_admin)):
    return await db.blog_posts.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)


# ---------- Admin: Submissions ----------
@api.get("/admin/newsletter")
async def admin_newsletter(_=Depends(require_admin)):
    return await db.newsletter.find({}, {"_id": 0}).sort("created_at", -1).to_list(2000)

@api.get("/admin/contacts")
async def admin_contacts(_=Depends(require_admin)):
    return await db.contacts.find({}, {"_id": 0}).sort("created_at", -1).to_list(2000)


@api.patch("/admin/contacts/{cid}/status")
async def admin_set_contact_status(cid: str, payload: ContactStatusIn, _=Depends(require_admin)):
    if payload.status not in VALID_CONTACT_STATUSES:
        raise HTTPException(status_code=400, detail="Invalid status")
    res = await db.contacts.update_one(
        {"id": cid},
        {"$set": {"status": payload.status, "updated_at": now_iso()}},
    )
    if res.matched_count == 0:
        raise HTTPException(status_code=404, detail="Contact not found")
    return {"ok": True, "id": cid, "status": payload.status}


# ---------- Admin: Settings (social) ----------
@api.put("/admin/settings/social")
async def admin_set_social(payload: SocialIn, _=Depends(require_admin)):
    data = payload.model_dump()
    await db.settings.update_one(
        {"key": "social"},
        {"$set": {"key": "social", **data}},
        upsert=True,
    )
    return {"ok": True, **data}


# ---------- Admin: Settings (brand identity) ----------
@api.put("/admin/settings/brand")
async def admin_set_brand(payload: BrandIn, _=Depends(require_admin)):
    data = payload.model_dump()
    await db.settings.update_one(
        {"key": "brand"},
        {"$set": {"key": "brand", **data}},
        upsert=True,
    )
    return {"ok": True, **data}


# ---------- Admin: Featured ----------
@api.put("/admin/featured")
async def admin_set_featured(payload: FeaturedIn, _=Depends(require_admin)):
    book = await db.books.find_one({"slug": payload.book_slug}, {"_id": 0})
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")
    await db.settings.update_one(
        {"key": "featured_book"},
        {"$set": {"key": "featured_book", "book_slug": payload.book_slug}},
        upsert=True,
    )
    return {"ok": True, "book_slug": payload.book_slug}


# ---------- Reader User: Auth ----------
@api.post("/users/signup", response_model=UserAuthOut)
async def user_signup(payload: UserSignupIn, request: Request, response: Response):
    name = payload.name.strip()
    email = payload.email.lower().strip()
    if not name:
        raise HTTPException(status_code=400, detail="Name is required")
    if len(payload.password) < 8:
        raise HTTPException(status_code=400, detail="Password must be at least 8 characters")
    if await db.users.find_one({"email": email}):
        raise HTTPException(status_code=400, detail="An account with this email already exists")
    doc = {
        "id": str(uuid.uuid4()),
        "name": name,
        "email": email,
        "password_hash": hash_password(payload.password),
        "role": "user",
        "auth_provider": "email",
        "reading_seconds_balance": 0,
        "status": "active",
        "created_at": now_iso(),
    }
    await db.users.insert_one(doc)
    token = await _create_user_session(doc, request, response)
    return UserAuthOut(token=token, user=UserOut(**_user_public(doc)))


@api.post("/users/login", response_model=UserAuthOut)
async def user_login(payload: UserLoginIn, request: Request, response: Response):
    email = payload.email.lower().strip()
    user = await db.users.find_one({"email": email}, {"_id": 0})
    if not user or user.get("role") != "user":
        raise HTTPException(status_code=401, detail="Invalid email or password")
    if user.get("status") == "blocked":
        raise HTTPException(status_code=403, detail="Account is blocked. Please contact support.")
    if not _has_password_credential(user) and user.get("auth_provider") not in {None, "", "email"}:
        raise HTTPException(status_code=403, detail=_password_login_unavailable_detail(user))
    if not verify_password(payload.password, user.get("password_hash")):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    token = await _create_user_session(user, request, response)
    return UserAuthOut(token=token, user=UserOut(**_user_public(user)))


@api.post("/users/logout")
async def user_logout(response: Response, user=Depends(require_user)):
    session_id = user.get("session_id")
    await db.user_sessions.update_one(
        {"id": session_id, "user_id": user["id"]},
        {"$set": {"status": "logged_out", "logged_out_at": datetime.now(timezone.utc)}},
    )
    await db.users.update_one({"id": user["id"], "active_user_session_id": session_id}, {"$unset": {"active_user_session_id": ""}})
    await _invalidate_user_cache(user["id"], session_ids=[session_id])
    _clear_user_refresh_cookie(response)
    return {"ok": True}


@api.post("/users/refresh", response_model=UserAuthOut)
async def user_refresh(
    request: Request,
    response: Response,
    refresh_token: Optional[str] = Cookie(default=None, alias=USER_REFRESH_COOKIE),
):
    if not refresh_token:
        raise HTTPException(status_code=401, detail="Session expired, please login again.")
    refreshed = await _refresh_user_session(refresh_token, request, response)
    if not refreshed:
        raise HTTPException(status_code=401, detail="Session expired, please login again.")
    return refreshed


# ---------- Social auth: Google + Mobile OTP ----------
GOOGLE_CLIENT_ID = os.environ.get("GOOGLE_CLIENT_ID", "").strip()
GOOGLE_CLIENT_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET", "").strip()
MSG91_AUTH_KEY = os.environ.get("MSG91_AUTH_KEY", "").strip()
MSG91_TEMPLATE_ID = os.environ.get("MSG91_TEMPLATE_ID", "").strip()
GOOGLE_TOKENINFO_URL = "https://oauth2.googleapis.com/tokeninfo"
GOOGLE_USERINFO_URL = "https://openidconnect.googleapis.com/v1/userinfo"


def _phone_email(mobile: str) -> str:
    digits = re.sub(r"\D", "", mobile)
    return f"mobile+{digits}@earnalism.local"


def _google_json(url: str, headers: Optional[Dict[str, str]] = None) -> dict:
    request = UrlRequest(url, headers=headers or {})
    with urlopen(request, timeout=10) as response:
        return _json.loads(response.read().decode("utf-8"))


def _google_audience_matches(info: dict) -> bool:
    audience = info.get("aud") or info.get("issued_to")
    if isinstance(audience, list):
        return GOOGLE_CLIENT_ID in audience
    return audience == GOOGLE_CLIENT_ID


def _google_email_verified(info: dict) -> bool:
    value = info.get("email_verified", info.get("verified_email", True))
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in {"1", "true", "yes"}


def _verify_google_access_token(access_token: str) -> dict:
    tokeninfo = _google_json(f"{GOOGLE_TOKENINFO_URL}?{urlencode({'access_token': access_token})}")
    if not _google_audience_matches(tokeninfo):
        raise ValueError("Google token audience mismatch")
    if not _google_email_verified(tokeninfo):
        raise ValueError("Google email is not verified")

    userinfo = {}
    if not tokeninfo.get("email") or not tokeninfo.get("name"):
        userinfo = _google_json(GOOGLE_USERINFO_URL, headers={"Authorization": f"Bearer {access_token}"})
        if not _google_email_verified(userinfo):
            raise ValueError("Google email is not verified")

    return {
        **userinfo,
        **tokeninfo,
        "email": (userinfo.get("email") or tokeninfo.get("email") or "").lower().strip(),
        "name": userinfo.get("name") or tokeninfo.get("name") or "",
        "picture": userinfo.get("picture") or tokeninfo.get("picture") or "",
    }


def _verify_google_credential(credential: str) -> dict:
    credential = (credential or "").strip()
    if not credential:
        raise ValueError("Missing Google credential")
    try:
        from google.oauth2 import id_token  # type: ignore
        from google.auth.transport import requests as google_requests  # type: ignore
    except Exception as e:
        raise RuntimeError(f"Google auth unavailable: {e}")

    try:
        idinfo = id_token.verify_oauth2_token(credential, google_requests.Request(), GOOGLE_CLIENT_ID)
        if not _google_email_verified(idinfo):
            raise ValueError("Google email is not verified")
        return idinfo
    except ValueError:
        return _verify_google_access_token(credential)


@api.post("/auth/google", response_model=UserAuthOut)
async def auth_google(payload: GoogleAuthIn, request: Request, response: Response):
    if not GOOGLE_CLIENT_ID:
        raise HTTPException(status_code=503, detail="Google sign-in is not configured")
    try:
        idinfo = _verify_google_credential(payload.credential)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=f"Google auth unavailable: {e}")
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid Google token")
    email = (idinfo.get("email") or "").lower().strip()
    if not email:
        raise HTTPException(status_code=401, detail="Google account has no email")
    name = idinfo.get("name") or email.split("@", 1)[0]
    picture = idinfo.get("picture") or ""

    user = await db.users.find_one({"email": email}, {"_id": 0})
    if not user:
        doc = {
            "id": str(uuid.uuid4()),
            "name": name,
            "email": email,
            "password_hash": "",
            "role": "user",
            "status": "active",
            "auth_provider": "google",
            "picture": picture,
            "reading_seconds_balance": 0,
            "created_at": now_iso(),
        }
        await db.users.insert_one(doc)
        user = doc
    elif user.get("status") == "blocked":
        raise HTTPException(status_code=403, detail="Account is blocked. Please contact support.")

    token = await _create_user_session(user, request, response)
    return UserAuthOut(token=token, user=UserOut(**_user_public(user)))


@api.post("/auth/otp/request")
async def auth_otp_request(payload: OTPRequestIn):
    mobile = payload.mobile.strip()
    if not re.match(r"^\+91[6-9]\d{9}$", mobile):
        raise HTTPException(status_code=400, detail="Invalid mobile number")
    import random
    otp = str(random.randint(100000, 999999))
    expires = datetime.now(timezone.utc) + timedelta(minutes=10)
    await db.otp_store.update_one(
        {"mobile": mobile},
        {"$set": {"otp": otp, "expires": expires, "attempts": 0}},
        upsert=True,
    )
    if MSG91_AUTH_KEY and MSG91_TEMPLATE_ID:
        try:
            import httpx  # type: ignore
            async with httpx.AsyncClient(timeout=10.0) as client:
                await client.post(
                    "https://api.msg91.com/api/v5/otp",
                    params={
                        "authkey": MSG91_AUTH_KEY,
                        "mobile": mobile.lstrip("+"),
                        "template_id": MSG91_TEMPLATE_ID,
                        "otp": otp,
                    },
                )
        except Exception as e:
            logger.warning(f"MSG91 send failed: {e}")
    return {"success": True, "message": "OTP sent"}


@api.post("/auth/otp/verify", response_model=UserAuthOut)
async def auth_otp_verify(payload: OTPVerifyIn, request: Request, response: Response):
    mobile = payload.mobile.strip()
    record = await db.otp_store.find_one({"mobile": mobile})
    if not record:
        raise HTTPException(status_code=400, detail="OTP not requested")
    expires = record.get("expires")
    if isinstance(expires, datetime) and expires.tzinfo is None:
        expires = expires.replace(tzinfo=timezone.utc)
    if not expires or expires < datetime.now(timezone.utc):
        raise HTTPException(status_code=400, detail="OTP expired")
    if int(record.get("attempts", 0)) >= 3:
        raise HTTPException(status_code=429, detail="Too many attempts")
    if record.get("otp") != payload.otp.strip():
        await db.otp_store.update_one({"mobile": mobile}, {"$inc": {"attempts": 1}})
        raise HTTPException(status_code=400, detail="Incorrect OTP")
    await db.otp_store.delete_one({"mobile": mobile})

    synthetic_email = _phone_email(mobile)
    user = await db.users.find_one({"$or": [{"mobile": mobile}, {"email": synthetic_email}]}, {"_id": 0})
    if not user:
        doc = {
            "id": str(uuid.uuid4()),
            "name": mobile,
            "email": synthetic_email,
            "mobile": mobile,
            "password_hash": "",
            "role": "user",
            "status": "active",
            "auth_provider": "mobile_otp",
            "reading_seconds_balance": 0,
            "created_at": now_iso(),
        }
        await db.users.insert_one(doc)
        user = doc
    elif user.get("status") == "blocked":
        raise HTTPException(status_code=403, detail="Account is blocked. Please contact support.")

    token = await _create_user_session(user, request, response)
    return UserAuthOut(token=token, user=UserOut(**_user_public(user)))


@api.get("/users/me", response_model=UserOut)
async def user_me(user=Depends(require_user)):
    return UserOut(**_user_public(user))


@api.get("/users/me/wallet")
async def user_wallet(user=Depends(require_user)):
    return {"wallet_seconds": await _cached_user_wallet_seconds(user["id"])}


@api.get("/users/me/transactions", response_model=List[WalletTransactionOut])
async def user_my_transactions(user=Depends(require_user)):
    cache_key = _user_transactions_cache_id(user["id"])
    cached = await _redis_cache_get("user-private", cache_key)
    if cached is not None:
        return cached
    rows = await db.wallet_transactions.find({"user_id": user["id"]}, {"_id": 0}).sort("created_at", -1).to_list(100)
    await _redis_cache_set("user-private", cache_key, rows, USER_TRANSACTIONS_CACHE_TTL_SECONDS)
    return rows


@api.get("/users/me/rewards")
async def user_reward_state(user=Depends(require_user)):
    return await _reader_reward_state(user["id"])


@api.post("/users/me/rewards/completion")
async def user_record_reader_completion(payload: ReaderCompletionIn, user=Depends(require_user)):
    progress = max(0, min(100, int(payload.progress or 0)))
    if progress < 90:
        state = await _reader_reward_state(user["id"])
        return {**state, "recorded": False}

    book_slug = payload.book_slug.strip()[:120]
    chapter_id = payload.chapter_id.strip()[:120]
    if not book_slug or not chapter_id:
        raise HTTPException(status_code=400, detail="Book and chapter are required")

    await db.reader_completions.update_one(
        {
            "user_id": user["id"],
            "book_slug": book_slug,
            "chapter_id": chapter_id,
            "completed_on": _utc_day(),
        },
        {
            "$set": {
                "chapter_title": payload.chapter_title.strip()[:180],
                "progress": progress,
                "updated_at": now_iso(),
            },
            "$setOnInsert": {
                "id": str(uuid.uuid4()),
                "created_at": now_iso(),
            },
        },
        upsert=True,
    )
    state = await _reader_reward_state(user["id"])
    return {**state, "recorded": True}


@api.post("/users/me/rewards/claim")
async def user_claim_reader_reward(user=Depends(require_user)):
    state = await _reader_reward_state(user["id"])
    if state["claimed"]:
        return {
            **state,
            "claimed_now": False,
            "wallet_seconds": await _cached_user_wallet_seconds(user["id"]),
        }
    if not state["eligible"]:
        raise HTTPException(status_code=400, detail="Reward is not available yet")

    res = await db.reward_claims.update_one(
        {"user_id": user["id"], "reward_key": READER_STREAK_REWARD_KEY},
        {
            "$setOnInsert": {
                "id": str(uuid.uuid4()),
                "reward_key": READER_STREAK_REWARD_KEY,
                "user_id": user["id"],
                "credited_seconds": READER_STREAK_REWARD_SECONDS,
                "created_at": now_iso(),
            }
        },
        upsert=True,
    )
    if res.upserted_id is None:
        return {
            **await _reader_reward_state(user["id"]),
            "claimed_now": False,
            "wallet_seconds": await _cached_user_wallet_seconds(user["id"]),
        }

    await db.users.update_one(
        {"id": user["id"]},
        {
            "$inc": {
                "reading_seconds_balance": READER_STREAK_REWARD_SECONDS,
                "wallet_seconds": READER_STREAK_REWARD_SECONDS,
            }
        },
    )
    fresh = await db.users.find_one({"id": user["id"]}, {"_id": 0, "reading_seconds_balance": 1, "wallet_seconds": 1}) or {}
    await _record_wallet_ledger(
        user_id=user["id"],
        action="reward_credit",
        seconds_delta=READER_STREAK_REWARD_SECONDS,
        reason="The Reader’s Reserve streak credit",
        actor="system",
        balance_after=int(fresh.get("reading_seconds_balance", fresh.get("wallet_seconds", 0)) or 0),
    )
    return {
        **await _reader_reward_state(user["id"]),
        "claimed_now": True,
        "wallet_seconds": int(fresh.get("reading_seconds_balance", fresh.get("wallet_seconds", 0)) or 0),
    }


@api.get("/reader/book/{slug}/manifest")
async def reader_book_manifest(
    slug: str,
    request: Request,
    response: Response,
    preview: Optional[str] = None,
    principal: Optional[dict] = Depends(optional_principal),
):
    is_admin_preview = bool((preview or "").lower() == "admin" and principal and principal.get("role") == "admin")
    manifest = await _reader_book_manifest_doc(slug, admin_preview=is_admin_preview)
    if not manifest:
        raise HTTPException(status_code=404, detail="Book not found")

    access = {
        "role": "guest",
        "authenticated": False,
        "admin_preview": is_admin_preview,
        "wallet_seconds": 0,
        "can_read_paid": False,
    }
    if principal and principal.get("role") == "admin":
        access.update({
            "role": "admin",
            "authenticated": True,
            "admin_preview": is_admin_preview,
            "can_read_paid": True,
        })
    elif principal and principal.get("role") == "user":
        wallet_seconds = 0 if principal.get("status") == "blocked" else await _cached_user_wallet_seconds(principal["id"])
        access.update({
            "role": "user",
            "authenticated": True,
            "status": principal.get("status", "active"),
            "wallet_seconds": wallet_seconds,
            "can_read_paid": wallet_seconds > 0 and principal.get("status") != "blocked",
        })

    payload = {**manifest, "access": access}
    etag = f'W/"reader-manifest-{manifest["version"]}"'
    response.headers["ETag"] = etag
    response.headers["X-Reader-Manifest-Version"] = manifest["version"]
    if access["role"] == "guest":
        response.headers["Cache-Control"] = "public, max-age=60, stale-while-revalidate=300"
    else:
        response.headers["Cache-Control"] = "private, max-age=20, stale-while-revalidate=60"
    if _client_etag_matches(request, etag):
        return Response(
            status_code=304,
            headers={
                "ETag": etag,
                "X-Reader-Manifest-Version": manifest["version"],
                "Cache-Control": response.headers["Cache-Control"],
            },
        )
    return payload


_b2_s3_client: Any = None


def _b2_is_configured() -> bool:
    return bool(B2_S3_ENDPOINT and B2_REGION and B2_BUCKET and B2_ACCESS_KEY_ID and B2_SECRET_ACCESS_KEY)


def _b2_client():
    global _b2_s3_client
    if _b2_s3_client is not None:
        return _b2_s3_client
    if not _b2_is_configured():
        raise HTTPException(status_code=503, detail="B2 audiobook storage is not configured")
    try:
        import boto3  # type: ignore
        from botocore.config import Config  # type: ignore
    except Exception as exc:
        raise HTTPException(status_code=503, detail=f"B2 client unavailable: {exc}")
    _b2_s3_client = boto3.client(
        "s3",
        endpoint_url=B2_S3_ENDPOINT,
        region_name=B2_REGION,
        aws_access_key_id=B2_ACCESS_KEY_ID,
        aws_secret_access_key=B2_SECRET_ACCESS_KEY,
        config=Config(s3={"addressing_style": "path"}),
    )
    return _b2_s3_client


def _b2_key_from_url(url: str) -> str:
    parsed = urlparse(url or "")
    path_parts = [unquote(part) for part in parsed.path.split("/") if part]
    if not path_parts:
        return ""
    if parsed.netloc.startswith(f"{B2_BUCKET}."):
        return "/".join(path_parts)
    if path_parts[0] == B2_BUCKET:
        return "/".join(path_parts[1:])
    return "/".join(path_parts)


def _parse_byte_range(range_header: str, total_size: int) -> Tuple[Optional[str], int]:
    value = (range_header or "").strip()
    if not value:
        return None, 200
    match = re.match(r"^bytes=(\d*)-(\d*)$", value)
    if not match:
        return None, 416
    start_raw, end_raw = match.groups()
    if not start_raw and not end_raw:
        return None, 416
    if start_raw:
        start = int(start_raw)
        end = int(end_raw) if end_raw else total_size - 1
    else:
        suffix = int(end_raw)
        if suffix <= 0:
            return None, 416
        start = max(0, total_size - suffix)
        end = total_size - 1
    if start < 0 or end < start or start >= total_size:
        return None, 416
    end = min(end, total_size - 1)
    return f"bytes={start}-{end}", 206


def _content_range_header(byte_range: str, total_size: int) -> str:
    match = re.match(r"^bytes=(\d+)-(\d+)$", byte_range or "")
    if not match:
        return f"bytes */{total_size}"
    return f"bytes {match.group(1)}-{match.group(2)}/{total_size}"


def _range_content_length(byte_range: str, total_size: int) -> int:
    match = re.match(r"^bytes=(\d+)-(\d+)$", byte_range or "")
    if not match:
        return total_size
    return max(0, int(match.group(2)) - int(match.group(1)) + 1)


def _content_range_total_size(content_range: str) -> int:
    match = re.search(r"/(\d+)$", content_range or "")
    if not match:
        return 0
    try:
        return int(match.group(1))
    except ValueError:
        return 0


def _streaming_body_iterator(body):
    try:
        while True:
            chunk = body.read(1024 * 1024)
            if not chunk:
                break
            yield chunk
    finally:
        close = getattr(body, "close", None)
        if callable(close):
            close()


def _audio_asset_content_type(asset_key: str, fallback: str = "") -> str:
    if fallback and fallback != "application/octet-stream":
        return fallback
    return {
        "mp3": "audio/mpeg",
        "timestamps": "application/json",
        "vtt": "text/vtt",
        "chapters": "application/json",
        "meta": "application/json",
        "manifest": "application/json",
    }.get(asset_key, "application/octet-stream")


def _audio_asset_cache_control(asset_key: str) -> str:
    # Browser-private because public audio remains separately gated; long enough
    # to keep playback and sidecar navigation warm during a reading session.
    if asset_key == "mp3":
        return "private, max-age=600, stale-while-revalidate=3600"
    return "private, max-age=3600, stale-while-revalidate=86400"


async def _b2_head_object(s3, *, bucket: str, key: str) -> dict:
    return await asyncio.to_thread(s3.head_object, Bucket=bucket, Key=key)


async def _b2_get_object(s3, *, bucket: str, key: str, byte_range: Optional[str]) -> dict:
    kwargs = {"Bucket": bucket, "Key": key}
    if byte_range:
        kwargs["Range"] = byte_range
    return await asyncio.to_thread(s3.get_object, **kwargs)


async def _reader_book_audiobook_asset(
    slug: str,
    asset_key: str,
    request: Request,
):
    normalized_key = str(asset_key or "mp3").strip().lower()
    if normalized_key not in ALLOWED_AUDIO_ASSET_KEYS:
        raise HTTPException(status_code=404, detail="Audiobook asset not found")
    if not _is_controlled_public_slug(slug):
        raise HTTPException(status_code=404, detail="Audiobook asset not found")
    book = await db.books.find_one(
        _controlled_public_book_query({"slug": slug}),
        {
            "_id": 0,
            "slug": 1,
            "is_published": 1,
            "audiobook": 1,
            "audiobook_assets": 1,
            "audiobook_provider": 1,
            "audiobook_enabled": 1,
            "generate_audiobook": 1,
            "rights_metadata": 1,
            "qa_status": 1,
            "source_hash": 1,
            "content_hash": 1,
            "provenance_hash": 1,
            "source_url": 1,
            "source_name": 1,
            "source_license": 1,
            "publication_status": 1,
            "approved_to_publish": 1,
        },
    )
    book = _reader_audio_truth_doc(book, slug)
    if not book or not can_expose_audio({**book, "slug": slug}):
        raise HTTPException(status_code=404, detail="Audiobook asset not found")
    asset_url = _book_audiobook_asset_url(book, normalized_key)
    if not _audio_asset_looks_like_b2(asset_url):
        if asset_url:
            return RedirectResponse(asset_url, status_code=307)
        raise HTTPException(status_code=404, detail="Audiobook asset not found")
    key = _b2_key_from_url(asset_url)
    if not key:
        raise HTTPException(status_code=404, detail="B2 audiobook asset key not found")

    s3 = _b2_client()
    range_header = (request.headers.get("range") or "").strip()
    if request.method != "HEAD" and range_header:
        try:
            obj = await _b2_get_object(s3, bucket=B2_BUCKET, key=key, byte_range=range_header)
        except Exception as exc:
            logger.warning("B2 audiobook asset ranged get failed for %s/%s/%s: %s", slug, normalized_key, key, exc)
            raise HTTPException(status_code=502, detail="Audiobook asset storage unavailable")
        content_range = obj.get("ContentRange") or ""
        content_type = _audio_asset_content_type(normalized_key, obj.get("ContentType") or "")
        content_length = int(obj.get("ContentLength") or 0)
        headers = {
            "Accept-Ranges": "bytes",
            "Content-Type": content_type,
            "Cache-Control": _audio_asset_cache_control(normalized_key),
            "Content-Length": str(content_length),
        }
        if content_range:
            headers["Content-Range"] = content_range
        etag = str(obj.get("ETag") or "").strip()
        if etag:
            headers["ETag"] = etag
        return StreamingResponse(
            _streaming_body_iterator(obj["Body"]),
            status_code=206 if content_range else 200,
            media_type=content_type,
            headers=headers,
        )

    try:
        head = await _b2_head_object(s3, bucket=B2_BUCKET, key=key)
    except Exception as exc:
        logger.warning("B2 audiobook asset head failed for %s/%s/%s: %s", slug, normalized_key, key, exc)
        raise HTTPException(status_code=404, detail="Audiobook asset object not found")

    total_size = int(head.get("ContentLength") or 0)
    content_type = _audio_asset_content_type(normalized_key, head.get("ContentType") or "")
    byte_range, status_code = _parse_byte_range(request.headers.get("range", ""), total_size)
    etag = str(head.get("ETag") or "").strip()
    base_headers = {
        "Accept-Ranges": "bytes",
        "Content-Type": content_type,
        "Cache-Control": _audio_asset_cache_control(normalized_key),
    }
    if etag:
        base_headers["ETag"] = etag
    if status_code == 416:
        return Response(status_code=416, headers={**base_headers, "Content-Range": f"bytes */{total_size}"})
    if not byte_range and _client_etag_matches(request, etag):
        return Response(status_code=304, headers=base_headers)

    if request.method == "HEAD":
        headers = {
            **base_headers,
            "Content-Length": str(_range_content_length(byte_range, total_size) if byte_range else total_size),
        }
        if status_code == 206 and byte_range:
            headers["Content-Range"] = _content_range_header(byte_range, total_size)
        return Response(status_code=status_code, headers=headers)

    try:
        obj = await _b2_get_object(s3, bucket=B2_BUCKET, key=key, byte_range=byte_range)
    except Exception as exc:
        logger.warning("B2 audiobook asset get failed for %s/%s/%s: %s", slug, normalized_key, key, exc)
        raise HTTPException(status_code=502, detail="Audiobook asset storage unavailable")

    content_length = int(obj.get("ContentLength") or total_size)
    headers = {
        **base_headers,
        "Content-Length": str(content_length),
    }
    if status_code == 206:
        headers["Content-Range"] = obj.get("ContentRange") or _content_range_header(byte_range, total_size)
    return StreamingResponse(
        _streaming_body_iterator(obj["Body"]),
        status_code=status_code,
        media_type=content_type,
        headers=headers,
    )


@api.api_route("/reader/book/{slug}/audiobook", methods=["GET", "HEAD"])
async def reader_book_audiobook(
    slug: str,
    request: Request,
    principal: Optional[dict] = Depends(optional_principal),
):
    return await _reader_book_audiobook_asset(slug, "mp3", request)


@api.api_route("/reader/book/{slug}/audiobook/{asset_key}", methods=["GET", "HEAD"])
async def reader_book_audiobook_sidecar(
    slug: str,
    asset_key: str,
    request: Request,
    principal: Optional[dict] = Depends(optional_principal),
):
    return await _reader_book_audiobook_asset(slug, asset_key, request)


@api.post("/reader/metrics")
async def reader_metrics(
    payload: ReaderMetricIn,
    request: Request,
    principal: Optional[dict] = Depends(optional_principal),
):
    if not READER_RUM_ENABLED:
        return {"ok": True, "recorded": False}
    event = _safe_metric_name(payload.event, fallback="reader_metric")
    timings = _safe_numeric_map(payload.timings, maximum=120_000)
    metrics = _safe_numeric_map(payload.metrics)
    tags = _safe_tag_map(payload.tags)
    await _aggregate_reader_metric(event, timings, metrics)

    should_persist = _should_persist_reader_metric(timings)
    if should_persist:
        await db.reader_experience_events.insert_one({
            "id": str(uuid.uuid4()),
            "event": event,
            "session_id": normalize_text(payload.session_id)[:120],
            "book_slug": normalize_text(payload.book_slug)[:160],
            "chapter_id": normalize_text(payload.chapter_id)[:160],
            "route": normalize_text(payload.route)[:240],
            "timings": timings,
            "metrics": metrics,
            "tags": tags,
            "role": (principal or {}).get("role", "guest") if principal else "guest",
            "user_id": (principal or {}).get("id", "") if principal and principal.get("role") == "user" else "",
            "client_country": request.headers.get("x-vercel-ip-country", request.headers.get("cf-ipcountry", ""))[:8],
            "created_at": now_iso(),
        })
    return {"ok": True, "recorded": should_persist}


# ---------- Reader: Sessions + Heartbeat ----------
HEARTBEAT_TICK_SECONDS = 30
LOW_BALANCE_THRESHOLD = 300


def _billable_reading_seconds(
    last_debit_at,
    now: datetime,
    *,
    visible: bool = True,
    idle: bool = False,
) -> int:
    """Return seconds to bill for this heartbeat.

    Billing is pulse-based, not catch-up based. A single heartbeat can never
    debit more than one 30-second pulse; long gaps usually mean the tab slept,
    the screen was hidden, or the user was away.
    """
    if not visible or idle:
        return 0
    last = _as_utc_dt(last_debit_at) or now
    elapsed = max(0, int((now - last).total_seconds()))
    if elapsed > READING_SESSION_IDLE_GRACE_SECONDS:
        return 0
    if elapsed + READING_HEARTBEAT_EARLY_GRACE_SECONDS < HEARTBEAT_TICK_SECONDS:
        return 0
    return HEARTBEAT_TICK_SECONDS


def _should_reset_reading_clock(last_debit_at, now: datetime, *, visible: bool = True, idle: bool = False) -> bool:
    if not visible or idle:
        return True
    last = _as_utc_dt(last_debit_at) or now
    elapsed = max(0, int((now - last).total_seconds()))
    return elapsed > READING_SESSION_IDLE_GRACE_SECONDS


async def _apply_reading_debit(user_id: str, session_id: str, book_title: str, seconds: int) -> Tuple[int, int]:
    if seconds <= 0:
        return 0, await _cached_user_wallet_seconds(user_id)

    fresh = await db.users.find_one({"id": user_id}, {"_id": 0, "reading_seconds_balance": 1, "wallet_seconds": 1}) or {}
    wallet = int(fresh.get("reading_seconds_balance", fresh.get("wallet_seconds", 0)) or 0)
    debit = min(seconds, max(0, wallet))
    remaining = max(0, wallet - debit)
    if debit <= 0:
        return 0, remaining

    res = await db.users.update_one(
        {"id": user_id, "reading_seconds_balance": wallet},
        {"$set": {"reading_seconds_balance": remaining, "wallet_seconds": remaining}},
    )
    if res.modified_count != 1:
        fresh = await db.users.find_one({"id": user_id}, {"_id": 0, "reading_seconds_balance": 1, "wallet_seconds": 1}) or {}
        wallet = int(fresh.get("reading_seconds_balance", fresh.get("wallet_seconds", 0)) or 0)
        await _set_user_wallet_cache(user_id, wallet)
        return 0, wallet
    await _record_wallet_ledger(
        user_id=user_id,
        session_id=session_id,
        action="reading_debit",
        seconds_delta=-debit,
        reason=f"Reading {book_title[:40]}",
        actor="system",
        balance_after=remaining,
    )
    return debit, remaining


async def _settle_active_reading_session(
    user_id: str,
    session_id: str,
    book_title: str = "",
    *,
    visible: bool = True,
    idle: bool = False,
) -> Tuple[int, int, str]:
    now = datetime.now(timezone.utc)
    active = await db.users.find_one({"id": user_id}, {"_id": 0, "active_reading_session": 1})
    session = (active or {}).get("active_reading_session") or {}
    if session.get("session_id") != session_id:
        return 0, await _cached_user_wallet_seconds(user_id), "session_invalid"

    last_debit_at = session.get("last_debit_at") or session.get("last_pulse_at") or session.get("started_at")
    billable = _billable_reading_seconds(last_debit_at, now, visible=visible, idle=idle)
    reset_clock = _should_reset_reading_clock(last_debit_at, now, visible=visible, idle=idle)
    if billable <= 0:
        if reset_clock:
            await db.users.update_one(
                {"id": user_id, "active_reading_session.session_id": session_id},
                {"$set": {"active_reading_session.last_debit_at": now, "active_reading_session.last_pulse_at": now}},
            )
        else:
            await db.users.update_one(
                {"id": user_id, "active_reading_session.session_id": session_id},
                {"$set": {"active_reading_session.last_pulse_at": now}},
            )
        status = "paused" if reset_clock else "active"
        return 0, await _cached_user_wallet_seconds(user_id), status

    debit_filter = {"id": user_id, "active_reading_session.session_id": session_id}
    if session.get("last_debit_at") is not None:
        debit_filter["active_reading_session.last_debit_at"] = session.get("last_debit_at")
    res = await db.users.update_one(
        debit_filter,
        {"$set": {"active_reading_session.last_debit_at": now, "active_reading_session.last_pulse_at": now}},
    )
    if res.modified_count != 1:
        return 0, await _cached_user_wallet_seconds(user_id), "session_invalid"
    deducted, remaining = await _apply_reading_debit(user_id, session_id, book_title or "Earnalism", billable)
    return deducted, remaining, "active"


@api.post("/reader/session/start")
async def reader_session_start(payload: ReaderSessionStartIn, user=Depends(require_user)):
    book_slug = payload.book_slug or payload.book_id
    if not book_slug:
        raise HTTPException(status_code=400, detail="Book id is required")
    book = await _reader_book_access_doc(book_slug)
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")
    balance = await _cached_user_wallet_seconds(user["id"])
    is_preview = _is_free_preview_chapter(book, payload.chapter_id)
    if balance <= 0 and not is_preview:
        raise HTTPException(status_code=402, detail="Insufficient reading time. Top up to continue.")
    sid = str(uuid.uuid4())
    await db.reading_sessions.update_many(
        {"user_id": user["id"], "status": "active"},
        {"$set": {"status": "replaced", "ended_at": now_iso()}},
    )
    now = datetime.now(timezone.utc)
    await db.reading_sessions.insert_one({
        "id": sid,
        "user_id": user["id"],
        "book_slug": book_slug,
        "chapter_id": payload.chapter_id,
        "started_at": now,
        "last_heartbeat_at": now,
        "last_debit_at": now,
        "seconds_consumed": 0,
        "status": "active",
        "auth_session_id": user.get("session_id", ""),
    })
    return {
        "session_id": sid,
        "remaining_seconds": balance,
        "is_preview": is_preview,
        "tick_seconds": HEARTBEAT_TICK_SECONDS,
    }


@api.post("/reader/heartbeat")
async def reader_heartbeat(payload: ReaderHeartbeatIn, user=Depends(require_user)):
    session = await db.reading_sessions.find_one({"id": payload.session_id, "user_id": user["id"]}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    if session.get("status") != "active":
        return {"deducted_seconds": 0, "remaining_seconds": 0, "status": session.get("status", "ended"), "is_preview": False}
    if session.get("auth_session_id") and session.get("auth_session_id") != user.get("session_id"):
        await db.reading_sessions.update_one(
            {"id": payload.session_id},
            {"$set": {"status": "replaced", "ended_at": now_iso(), "ended_reason": "auth_session_mismatch"}},
        )
        return {"deducted_seconds": 0, "remaining_seconds": 0, "status": "session_invalid", "is_preview": False}

    book = await _reader_book_access_doc(session["book_slug"])
    if not book:
        await db.reading_sessions.update_one({"id": payload.session_id}, {"$set": {"status": "ended"}})
        raise HTTPException(status_code=404, detail="Book not found")

    chapter_id = payload.chapter_id or session.get("chapter_id")
    is_preview = _is_free_preview_chapter(book, chapter_id)

    balance = await _cached_user_wallet_seconds(user["id"])

    deducted = 0
    status = "active"

    if is_preview:
        status = "preview"
    elif balance <= 0:
        status = "depleted"
    else:
        now = datetime.now(timezone.utc)
        last_debit_at = session.get("last_debit_at") or session.get("last_heartbeat_at")
        billable = _billable_reading_seconds(
            last_debit_at,
            now,
            visible=payload.visible,
            idle=payload.idle,
        )
        reset_clock = _should_reset_reading_clock(
            last_debit_at,
            now,
            visible=payload.visible,
            idle=payload.idle,
        )
        if billable > 0:
            heartbeat_filter = {"id": payload.session_id}
            if session.get("last_debit_at") is not None:
                heartbeat_filter["last_debit_at"] = session.get("last_debit_at")
            res = await db.reading_sessions.update_one(
                heartbeat_filter,
                {"$set": {"last_debit_at": now}},
            )
            if res.modified_count == 1:
                deducted, balance = await _apply_reading_debit(user["id"], payload.session_id, book.get("title", "Earnalism"), billable)
                status = "depleted" if balance <= 0 else "active"
        elif reset_clock:
            await db.reading_sessions.update_one(
                {"id": payload.session_id},
                {"$set": {"last_debit_at": now}},
            )
            status = "paused"

    await db.reading_sessions.update_one(
        {"id": payload.session_id},
        {
            "$set": {"last_heartbeat_at": datetime.now(timezone.utc), "chapter_id": chapter_id},
            "$inc": {"seconds_consumed": deducted},
        },
    )

    return {
        "deducted_seconds": deducted,
        "remaining_seconds": balance,
        "status": status,
        "is_preview": is_preview,
    }


@api.post("/reader/session/end")
async def reader_session_end(payload: ReaderSessionEndIn, user=Depends(require_user)):
    res = await db.reading_sessions.update_one(
        {"id": payload.session_id, "user_id": user["id"], "status": "active"},
        {"$set": {"status": "ended", "ended_at": now_iso()}},
    )
    return {"ended": res.modified_count}


# Aliases for /reading/session/ endpoints (frontend compatibility)
@api.post("/reading/session/start")
async def reading_session_start_v2(payload: ReaderSessionStartIn, request: Request, user=Depends(require_user)):
    session_id = payload.session_id or str(uuid.uuid4())
    book_id = payload.book_id or payload.book_slug
    if not book_id:
        raise HTTPException(status_code=400, detail="Book id is required")
    fresh = await db.users.find_one({"id": user["id"]}, {"_id": 0, "active_reading_session": 1}) or {}
    existing = (fresh.get("active_reading_session") or {})
    if existing.get("session_id") and existing.get("session_id") != session_id:
        old_book = await _reader_book_access_doc(existing.get("book_id", "")) or {}
        await _settle_active_reading_session(user["id"], existing["session_id"], old_book.get("title", "Earnalism"))
    now = datetime.utcnow()
    await db.users.update_one(
        {"id": user["id"]},
        {
            "$set": {
                "active_reading_session": {
                    "session_id": session_id,
                    "book_id": book_id,
                    "chapter_id": payload.chapter_id,
                    "started_at": now,
                    "last_pulse_at": now,
                    "last_debit_at": now,
                    "device": request.headers.get("user-agent", "unknown")[:100],
                    "auth_session_id": user.get("session_id", ""),
                }
            }
        },
    )
    await _invalidate_user_cache(user["id"])
    return {"success": True, "session_id": session_id}


@api.post("/reading/session/end")
async def reading_session_end_v2(payload: ReaderSessionEndIn, principal: Optional[dict] = Depends(optional_principal)):
    if not principal or principal.get("role") != "user":
        return {"success": False, "status": "session_invalid"}
    user = principal
    fresh = await db.users.find_one({"id": user["id"]}, {"_id": 0, "active_reading_session": 1}) or {}
    active = fresh.get("active_reading_session") or {}
    if active.get("session_id") != payload.session_id:
        return {"success": False, "status": "session_invalid"}
    if active.get("auth_session_id") and active.get("auth_session_id") != user.get("session_id"):
        return {"success": False, "status": "session_invalid"}
    book = await _reader_book_access_doc(active.get("book_id", "")) or {}
    await _settle_active_reading_session(user["id"], payload.session_id, book.get("title", "Earnalism"))
    await db.users.update_one({"id": user["id"]}, {"$unset": {"active_reading_session": ""}})
    await _invalidate_user_cache(user["id"])
    return {"success": True}


@api.post("/reading/pulse")
async def reading_pulse(payload: ReadingPulseIn, principal: Optional[dict] = Depends(optional_principal)):
    if not principal or principal.get("role") != "user" or principal.get("status") == "blocked":
        return {"success": False, "status": "session_invalid"}
    user = principal
    fresh = await db.users.find_one({"id": user["id"]}, {"_id": 0, "active_reading_session": 1}) or {}
    active = fresh.get("active_reading_session") or {}
    if active.get("session_id") != payload.session_id:
        return {"success": False, "status": "session_invalid"}
    if active.get("auth_session_id") and active.get("auth_session_id") != user.get("session_id"):
        return {"success": False, "status": "session_invalid"}

    wallet = await _cached_user_wallet_seconds(user["id"])
    if wallet <= 0:
        return {"success": False, "status": "wallet_empty", "wallet_seconds": 0}

    book = await _reader_book_access_doc(active.get("book_id", "")) or {}
    deducted, remaining, settle_status = await _settle_active_reading_session(
        user["id"],
        payload.session_id,
        book.get("title", "Earnalism"),
        visible=payload.visible,
        idle=payload.idle,
    )
    if settle_status in {"paused", "session_invalid"}:
        status = settle_status
    else:
        status = "wallet_empty" if remaining <= 0 else ("low_balance" if remaining <= LOW_BALANCE_THRESHOLD else "ok")
    if deducted <= 0:
        await db.users.update_one(
            {"id": user["id"]},
            {"$set": {"active_reading_session.last_pulse_at": datetime.now(timezone.utc)}},
        )
    return {"success": status != "wallet_empty", "status": status, "wallet_seconds": remaining, "deducted_seconds": deducted}


@api.get("/reading/packs")
async def reading_packs():
    cache_key = _public_cache_key("reading_packs")
    cached = await _public_cache_get(cache_key)
    if cached is not None:
        return cached
    result = [
        {"id": p["id"], "minutes": p["minutes"], "price": p["price_inr"], "label": p["label"]}
        for p in PACKS
    ]
    await _public_cache_set(cache_key, result)
    return result


# ---------- Reader: Gated chapter content ----------
# Returns chapter BODY only when the caller is authorised. Guests, blocked
# users and users with no reading time get a locked metadata-only response.
# Admin tokens always pass (used for the admin reader preview).
@api.get("/reader/chapter/{slug}/{chapter_id}")
async def reader_get_chapter(
    slug: str,
    chapter_id: str,
    request: Request,
    response: Response,
    v: Optional[str] = None,
    principal: Optional[dict] = Depends(optional_principal),
):
    is_admin_preview = bool(principal and principal.get("role") == "admin")
    book_meta = await _reader_book_access_doc(slug, admin_preview=is_admin_preview)
    if not book_meta:
        raise HTTPException(status_code=404, detail="Book not found")
    chapters = sorted((book_meta.get("chapters") or []), key=lambda c: c.get("order", 0))
    target_meta = next((c for c in chapters if c.get("id") == chapter_id), None)
    if not target_meta:
        raise HTTPException(status_code=404, detail="Chapter not found")

    meta = {
        "id": target_meta["id"],
        "title": target_meta.get("title", ""),
        "order": target_meta.get("order", 0),
        "is_preview": target_meta.get("is_preview", False),
        "content_version": target_meta.get("content_version") or v or "",
    }
    is_preview = _is_free_preview_chapter(book_meta, chapter_id)
    content_version = meta["content_version"] or _stable_digest({
        "slug": slug,
        "chapter_id": chapter_id,
        "book_generation": await _reader_content_cache_generation_value(),
    }, length=20)
    response.headers["ETag"] = f'W/"reader-chapter-{content_version}"'
    response.headers["X-Reader-Chapter-Version"] = content_version

    def not_modified_response(cache_control: str) -> Optional[Response]:
        etag = response.headers["ETag"]
        if not _client_etag_matches(request, etag):
            return None
        return Response(
            status_code=304,
            headers={
                "ETag": etag,
                "X-Reader-Chapter-Version": content_version,
                "Cache-Control": cache_control,
            },
        )

    async def unlocked_chapter_response(preview: bool):
        content = await _reader_chapter_content(slug, chapter_id, admin_preview=is_admin_preview)
        return {
            "locked": False,
            "is_preview": preview,
            "chapter": {**meta, "content_version": content_version, "is_preview": preview, "content": content},
        }

    # Free preview is open to everyone — no auth, no deduction.
    if is_preview and not is_admin_preview:
        response.headers["Cache-Control"] = "public, max-age=60, stale-while-revalidate=300"
        not_modified = not_modified_response(response.headers["Cache-Control"])
        if not_modified:
            return not_modified
        cache_key = _public_cache_key("reader_preview_chapter", slug=slug, chapter_id=chapter_id)
        cached = await _public_cache_get(cache_key)
        if cached is not None:
            return cached
        result = await unlocked_chapter_response(True)
        await _public_cache_set(cache_key, result)
        return result

    # Admin always has full access (admin reader preview).
    if is_admin_preview:
        response.headers["Cache-Control"] = "private, max-age=20"
        not_modified = not_modified_response(response.headers["Cache-Control"])
        if not_modified:
            return not_modified
        return await unlocked_chapter_response(False)

    # Reader user
    if principal and principal.get("role") == "user":
        if principal.get("status") == "blocked":
            return {
                "locked": True,
                "reason": "BLOCKED",
                "message": "Account is blocked. Please contact support.",
                "chapter": meta,
            }
        if await _cached_user_wallet_seconds(principal["id"]) > 0:
            response.headers["Cache-Control"] = "private, max-age=45, stale-while-revalidate=180"
            not_modified = not_modified_response(response.headers["Cache-Control"])
            if not_modified:
                return not_modified
            return await unlocked_chapter_response(False)
        return {
            "locked": True,
            "reason": "INSUFFICIENT_READING_TIME",
            "message": "Your reading time has ended. Top up to continue reading.",
            "chapter": meta,
        }

    # Guest
    return {
        "locked": True,
        "reason": "AUTH_REQUIRED",
        "message": "Sign in and add reading time to continue.",
        "chapter": meta,
    }


# ---------- Admin: Reader Users + Wallet ----------
@api.get("/admin/users")
async def admin_list_users(_=Depends(require_admin)):
    rows = await db.users.find(
        {"role": "user"},
        {"_id": 0, "password_hash": 0},
    ).sort("created_at", -1).to_list(2000)
    return [_user_public(r) for r in rows]


@api.post("/admin/users/{uid}/wallet/adjust")
async def admin_wallet_adjust(uid: str, payload: WalletAdjustIn, admin=Depends(require_admin)):
    user = await db.users.find_one({"id": uid, "role": "user"}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    seconds_delta = int(payload.minutes) * 60
    current = int(user.get("reading_seconds_balance", 0))
    new_balance = max(0, current + seconds_delta)
    actual_delta = new_balance - current
    await db.users.update_one(
        {"id": uid},
        {"$set": {"reading_seconds_balance": new_balance, "wallet_seconds": new_balance}},
    )
    await _record_wallet_ledger(
        user_id=uid,
        action="admin_adjustment",
        seconds_delta=actual_delta,
        reason=payload.reason.strip() or ("Admin top-up" if actual_delta >= 0 else "Admin deduction"),
        actor=f"admin:{admin.get('email','')}",
        balance_after=new_balance,
    )
    return {
        "ok": True,
        "user_id": uid,
        "applied_seconds": actual_delta,
        "reading_seconds_balance": new_balance,
    }


@api.get("/admin/users/{uid}/transactions", response_model=List[WalletTransactionOut])
async def admin_user_transactions(uid: str, _=Depends(require_admin)):
    rows = await db.wallet_transactions.find(
        {"user_id": uid}, {"_id": 0}
    ).sort("created_at", -1).to_list(500)
    return rows


@api.get("/admin/users/{uid}/wallet/refund-review")
async def admin_wallet_refund_review(uid: str, lookback_days: int = 30, _=Depends(require_admin)):
    user = await db.users.find_one({"id": uid, "role": "user"}, {"_id": 0, "password_hash": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    days = max(1, min(365, int(lookback_days or 30)))
    cutoff = datetime.now(timezone.utc) - timedelta(days=days)
    txs = await db.wallet_transactions.find({"user_id": uid}, {"_id": 0}).sort("created_at", -1).to_list(2000)
    txs = [
        tx for tx in txs
        if (_as_utc_dt(tx.get("created_at")) or datetime.max.replace(tzinfo=timezone.utc)) >= cutoff
    ]
    refund_rows = await db.wallet_refunds.find({"user_id": uid}, {"_id": 0}).sort("created_at", -1).to_list(500)
    refunded_ids = {row.get("candidate_id") for row in refund_rows if row.get("candidate_id")}
    candidates = _wallet_refund_candidates(txs, refunded_ids)
    stored_balance = int(user.get("reading_seconds_balance", user.get("wallet_seconds", 0)) or 0)
    ledger_balance = await _ledger_balance(uid)
    return {
        "user": _user_public(user),
        "lookback_days": days,
        "stored_balance_seconds": stored_balance,
        "ledger_balance_seconds": ledger_balance,
        "wallet_divergence_seconds": stored_balance - ledger_balance,
        "candidate_count": len(candidates),
        "refundable_seconds": sum(int(c.get("refundable_seconds", 0) or 0) for c in candidates),
        "candidates": candidates,
        "previous_refunds": refund_rows[:25],
    }


@api.post("/admin/users/{uid}/wallet/refund-approve")
async def admin_wallet_refund_approve(uid: str, payload: WalletRefundApproveIn, admin=Depends(require_admin)):
    user = await db.users.find_one({"id": uid, "role": "user"}, {"_id": 0, "password_hash": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    requested = {candidate_id.strip() for candidate_id in payload.candidate_ids if candidate_id.strip()}
    if not requested:
        raise HTTPException(status_code=400, detail="Select at least one refund candidate before approval")

    txs = await db.wallet_transactions.find({"user_id": uid}, {"_id": 0}).sort("created_at", -1).to_list(2000)
    refunded_rows = await db.wallet_refunds.find({"user_id": uid}, {"_id": 0, "candidate_id": 1}).to_list(1000)
    refunded_ids = {row.get("candidate_id") for row in refunded_rows if row.get("candidate_id")}
    open_candidates = _wallet_refund_candidates(txs, refunded_ids)
    candidates_by_id = {c["candidate_id"]: c for c in open_candidates}
    missing = sorted(requested - set(candidates_by_id))
    if missing:
        raise HTTPException(status_code=409, detail={
            "message": "Some selected candidates were already refunded or are no longer eligible.",
            "candidate_ids": missing,
        })

    batch_id = str(uuid.uuid4())
    now = now_iso()
    approved: List[dict] = []
    for candidate_id in sorted(requested):
        candidate = candidates_by_id[candidate_id]
        refund_doc = {
            "id": str(uuid.uuid4()),
            "batch_id": batch_id,
            "candidate_id": candidate_id,
            "user_id": uid,
            "status": "approved",
            "source_transaction_id": candidate.get("source_transaction_id", ""),
            "session_id": candidate.get("session_id", ""),
            "issue": candidate.get("issue", ""),
            "charged_seconds": int(candidate.get("charged_seconds", 0) or 0),
            "refunded_seconds": int(candidate.get("refundable_seconds", 0) or 0),
            "evidence": candidate.get("evidence", ""),
            "admin_note": payload.note.strip()[:500],
            "approved_by": f"admin:{admin.get('email','')}",
            "created_at": now,
            "credited": False,
        }
        result = await db.wallet_refunds.update_one(
            {"candidate_id": candidate_id},
            {"$setOnInsert": refund_doc},
            upsert=True,
        )
        if result.upserted_id is not None:
            approved.append(refund_doc)

    applied_seconds = sum(int(row.get("refunded_seconds", 0) or 0) for row in approved)
    if applied_seconds <= 0:
        fresh = await db.users.find_one({"id": uid}, {"_id": 0, "reading_seconds_balance": 1, "wallet_seconds": 1}) or {}
        return {
            "ok": True,
            "applied_seconds": 0,
            "applied_candidates": [],
            "reading_seconds_balance": int(fresh.get("reading_seconds_balance", fresh.get("wallet_seconds", 0)) or 0),
            "message": "Selected candidates were already approved.",
        }

    await db.users.update_one(
        {"id": uid},
        {"$inc": {"reading_seconds_balance": applied_seconds, "wallet_seconds": applied_seconds}},
    )
    fresh = await db.users.find_one({"id": uid}, {"_id": 0, "reading_seconds_balance": 1, "wallet_seconds": 1}) or {}
    balance_after = int(fresh.get("reading_seconds_balance", fresh.get("wallet_seconds", 0)) or 0)
    tx = await _record_wallet_ledger(
        user_id=uid,
        action="refund_credit",
        seconds_delta=applied_seconds,
        reason=f"Admin-approved billing refund ({len(approved)} finding{'s' if len(approved) != 1 else ''})",
        actor=f"admin:{admin.get('email','')}",
        balance_after=balance_after,
        extra={
            "refund_batch_id": batch_id,
            "candidate_ids": ",".join(row["candidate_id"] for row in approved),
            "admin_note": payload.note.strip()[:240],
        },
    )
    await db.wallet_refunds.update_many(
        {"batch_id": batch_id, "candidate_id": {"$in": [row["candidate_id"] for row in approved]}},
        {"$set": {"credited": True, "credited_transaction_id": tx["id"], "balance_after": balance_after}},
    )
    return {
        "ok": True,
        "batch_id": batch_id,
        "applied_seconds": applied_seconds,
        "applied_candidates": approved,
        "transaction_id": tx["id"],
        "reading_seconds_balance": balance_after,
    }


@api.patch("/admin/users/{uid}/status")
async def admin_user_status(uid: str, payload: UserStatusIn, _=Depends(require_admin)):
    if payload.status not in {"active", "blocked"}:
        raise HTTPException(status_code=400, detail="Invalid status")
    res = await db.users.update_one({"id": uid, "role": "user"}, {"$set": {"status": payload.status}})
    if res.matched_count == 0:
        raise HTTPException(status_code=404, detail="User not found")
    await _invalidate_user_cache(uid)
    return {"ok": True, "id": uid, "status": payload.status}


# =====================================================================
# Razorpay test-mode wallet top-up
# =====================================================================

# ---------- Public: Pack catalogue ----------
@api.get("/payments/packs", response_model=List[PackOut])
async def payments_list_packs():
    cache_key = _public_cache_key("payment_packs")
    cached = await _public_cache_get(cache_key)
    if cached is not None:
        return cached
    result = [PackOut(**p).model_dump() for p in PACKS]
    await _public_cache_set(cache_key, result)
    return result


@api.get("/payments/config")
async def payments_config():
    """Lightweight config shim used by frontend to know if Razorpay is wired."""
    cache_key = _public_cache_key("payment_config")
    cached = await _public_cache_get(cache_key)
    if cached is not None:
        return cached
    result = {
        "configured": razorpay_keys_configured(),
        "mode": RAZORPAY_MODE,
        "key_id": RAZORPAY_KEY_ID if razorpay_keys_configured() else "",
    }
    await _public_cache_set(cache_key, result)
    return result


# ---------- Wallet credit helper (idempotent) ----------
async def _credit_wallet_for_intent(intent: dict, payment_id: Optional[str], source: str) -> dict:
    """Atomically transition a top-up intent to 'credited' and add minutes to
    the user's wallet. Safe against concurrent webhook + verify calls.
    Returns the refreshed intent.
    """
    if _topup_intent_is_expired(intent):
        await db.topup_intents.update_one(
            {"id": intent["id"], "status": {"$nin": ["credited", "expired"]}},
            {
                "$set": {
                    "status": "expired",
                    "failed_reason": "intent_expired",
                    "expired_at": now_iso(),
                }
            },
        )
        return await db.topup_intents.find_one({"id": intent["id"]}, {"_id": 0}) or {**intent, "status": "expired"}

    # Atomic transition: only ONE caller can flip status from a non-credited
    # state to 'credited'. Prevents double-credit when the verify endpoint and
    # the webhook both fire for the same payment.
    set_doc: dict = {
        "status": "credited",
        "credited_at": now_iso(),
        "credited_by": source,  # "verify" | "webhook" | "admin_reconcile"
    }
    if payment_id:
        set_doc["razorpay_payment_id"] = payment_id
    res = await db.topup_intents.update_one(
        {"id": intent["id"], "status": {"$ne": "credited"}},
        {"$set": set_doc},
    )
    if res.modified_count != 1:
        # Already credited by another path — return the existing record.
        return await db.topup_intents.find_one({"id": intent["id"]}, {"_id": 0}) or intent

    seconds = int(intent["minutes"]) * 60
    user_id = intent["user_id"]
    # Race-safe credit using $inc.
    await db.users.update_one(
        {"id": user_id, "role": "user"},
        {"$inc": {"reading_seconds_balance": seconds, "wallet_seconds": seconds}},
    )
    fresh = await db.users.find_one({"id": user_id}, {"_id": 0, "reading_seconds_balance": 1, "wallet_seconds": 1}) or {}
    balance_after = int(fresh.get("reading_seconds_balance", fresh.get("wallet_seconds", 0)) or 0)
    await _record_wallet_ledger(
        user_id=user_id,
        action="topup_credit",
        seconds_delta=seconds,
        reason=f"Razorpay top-up · {intent.get('pack_id')} · {intent['minutes']} min",
        actor=f"razorpay:{source}",
        balance_after=balance_after,
        extra={"topup_intent_id": intent["id"]},
    )
    return await db.topup_intents.find_one({"id": intent["id"]}, {"_id": 0}) or {**intent, **set_doc}


# ---------- Reader: create a top-up intent + Razorpay order ----------
@api.post("/payments/topup", response_model=TopUpCreateOut)
async def payments_create_topup(payload: TopUpCreateIn, user=Depends(require_user)):
    pack = PACKS_BY_ID.get(payload.pack_id)
    if not pack:
        raise HTTPException(status_code=400, detail="Unknown pack")
    if not razorpay_keys_configured():
        raise HTTPException(
            status_code=503,
            detail="Razorpay is not configured yet. Add RAZORPAY_KEY_ID and RAZORPAY_KEY_SECRET to enable payments.",
        )
    client = get_razorpay_client()
    if client is None:
        raise HTTPException(status_code=503, detail="Razorpay SDK is unavailable on this server.")

    intent_id = str(uuid.uuid4())
    receipt = f"earn-{intent_id[:24]}"  # Razorpay receipt must be <= 40 chars

    try:
        order = client.order.create({
            "amount": pack["amount_paise"],
            "currency": "INR",
            "receipt": receipt,
            "payment_capture": 1,
            "notes": {
                "intent_id": intent_id,
                "user_id": user["id"],
                "pack_id": pack["id"],
                "minutes": str(pack["minutes"]),
            },
        })
    except Exception as exc:  # razorpay.errors.* — swallow to a clean 502
        logger.exception("Razorpay order.create failed")
        raise HTTPException(status_code=502, detail=f"Could not start payment: {exc}")

    await db.topup_intents.insert_one({
        "id": intent_id,
        "user_id": user["id"],
        "user_email": user.get("email", ""),
        "pack_id": pack["id"],
        "minutes": pack["minutes"],
        "amount_paise": pack["amount_paise"],
        "currency": "INR",
        "razorpay_order_id": order["id"],
        "razorpay_payment_id": None,
        "status": "created",  # created -> paid -> credited (or failed)
        "mode": RAZORPAY_MODE,
        "created_at": now_iso(),
        "expires_at": topup_intent_expires_at(),
        "credited_at": None,
        "credited_by": None,
    })
    await _invalidate_user_cache(user["id"])

    return TopUpCreateOut(
        intent_id=intent_id,
        razorpay_order_id=order["id"],
        key_id=RAZORPAY_KEY_ID,
        amount=pack["amount_paise"],
        currency="INR",
        name="The Earnalism Digital Library",
        description=f'{pack["label"]} · {pack["minutes"]} minutes',
        pack=PackOut(**pack),
        prefill={
            "name": user.get("name", ""),
            "email": user.get("email", ""),
        },
    )


# ---------- Reader: verify Razorpay checkout signature & credit ----------
@api.post("/payments/verify")
async def payments_verify(payload: PaymentVerifyIn, user=Depends(require_user)):
    intent = await db.topup_intents.find_one(
        {"razorpay_order_id": payload.razorpay_order_id, "user_id": user["id"]},
        {"_id": 0},
    )
    if not intent:
        raise HTTPException(status_code=404, detail="Top-up intent not found")
    if not razorpay_keys_configured():
        raise HTTPException(status_code=503, detail="Razorpay is not configured")

    # Verify HMAC-SHA256(order_id|payment_id, KEY_SECRET)
    expected = _hmac_sha256_hex(
        RAZORPAY_KEY_SECRET,
        f"{payload.razorpay_order_id}|{payload.razorpay_payment_id}".encode("utf-8"),
    )
    if not hmac.compare_digest(expected, payload.razorpay_signature):
        await db.topup_intents.update_one({"id": intent["id"]}, {"$set": {"status": "failed", "failed_reason": "bad_signature"}})
        await _invalidate_user_cache(user["id"])
        raise HTTPException(status_code=400, detail="Invalid payment signature")

    refreshed = await _credit_wallet_for_intent(intent, payload.razorpay_payment_id, "verify")
    if refreshed.get("status") != "credited":
        await _invalidate_user_cache(user["id"])
        raise HTTPException(status_code=409, detail="Top-up intent is expired or not creditable")
    return {
        "ok": True,
        "intent": refreshed,
        "reading_seconds_balance": await _cached_user_wallet_seconds(user["id"]),
    }


# ---------- Razorpay webhook (authoritative) ----------
@api.post("/payments/webhook")
async def payments_webhook(request: Request):
    raw_body = await request.body()
    signature = request.headers.get("X-Razorpay-Signature", "")

    if not RAZORPAY_WEBHOOK_SECRET:
        # Refuse to silently accept unsigned events.
        raise HTTPException(status_code=503, detail="Webhook secret not configured")
    expected = _hmac_sha256_hex(RAZORPAY_WEBHOOK_SECRET, raw_body)
    if not hmac.compare_digest(expected, signature):
        # Store the rejected event so admins can audit attempts.
        try:
            await db.payment_webhook_events.insert_one({
                "id": str(uuid.uuid4()),
                "event_id": f"rejected:{uuid.uuid4()}",
                "event": "unknown",
                "status": "rejected_bad_signature",
                "raw": raw_body.decode("utf-8", errors="replace")[:8000],
                "created_at": now_iso(),
            })
        except Exception:
            logger.warning("Could not persist rejected Razorpay webhook audit event", exc_info=True)
        raise HTTPException(status_code=400, detail="Invalid webhook signature")

    try:
        body = _json.loads(raw_body)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON")

    event = body.get("event", "")
    event_id = request.headers.get("X-Razorpay-Event-Id") or body.get("id")
    payment = (body.get("payload") or {}).get("payment", {}).get("entity") or {}
    payment_id = payment.get("id")
    order_id = payment.get("order_id")

    # Idempotency: reject duplicate event ids quickly.
    if event_id:
        seen = await db.payment_webhook_events.find_one({"event_id": event_id}, {"_id": 0})
        if seen:
            try:
                await db.payment_webhook_events.insert_one({
                    "id": str(uuid.uuid4()),
                    "event_id": f"duplicate:{uuid.uuid4()}",
                    "event": event or seen.get("event", ""),
                    "status": "duplicate_replay_blocked",
                    "event_id_hash": hashlib.sha256(str(event_id).encode("utf-8")).hexdigest()[:24],
                    "created_at": now_iso(),
                })
            except Exception:
                logger.warning("Could not persist duplicate Razorpay webhook metric", exc_info=True)
            return {"ok": True, "duplicate": True, "event_id": event_id}

    log_doc = {
        "id": str(uuid.uuid4()),
        "event_id": event_id,
        "event": event,
        "status": "received",
        "razorpay_order_id": order_id,
        "razorpay_payment_id": payment_id,
        "raw": raw_body.decode("utf-8", errors="replace")[:8000],
        "created_at": now_iso(),
    }

    intent = None
    if order_id:
        intent = await db.topup_intents.find_one({"razorpay_order_id": order_id}, {"_id": 0})

    if event == "payment.captured" and intent:
        refreshed = await _credit_wallet_for_intent(intent, payment_id, "webhook")
        log_doc["status"] = "credited" if refreshed.get("status") == "credited" else refreshed.get("status", "ignored")
    elif event == "payment.failed" and intent:
        await db.topup_intents.update_one(
            {"id": intent["id"], "status": "created"},
            {"$set": {"status": "failed", "razorpay_payment_id": payment_id, "failed_reason": payment.get("error_description", "payment_failed")}},
        )
        await _invalidate_user_cache(intent.get("user_id", ""))
        log_doc["status"] = "marked_failed"
    elif not intent:
        log_doc["status"] = "unmatched_intent"
    else:
        log_doc["status"] = "ignored"

    await db.payment_webhook_events.insert_one(log_doc)
    return {"ok": True, "event": event, "status": log_doc["status"]}


# ---------- Dev-only: simulate a webhook to exercise the credit flow ----------
@api.post("/payments/_simulate_topup")
async def payments_simulate_topup(payload: TopUpCreateIn, user=Depends(require_user)):
    """Dev-only mirror of /payments/topup that does NOT call Razorpay. Creates
    a top-up intent with a synthetic order id so the credit path can be
    exercised end-to-end before real keys are wired up.

    Disabled when RAZORPAY_MODE is not 'test'.
    """
    if RAZORPAY_MODE != "test":
        raise HTTPException(status_code=403, detail="Simulator disabled outside test mode")
    pack = PACKS_BY_ID.get(payload.pack_id)
    if not pack:
        raise HTTPException(status_code=400, detail="Unknown pack")
    intent_id = str(uuid.uuid4())
    fake_order_id = f"order_test_{uuid.uuid4().hex[:18]}"
    await db.topup_intents.insert_one({
        "id": intent_id,
        "user_id": user["id"],
        "user_email": user.get("email", ""),
        "pack_id": pack["id"],
        "minutes": pack["minutes"],
        "amount_paise": pack["amount_paise"],
        "currency": "INR",
        "razorpay_order_id": fake_order_id,
        "razorpay_payment_id": None,
        "status": "created",
        "mode": "test_simulated",
        "created_at": now_iso(),
        "expires_at": topup_intent_expires_at(),
        "credited_at": None,
        "credited_by": None,
    })
    await _invalidate_user_cache(user["id"])
    return {
        "intent_id": intent_id,
        "razorpay_order_id": fake_order_id,
        "amount": pack["amount_paise"],
        "currency": "INR",
        "pack": PackOut(**pack).model_dump(),
        "simulated": True,
    }


@api.post("/payments/_simulate_webhook")
async def payments_simulate_webhook(
    intent_id: str,
    user=Depends(require_user),
):
    """Test helper. Only enabled when RAZORPAY_MODE=='test'. Lets a logged-in
    user (or admin via curl) drive an intent to 'credited' WITHOUT real
    Razorpay — useful when keys are not configured yet.
    """
    if RAZORPAY_MODE != "test":
        raise HTTPException(status_code=403, detail="Simulator disabled outside test mode")
    intent = await db.topup_intents.find_one({"id": intent_id}, {"_id": 0})
    if not intent:
        raise HTTPException(status_code=404, detail="Intent not found")
    if intent["user_id"] != user["id"]:
        raise HTTPException(status_code=403, detail="Not your intent")
    if intent["status"] == "credited":
        return {"ok": True, "duplicate": True, "intent": intent}
    fake_payment_id = f"pay_test_{uuid.uuid4().hex[:20]}"
    refreshed = await _credit_wallet_for_intent(intent, fake_payment_id, "simulate")
    refreshed_status = refreshed.get("status", "unknown")
    await db.payment_webhook_events.insert_one({
        "id": str(uuid.uuid4()),
        "event_id": f"evt_sim_{uuid.uuid4().hex[:20]}",
        "event": "payment.captured",
        "status": "credited_simulated" if refreshed_status == "credited" else f"{refreshed_status}_simulated",
        "razorpay_order_id": intent.get("razorpay_order_id"),
        "razorpay_payment_id": fake_payment_id,
        "raw": "{simulated}",
        "created_at": now_iso(),
    })
    return {
        "ok": True,
        "simulated": True,
        "intent": refreshed,
        "reading_seconds_balance": await _cached_user_wallet_seconds(user["id"]),
    }


# ---------- Reader: own top-up history ----------
@api.get("/payments/me/intents")
async def payments_my_intents(user=Depends(require_user)):
    cache_key = _user_payment_intents_cache_id(user["id"])
    cached = await _redis_cache_get("user-private", cache_key)
    if cached is not None:
        return cached
    rows = await db.topup_intents.find({"user_id": user["id"]}, {"_id": 0}).sort("created_at", -1).to_list(100)
    await _redis_cache_set("user-private", cache_key, rows, USER_PAYMENT_INTENTS_CACHE_TTL_SECONDS)
    return rows


# ---------- Admin: payments dashboard ----------
@api.get("/admin/payments/intents")
async def admin_list_intents(_=Depends(require_admin)):
    rows = await db.topup_intents.find({}, {"_id": 0}).sort("created_at", -1).to_list(2000)
    return rows


@api.get("/admin/payments/webhooks")
async def admin_list_webhooks(_=Depends(require_admin)):
    rows = await db.payment_webhook_events.find({}, {"_id": 0}).sort("created_at", -1).to_list(2000)
    return rows


@api.get("/admin/secure-reader/alerts")
async def admin_secure_reader_alerts(_=Depends(require_admin)):
    rows = await db.reader_security_events.find(
        {"event_type": {"$in": sorted(SECURE_READER_RECORDED_EVENTS)}},
        {"_id": 0},
    ).sort("created_at", -1).to_list(500)

    grouped: Dict[str, dict] = {}
    for row in rows:
        key = row.get("session_id") or row.get("id")
        alert = grouped.setdefault(key, {
            "session_id": key,
            "user_email": row.get("user_email", ""),
            "book_slug": row.get("book_slug", ""),
            "chapter_id": row.get("chapter_id", ""),
            "latest_at": row.get("created_at"),
            "total_attempts": 0,
            "events": {},
        })
        alert["total_attempts"] += 1
        event_type = row.get("event_type", "unknown")
        alert["events"][event_type] = alert["events"].get(event_type, 0) + 1
        if str(row.get("created_at", "")) > str(alert.get("latest_at", "")):
            alert["latest_at"] = row.get("created_at")

    alerts = sorted(grouped.values(), key=lambda item: item.get("latest_at") or "", reverse=True)
    high_risk = [item for item in alerts if item["total_attempts"] >= 3]
    return {
        "summary": {
            "events": len(rows),
            "sessions": len(alerts),
            "high_risk_sessions": len(high_risk),
        },
        "alerts": alerts,
    }


@api.get("/admin/launch-monitor/summary")
async def admin_launch_monitor_summary(_=Depends(require_admin)):
    return await build_launch_monitor_summary()


@api.post("/admin/payments/intents/{intent_id}/reconcile")
async def admin_reconcile_intent(intent_id: str, payload: PaymentReconcileIn, admin=Depends(require_admin)):
    """Manually mark an intent as credited (e.g. webhook never fired). Idempotent."""
    intent = await db.topup_intents.find_one({"id": intent_id}, {"_id": 0})
    if not intent:
        raise HTTPException(status_code=404, detail="Intent not found")
    if intent["status"] == "credited":
        return {"ok": True, "duplicate": True, "intent": intent}
    refreshed = await _credit_wallet_for_intent(
        intent,
        intent.get("razorpay_payment_id"),
        f"admin_reconcile:{admin.get('email','')}:{payload.note[:80]}",
    )
    return {"ok": True, "intent": refreshed}


def _current_rss_bytes() -> Optional[int]:
    statm = Path("/proc/self/statm")
    if not statm.exists():
        return None
    try:
        parts = statm.read_text(encoding="utf-8").split()
        if len(parts) < 2:
            return None
        page_size = os.sysconf("SC_PAGE_SIZE")
        return int(parts[1]) * int(page_size)
    except Exception:
        return None


def _process_memory_snapshot() -> Dict[str, Any]:
    max_rss_raw = resource.getrusage(resource.RUSAGE_SELF).ru_maxrss
    # Linux reports KiB; macOS reports bytes. Railway runs Linux, but keep local
    # diagnostics readable on developer machines too.
    max_rss_bytes = max_rss_raw if max_rss_raw > 10 * 1024 * 1024 else max_rss_raw * 1024
    return {
        "rss_bytes": _current_rss_bytes(),
        "max_rss_bytes": max_rss_bytes,
        "public_cache_entries": len(_public_cache),
        "rate_limit_buckets": len(_rate_limit_hits),
    }


@api.get("/admin/system/cost-control")
async def admin_system_cost_control(_=Depends(require_admin)):
    return {
        "environment": ENVIRONMENT,
        "uptime_seconds": round(time.monotonic() - _PROCESS_STARTED_AT, 2),
        "memory": _process_memory_snapshot(),
        "flags": _cost_control_flags(),
        "active_background_jobs": int(_expensive_job_state.get("active", 0)),
        "expensive_jobs_started": dict(_expensive_job_state.get("started", {})),
        "expensive_jobs_blocked": dict(_expensive_job_state.get("blocked", {})),
        "runtime_keys_detected": {
            "openai": bool(os.environ.get("OPENAI_API_KEY")),
            "cloudinary": _cloudinary_config_detected(),
            "b2_storage": _b2_is_configured(),
            "redis": bool(REDIS_URL),
        },
        "notes": {
            "secrets_exposed": False,
            "heavy_jobs_default_disabled": (
                not ENABLE_BACKGROUND_WORKERS
                and not ENABLE_AUDIOBOOK_PIPELINE
                and not ENABLE_BOOK_RENDERING_JOBS
                and not ENABLE_COVER_GENERATION
                and not ENABLE_SCHEDULED_JOBS
                and not ENABLE_QUEUE_CONSUMER
            ),
        },
    }


@api.get("/admin/cache/status")
async def admin_cache_status(_=Depends(require_admin)):
    memory: Dict[str, Any] = {}
    keyspace: Dict[str, Any] = {}
    config: Dict[str, Any] = {}
    dbsize: Optional[int] = None
    if _redis_state_enabled():
        try:
            raw_memory = await _redis_client.info("memory")
            memory = {
                key: raw_memory.get(key)
                for key in (
                    "used_memory",
                    "used_memory_human",
                    "used_memory_peak",
                    "used_memory_peak_human",
                    "maxmemory",
                    "maxmemory_human",
                    "maxmemory_policy",
                    "mem_fragmentation_ratio",
                )
                if key in raw_memory
            }
        except Exception:
            logger.warning("Redis memory info failed", exc_info=True)
        try:
            keyspace = await _redis_client.info("keyspace")
        except Exception:
            logger.warning("Redis keyspace info failed", exc_info=True)
        try:
            dbsize = int(await _redis_client.dbsize())
        except Exception:
            logger.warning("Redis dbsize failed", exc_info=True)
        try:
            config = await _redis_client.config_get("maxmemory*")
        except Exception:
            config = {}

    return {
        "enabled": REDIS_CACHE_ENABLED,
        "available": _redis_state_enabled(),
        "multi_replica_enabled": MULTI_REPLICA_ENABLED,
        "key_prefix": REDIS_KEY_PREFIX,
        "timeouts_seconds": {
            "connect": REDIS_SOCKET_CONNECT_TIMEOUT_SECONDS,
            "operation": REDIS_SOCKET_TIMEOUT_SECONDS,
        },
        "ttl_seconds": {
            "public": PUBLIC_CACHE_TTL_SECONDS,
            "user_auth": USER_AUTH_CACHE_TTL_SECONDS,
            "user_session": USER_SESSION_CACHE_TTL_SECONDS,
            "user_wallet": USER_WALLET_CACHE_TTL_SECONDS,
            "user_transactions": USER_TRANSACTIONS_CACHE_TTL_SECONDS,
            "user_payment_intents": USER_PAYMENT_INTENTS_CACHE_TTL_SECONDS,
            "reader_manifest": READER_MANIFEST_CACHE_TTL_SECONDS,
            "reader_book": READER_BOOK_CACHE_TTL_SECONDS,
            "reader_chapter": READER_CHAPTER_CACHE_TTL_SECONDS,
        },
        "compression": {
            "min_payload_bytes": REDIS_CACHE_COMPRESS_MIN_BYTES,
            "ttl_jitter_seconds": REDIS_CACHE_TTL_JITTER_SECONDS,
        },
        "reader_rum": {
            "enabled": READER_RUM_ENABLED,
            "sample_rate": READER_RUM_SAMPLE_RATE,
            "slow_ms": READER_RUM_SLOW_MS,
            "aggregate_ttl_seconds": READER_RUM_AGGREGATE_TTL_SECONDS,
        },
        "stats": dict(_cache_stats),
        "policy": {
            "allowed_payloads": REDIS_CACHE_ALLOWED_PAYLOADS,
            "excluded_payloads": REDIS_CACHE_EXCLUDED_PAYLOADS,
            "media_binary_guard_enabled": True,
            "media_binary_note": (
                "Redis stores metadata, reader manifests, chapter text and short-lived state only. "
                "Book-cover images and audiobook binaries stay on Cloudinary/CDN/browser caches."
            ),
        },
        "redis": {
            "dbsize": dbsize,
            "memory": memory,
            "keyspace": keyspace,
            "config": config,
            "startup_config": _redis_config_status,
            "requested_maxmemory": REDIS_MAXMEMORY,
            "requested_maxmemory_policy": REDIS_MAXMEMORY_POLICY,
        },
        "architecture_note": (
            "This Redis instance is a shared backend cache in the Railway service region. "
            "It lowers MongoDB pressure and cross-replica misses, but India/UK latency still depends on backend placement/CDN strategy."
        ),
    }


app.include_router(api)

cors_origins = set()
frontend_url = os.getenv("FRONTEND_URL", "").strip()
if frontend_url:
    cors_origins.add(frontend_url)
elif ENVIRONMENT != "production":
    cors_origins.add("http://localhost:3000")
if ENVIRONMENT != "production":
    cors_origins.update({"http://localhost:3000", "http://127.0.0.1:3000"})
cors_origins.update(
    origin.strip()
    for origin in os.getenv("CORS_ORIGINS", "").split(",")
    if origin.strip()
)
if ENVIRONMENT == "production" and not cors_origins:
    logger.warning("No CORS origins configured for production")

app.add_middleware(GZipMiddleware, minimum_size=1024)
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=sorted(cors_origins),
    allow_methods=["GET", "POST", "PUT", "PATCH", "DELETE", "OPTIONS"],
    allow_headers=[
        "Authorization",
        "Cache-Control",
        "Content-Type",
        "Pragma",
        "X-Request-ID",
        "X-Requested-With",
        "X-Razorpay-Event-Id",
        "X-Razorpay-Signature",
    ],
    expose_headers=[
        "X-Request-ID",
        "X-Response-Time-ms",
        "Server-Timing",
        "Cache-Control",
        "ETag",
        "X-Reader-Manifest-Version",
        "X-Reader-Chapter-Version",
    ],
)


# ---------- Seed ----------
SEED_CATEGORIES = [
    {"slug": "bengali-classics", "name": "Bengali Classics", "description": "A cultural identity shelf for Bengali literature, short fiction, and clean digital editions for readers in India and beyond.", "order": 1, "image_url": "/assets/shelves/bengali-classics.jpg"},
    {"slug": "literary-fiction", "name": "Literary Fiction", "description": "Enduring novels and modern literary works prepared for focused, thoughtful reading.", "order": 2, "image_url": "/assets/shelves/literary-fiction.jpg"},
    {"slug": "young-readers", "name": "Young Readers", "description": "Classic and modern children's books for younger readers, families, and lifelong rereading.", "order": 3, "image_url": "/assets/shelves/young-readers.jpg"},
    {"slug": "business", "name": "Business & Entrepreneurship", "description": "Founder reads on building, growing, and sustaining ventures.", "order": 4, "image_url": "/assets/shelves/business.jpg"},
    {"slug": "technology", "name": "Technology & AI", "description": "Software, AI, data, digital systems, product thinking, and the future of work.", "order": 5, "image_url": "/assets/shelves/technology-ai.jpg"},
    {"slug": "history-strategy", "name": "History & Strategy", "description": "Modern history, geopolitics, statecraft, diplomacy, and strategic thought.", "order": 6, "image_url": "/assets/shelves/history-strategy.jpg"},
    {"slug": "adventure", "name": "Adventure", "description": "Survival, journeys, and high-stakes stories of movement and courage.", "order": 7, "image_url": "/assets/shelves/adventure.jpg"},
    {"slug": "science-fiction", "name": "Science Fiction", "description": "Speculative classics about invention, time, society, and possible futures.", "order": 8, "image_url": "/assets/shelves/science-fiction.jpg"},
    {"slug": "gothic-fiction", "name": "Gothic Fiction", "description": "Atmospheric classics of fear, invention, mystery, and moral tension.", "order": 9, "image_url": "/assets/shelves/gothic-fiction.jpg"},
]

RETIRED_SEED_BOOK_SLUGS = frozenset({
    "the-architecture-of-intelligent-systems",
    "brownies-to-break-even-and-beyond",
})

RETIRED_PUBLIC_BLOG_SLUGS = frozenset({
    "the-quiet-power-of-a-premium-bookstore-brand",
})

SEED_POSTS = [
    {
        "slug": "why-every-small-business-needs-a-story-before-a-strategy",
        "title": "Why Every Small Business Needs a Story Before a Strategy",
        "excerpt": "Strategy without story is choreography without music. Here's how to find the line that holds your business together.",
        "content": "Most small businesses arrive at strategy too early. They map the funnel before they have found the feeling — the singular line a customer carries home. A story is not a tagline; it is the gravitational center that makes every later decision feel inevitable. When you know the story, pricing becomes posture, marketing becomes memory, and operations become craft.\n\nBegin with the moment you would defend in a quiet room: the reason this venture is worth your unhurried years. Write it down without flourish. Read it aloud. Then build the strategy that protects it.",
        "category": "Business",
        "pull_quote": "Strategy without story is choreography without music.",
        "cover_image_url": "https://images.unsplash.com/photo-1764087957302-ef0756ed8e0a?crop=entropy&cs=srgb&fm=jpg&ixid=M3w3NDk1ODB8MHwxfHNlYXJjaHwxfHxsdXh1cnklMjBmb3VudGFpbiUyMHBlbiUyMHdyaXRpbmclMjBkZXNrfGVufDB8fHx8MTc3NzYxNzE3N3ww&ixlib=rb-4.1.0&q=85",
        "is_published": True,
    },
    {
        "slug": "the-quiet-power-of-a-premium-bookstore-brand",
        "title": "The Quiet Power of a Premium Bookstore Brand",
        "excerpt": "Why restraint, ritual, and refusal are the most underrated assets in running a modern independent bookstore.",
        "content": "A premium brand is not loud — it is precise. It refuses noisy launches, declines easy collaborations, and waits for the right shelf. In bookstores especially, restraint is the most expensive ingredient. The Earnalism is built around the same principle: fewer titles, deeper readings, a longer relationship with each book.\n\nTrust accrues like compound interest. Each careful decision becomes the next reader's reason to return.",
        "category": "Brand",
        "pull_quote": "A premium brand is not loud — it is precise.",
        "cover_image_url": "https://images.unsplash.com/photo-1565357457446-4705a9445e54?crop=entropy&cs=srgb&fm=jpg&ixid=M3w3NTY2NzB8MHwxfHNlYXJjaHwyfHxhbnRpcXVlJTIwYm9vayUyMHBhZ2VzJTIwbWFjcm8lMjBwaG90b2dyYXBoeXxlbnwwfHx8fDE3Nzc2MTcxNzd8MA&ixlib=rb-4.1.0&q=85",
        "is_published": False,
    },
    {
        "slug": "how-reading-shapes-better-founders",
        "title": "How Reading Shapes Better Founders",
        "excerpt": "The founders who endure are not always the loudest readers — they are the most patient ones.",
        "content": "Books slow the founder's mind into the right tempo. They argue with you, comfort you, and rearrange your assumptions while you sleep. Strategy decks fade in a quarter; a great chapter stays for a decade.\n\nReading is not a productivity hack. It is a posture — a way of returning to depth in a culture that rewards velocity. The founders who endure tend to read like they breathe: regularly, unhurriedly, and across worlds.",
        "category": "Self-Growth",
        "pull_quote": "Reading is not a productivity hack. It is a posture.",
        "cover_image_url": "https://images.unsplash.com/photo-1761237771835-1d2914546cea?crop=entropy&cs=srgb&fm=jpg&ixid=M3w3NTY2ODh8MHwxfHNlYXJjaHwzfHxjYWxtJTIwbWluaW1hbCUyMGFic3RyYWN0JTIwd2FybSUyMHRleHR1cmV8ZW58MHx8fHwxNzc3NjE3MTc3fDA&ixlib=rb-4.1.0&q=85",
        "is_published": True,
    },
]
