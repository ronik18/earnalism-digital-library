#!/usr/bin/env python3
"""Review a DOCX manuscript for publishing and compliance risks.

The checks in this script are heuristic. They are designed to catch issues for
human review, not to make final legal, security, authorship, or plagiarism
determinations.
"""

from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import html as html_lib
import json
import os
import re
import sys
import urllib.parse
import urllib.request
from collections import Counter, defaultdict
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from PIL import Image, ImageOps, UnidentifiedImageError


SEVERITY_INFO = "INFO"
SEVERITY_WARNING = "WARNING"
SEVERITY_ERROR = "ERROR"
EMU_PER_INCH = 914400


AI_DIRECT_PATTERNS = [
    r"\bas an ai language model\b",
    r"\bas a large language model\b",
    r"\bi cannot (?:browse|access|provide|assist)\b",
    r"\bi do not have (?:personal|real-time|access)\b",
    r"\bmy training data\b",
    r"\bi am not able to\b",
]

AI_CLICHES = [
    "it is important to note",
    "in conclusion",
    "delve",
    "delving",
    "realm of",
    "tapestry",
    "embark on a journey",
    "unlock the potential",
    "navigate the complexities",
    "in today's fast-paced world",
    "a testament to",
    "seamless integration",
    "robust framework",
    "furthermore",
    "moreover",
    "additionally",
    "ultimately",
]

STRONG_AI_CLICHES = {
    "it is important to note",
    "in conclusion",
    "delve",
    "delving",
    "realm of",
    "tapestry",
    "embark on a journey",
    "unlock the potential",
    "navigate the complexities",
}

TRANSITION_STARTS = [
    "furthermore",
    "moreover",
    "additionally",
    "in addition",
    "on the other hand",
    "in conclusion",
    "ultimately",
    "therefore",
]

CLASSIFICATION_PATTERNS = [
    r"\bTOP SECRET\b",
    r"\bSECRET\b",
    r"\bCONFIDENTIAL\b",
    r"\bFOUO\b",
    r"\bFOR OFFICIAL USE ONLY\b",
    r"\bCUI\b",
    r"\bNOFORN\b",
    r"\bSCI\b",
    r"\bSAP\b",
    r"\bORCON\b",
    r"\bNATO SECRET\b",
    r"\bREL TO\b",
]

PII_PATTERNS = {
    "ssn": r"\b\d{3}-\d{2}-\d{4}\b",
    "email": r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b",
    "phone": r"(?:\+?\d{1,3}[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?)\d{3}[-.\s]?\d{4}\b",
    "passport_like": r"\b[A-PR-WY][0-9]{7,9}\b",
}

INTELLIGENCE_MATERIAL_PATTERNS = [
    r"\bhuman intelligence\b",
    r"\bsignals intelligence\b",
    r"\bsource identity\b",
    r"\btarget package\b",
    r"\brules of engagement\b",
    r"\btroop movement\b",
    r"\boperational plan\b",
    r"\bclassified annex\b",
    r"\bintelligence asset\b",
    r"\bgrid coordinates\b",
]

DEFAMATION_NEGATIVE_PATTERNS = [
    r"\bcommitted fraud\b",
    r"\bis a fraud\b",
    r"\bdefrauded\b",
    r"\bstole\b",
    r"\bembezzled\b",
    r"\bcorrupt\b",
    r"\bcriminal\b",
    r"\bscam\b",
    r"\billegal\b",
    r"\bbribed\b",
    r"\bharassed\b",
    r"\babused\b",
    r"\blied\b",
    r"\bmurdered\b",
    r"\bterrorist\b",
]

DEFAMATION_QUALIFIERS = [
    "allegedly",
    "reportedly",
    "according to",
    "in my opinion",
    "i believe",
    "the complaint alleges",
    "the lawsuit alleges",
    "was accused",
    "has been accused",
    "is accused",
    "claimed",
    "claims",
    "suspected",
]

KNOWN_COMMON_PASSAGES = [
    "it was the best of times it was the worst of times",
    "all happy families are alike each unhappy family is unhappy in its own way",
    "call me ishmael",
    "four score and seven years ago",
    "to be or not to be that is the question",
    "we hold these truths to be self evident",
    "the quick brown fox jumps over the lazy dog",
    "two roads diverged in a yellow wood",
    "shall i compare thee to a summer's day",
]

DEFAULT_CHAPTER_PREFIX = "Chapter"
DEFAULT_PUBLISHER = "Reo Enterprise"
DEFAULT_READER_VIEWPORT = "390x844"
DEFAULT_READER_FONT_SIZE_INDEX = 1
READER_FONT_SIZES = ["15px", "17px", "19px", "21px"]
READER_STYLE_MAP = """
p[style-name='Heading 1'] => h2:fresh
p[style-name='Heading 2'] => h3:fresh
p[style-name='Heading 3'] => h4:fresh
p[style-name='Quote'] => blockquote:fresh
p[style-name='Block Text'] => blockquote:fresh
""".strip()
BOOK_UPLOAD_FORM_FIELDS = [
    "title",
    "subtitle",
    "author",
    "category_slug",
    "short_description",
    "description",
    "cover_image_url",
    "back_cover_image_url",
    "estimated_reading_time",
    "price_paperback",
    "price_ebook",
    "buy_url",
    "formats",
    "benefits",
    "who_for",
    "learnings",
    "about_author",
    "is_published",
]
KNOWN_CATEGORY_KEYWORDS = {
    "business": [
        "business",
        "founder",
        "venture",
        "entrepreneur",
        "startup",
        "pricing",
        "cash flow",
        "compliance",
        "customer",
        "profit",
    ],
    "self-growth": ["habit", "mindset", "personal growth", "discipline", "resilience", "self"],
    "technology": ["technology", "software", "ai", "data", "digital", "platform"],
    "spirituality": ["spiritual", "meditation", "faith", "soul", "devotion"],
    "literature": ["novel", "poem", "story", "literary", "fiction"],
    "bengali": ["বাংলা", "বাঙালি", "অধ্যায়"],
}
CHAPTER_NUMBER_WORDS = {
    "one": 1,
    "two": 2,
    "three": 3,
    "four": 4,
    "five": 5,
    "six": 6,
    "seven": 7,
    "eight": 8,
    "nine": 9,
    "ten": 10,
    "eleven": 11,
    "twelve": 12,
    "thirteen": 13,
    "fourteen": 14,
    "fifteen": 15,
    "sixteen": 16,
    "seventeen": 17,
    "eighteen": 18,
    "nineteen": 19,
    "twenty": 20,
}


def make_finding(
    severity: str,
    message: str,
    paragraph_index: Optional[int] = None,
    excerpt: str = "",
    metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """Create a JSON-serializable finding object."""

    finding: Dict[str, Any] = {
        "severity": severity,
        "message": message,
    }
    if paragraph_index is not None:
        finding["paragraph_index"] = paragraph_index
    if excerpt:
        finding["excerpt"] = excerpt[:500]
    if metadata:
        finding["metadata"] = metadata
    return finding


def normalize_spaces(value: str) -> str:
    """Collapse whitespace for stable pattern matching and report excerpts."""

    return re.sub(r"\s+", " ", value or "").strip()


def normalized_for_match(value: str) -> str:
    """Return lowercase alphanumeric-ish text for phrase matching."""

    return re.sub(r"[^a-z0-9'\s-]", " ", value.lower())


def split_sentences(text: str) -> List[str]:
    """Split text into rough sentence units for paragraph-local reporting."""

    text = normalize_spaces(text)
    if not text:
        return []
    parts = re.split(r"(?<=[.!?।])\s+", text)
    return [part.strip() for part in parts if part.strip()]


def parse_viewport(value: str) -> Tuple[int, int]:
    """Parse a viewport string like 390x844."""

    match = re.match(r"^\s*(\d{3,5})\s*x\s*(\d{3,5})\s*$", value or "", re.IGNORECASE)
    if not match:
        raise ValueError("Reader viewport must use WIDTHxHEIGHT format, for example 390x844")
    return int(match.group(1)), int(match.group(2))


def load_document(path: Path) -> Document:
    """Load a DOCX file with basic extension validation."""

    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")
    if path.suffix.lower() != ".docx":
        raise ValueError("Input file must have a .docx extension")
    return Document(str(path))


def extract_paragraphs(document: Document) -> List[Dict[str, Any]]:
    """Extract paragraph text while preserving original paragraph indices."""

    paragraphs: List[Dict[str, Any]] = []
    for index, paragraph in enumerate(document.paragraphs):
        style_name = paragraph.style.name if paragraph.style else ""
        paragraphs.append(
            {
                "index": index,
                "text": paragraph.text or "",
                "style": style_name,
            }
        )
    return paragraphs


def iter_sentences(paragraphs: Sequence[Dict[str, Any]]) -> Iterable[Tuple[int, str]]:
    """Yield sentence text with its source paragraph index."""

    for paragraph in paragraphs:
        for sentence in split_sentences(paragraph["text"]):
            yield paragraph["index"], sentence


def check_status(findings: Sequence[Dict[str, Any]], fail_on_warning: bool = False) -> str:
    """Convert findings into a PASS/FAIL status."""

    if any(item["severity"] == SEVERITY_ERROR for item in findings):
        return "FAIL"
    if fail_on_warning and any(item["severity"] == SEVERITY_WARNING for item in findings):
        return "FAIL"
    return "PASS"


def check_ai_authorship(paragraphs: Sequence[Dict[str, Any]]) -> Dict[str, Any]:
    """Flag common markers associated with generated or heavily templated text."""

    findings: List[Dict[str, Any]] = []
    sentence_starts: Dict[str, List[Tuple[int, str]]] = defaultdict(list)
    direct_patterns = [re.compile(pattern, re.IGNORECASE) for pattern in AI_DIRECT_PATTERNS]

    for paragraph_index, sentence in iter_sentences(paragraphs):
        lowered = sentence.lower()
        score = 0
        matched_markers: List[str] = []

        for pattern in direct_patterns:
            if pattern.search(sentence):
                score += 4
                matched_markers.append(pattern.pattern)

        for marker in AI_CLICHES:
            if marker in lowered:
                score += 2 if marker in STRONG_AI_CLICHES else 1
                matched_markers.append(marker)

        if any(lowered.startswith(start) for start in TRANSITION_STARTS):
            score += 1
            matched_markers.append("transition_start")

        words = re.findall(r"\b[\w'-]+\b", lowered)
        if len(words) >= 35 and sum(marker in lowered for marker in AI_CLICHES) >= 2:
            score += 1
            matched_markers.append("long_transitional_sentence")

        if len(words) >= 4:
            starter = " ".join(words[:4])
            sentence_starts[starter].append((paragraph_index, sentence))

        if score >= 2:
            severity = SEVERITY_ERROR if score >= 4 else SEVERITY_WARNING
            findings.append(
                make_finding(
                    severity,
                    "Possible AI-generated or templated sentence marker detected.",
                    paragraph_index,
                    sentence,
                    {"score": score, "markers": matched_markers},
                )
            )

    for starter, examples in sentence_starts.items():
        if len(examples) >= 3:
            for paragraph_index, sentence in examples[:5]:
                findings.append(
                    make_finding(
                        SEVERITY_WARNING,
                        "Repetitive sentence opening may indicate templated prose.",
                        paragraph_index,
                        sentence,
                        {"repeated_start": starter, "count": len(examples)},
                    )
                )

    findings.extend(optional_openai_ai_scoring(findings))
    return {
        "status": check_status(findings),
        "summary": f"{len(findings)} AI-authorship review finding(s).",
        "findings": findings,
    }


def optional_openai_ai_scoring(existing_findings: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Optionally ask OpenAI to score heuristic AI-authorship candidates.

    Only candidate excerpts already flagged by local heuristics are sent, and
    only when OPENAI_API_KEY is available. This avoids sending an entire
    manuscript to an external service.
    """

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key or not existing_findings:
        return []

    try:
        from openai import OpenAI  # type: ignore
    except Exception as exc:
        return [
            make_finding(
                SEVERITY_INFO,
                "OPENAI_API_KEY is set, but the openai package is unavailable; skipped optional AI scoring.",
                metadata={"error": str(exc)},
            )
        ]

    excerpts = [
        {
            "paragraph_index": item.get("paragraph_index"),
            "excerpt": item.get("excerpt", ""),
        }
        for item in existing_findings[:20]
        if item.get("excerpt")
    ]
    if not excerpts:
        return []

    try:
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Return compact JSON only. Score each excerpt for likelihood "
                        "of AI-generated prose from 0 to 1. Do not make a final authorship claim."
                    ),
                },
                {
                    "role": "user",
                    "content": json.dumps({"excerpts": excerpts}, ensure_ascii=False),
                },
            ],
        )
        content = response.choices[0].message.content or "{}"
        return [
            make_finding(
                SEVERITY_INFO,
                "Optional OpenAI AI-authorship scoring completed for heuristic candidates.",
                metadata={"model": os.getenv("OPENAI_MODEL", "gpt-4o-mini"), "raw_response": content[:2000]},
            )
        ]
    except Exception as exc:
        return [
            make_finding(
                SEVERITY_INFO,
                "Optional OpenAI AI-authorship scoring failed; local heuristic findings remain available.",
                metadata={"error": str(exc)},
            )
        ]


def check_sensitive_information(paragraphs: Sequence[Dict[str, Any]]) -> Dict[str, Any]:
    """Detect classification markings, PII patterns, and sensitive military terms."""

    findings: List[Dict[str, Any]] = []
    classification_regexes = [re.compile(pattern) for pattern in CLASSIFICATION_PATTERNS]
    pii_regexes = {name: re.compile(pattern, re.IGNORECASE) for name, pattern in PII_PATTERNS.items()}
    intelligence_regexes = [re.compile(pattern, re.IGNORECASE) for pattern in INTELLIGENCE_MATERIAL_PATTERNS]

    for paragraph in paragraphs:
        text = paragraph["text"]
        if not normalize_spaces(text):
            continue

        for pattern in classification_regexes:
            if pattern.search(text):
                findings.append(
                    make_finding(
                        SEVERITY_ERROR,
                        "Government classification or handling marker detected.",
                        paragraph["index"],
                        text,
                        {"pattern": pattern.pattern},
                    )
                )

        for kind, pattern in pii_regexes.items():
            for match in pattern.finditer(text):
                severity = SEVERITY_ERROR if kind in {"ssn", "passport_like"} else SEVERITY_WARNING
                findings.append(
                    make_finding(
                        severity,
                        f"Potential PII detected: {kind}.",
                        paragraph["index"],
                        text,
                        {"match": mask_sensitive(match.group(0)), "kind": kind},
                    )
                )

        for pattern in intelligence_regexes:
            if pattern.search(text):
                findings.append(
                    make_finding(
                        SEVERITY_WARNING,
                        "Possible military or intelligence-sensitive phrase detected.",
                        paragraph["index"],
                        text,
                        {"pattern": pattern.pattern},
                    )
                )

    return {
        "status": check_status(findings, fail_on_warning=True),
        "summary": f"{len(findings)} sensitive-information finding(s).",
        "findings": findings,
    }


def mask_sensitive(value: str) -> str:
    """Mask values in reports while leaving enough shape for review."""

    if len(value) <= 4:
        return "*" * len(value)
    return f"{value[:2]}{'*' * max(2, len(value) - 4)}{value[-2:]}"


def check_defamation_risk(paragraphs: Sequence[Dict[str, Any]]) -> Dict[str, Any]:
    """Flag direct negative factual claims about named entities without qualifiers."""

    findings: List[Dict[str, Any]] = []
    negative_regexes = [re.compile(pattern, re.IGNORECASE) for pattern in DEFAMATION_NEGATIVE_PATTERNS]
    person_pattern = re.compile(r"\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3}\b")
    org_pattern = re.compile(
        r"\b[A-Z][A-Za-z0-9&.'-]*(?:\s+[A-Z][A-Za-z0-9&.'-]*){0,5}\s+"
        r"(?:Inc\.?|LLC|Ltd\.?|Corporation|Corp\.?|Company|University|Bank|Foundation|Agency|Ministry|Government)\b"
    )

    for paragraph_index, sentence in iter_sentences(paragraphs):
        lower = sentence.lower()
        matched = [pattern.pattern for pattern in negative_regexes if pattern.search(sentence)]
        if not matched:
            continue

        entities = sorted(set(person_pattern.findall(sentence) + org_pattern.findall(sentence)))
        if not entities:
            continue

        if any(qualifier in lower for qualifier in DEFAMATION_QUALIFIERS):
            continue

        findings.append(
            make_finding(
                SEVERITY_WARNING,
                "Potential defamation risk: direct negative factual claim about a named person or organization without qualifying language.",
                paragraph_index,
                sentence,
                {"entities": entities, "negative_patterns": matched},
            )
        )

    return {
        "status": check_status(findings, fail_on_warning=True),
        "summary": f"{len(findings)} defamation-risk finding(s).",
        "findings": findings,
    }


def extract_unique_ngrams(text: str, minimum: int = 8, maximum: int = 12, limit: int = 50) -> List[str]:
    """Extract unique n-grams useful for manual plagiarism searches."""

    tokens = re.findall(r"[a-z0-9']+", normalized_for_match(text))
    counter: Counter[Tuple[str, ...]] = Counter()
    for n in range(minimum, maximum + 1):
        for index in range(0, max(0, len(tokens) - n + 1)):
            counter[tuple(tokens[index : index + n])] += 1
    unique = [" ".join(words) for words, count in counter.items() if count == 1]
    unique.sort(key=lambda value: (-len(value.split()), value))
    return unique[:limit]


def check_plagiarism_surface(full_text: str) -> Dict[str, Any]:
    """Run offline surface checks and optional external plagiarism integration."""

    findings: List[Dict[str, Any]] = []
    normalized = " ".join(re.findall(r"[a-z0-9']+", normalized_for_match(full_text)))
    unique_ngrams = extract_unique_ngrams(full_text)

    for passage in KNOWN_COMMON_PASSAGES:
        if passage in normalized:
            findings.append(
                make_finding(
                    SEVERITY_WARNING,
                    "Known common literary or factual passage detected; verify quotation, permission, or public-domain status.",
                    excerpt=passage,
                    metadata={"matched_passage": passage},
                )
            )

    findings.append(
        make_finding(
            SEVERITY_INFO,
            "Extracted statistically unique phrases for manual plagiarism searches.",
            metadata={"sample_ngrams": unique_ngrams[:20], "total_returned": len(unique_ngrams)},
        )
    )

    external = optional_copyscape_check(full_text)
    if external:
        findings.append(external)
    else:
        findings.append(
            make_finding(
                SEVERITY_INFO,
                "External plagiarism API check skipped. Set COPYSCAPE_USERNAME, COPYSCAPE_API_KEY, and DOCX_REVIEW_EXTERNAL_CHECKS=1 to enable.",
            )
        )

    return {
        "status": check_status(findings, fail_on_warning=True),
        "summary": f"{len(findings)} plagiarism-surface finding(s), including informational search phrases.",
        "findings": findings,
    }


def optional_copyscape_check(full_text: str) -> Optional[Dict[str, Any]]:
    """Run a basic Copyscape text search when explicitly enabled.

    The check sends manuscript text to Copyscape. It only runs when
    DOCX_REVIEW_EXTERNAL_CHECKS=1 and both COPYSCAPE_USERNAME and
    COPYSCAPE_API_KEY are set.
    """

    if os.getenv("DOCX_REVIEW_EXTERNAL_CHECKS") != "1":
        return None
    username = os.getenv("COPYSCAPE_USERNAME")
    api_key = os.getenv("COPYSCAPE_API_KEY")
    if not username or not api_key:
        return make_finding(
            SEVERITY_INFO,
            "External plagiarism checks requested, but Copyscape credentials are incomplete.",
        )

    # Keep payload bounded so the script does not accidentally send a full book.
    sample = full_text[:10000]
    data = urllib.parse.urlencode(
        {
            "u": username,
            "k": api_key,
            "o": "csearch",
            "t": sample,
        }
    ).encode("utf-8")
    try:
        request = urllib.request.Request("https://www.copyscape.com/api/", data=data)
        with urllib.request.urlopen(request, timeout=30) as response:
            body = response.read().decode("utf-8", errors="replace")
        severity = SEVERITY_WARNING if "<result>" in body.lower() else SEVERITY_INFO
        return make_finding(
            severity,
            "Copyscape API check completed.",
            metadata={"response_excerpt": body[:2000]},
        )
    except Exception as exc:
        return make_finding(
            SEVERITY_WARNING,
            "Copyscape API check failed; perform manual plagiarism review.",
            metadata={"error": str(exc)},
        )


def check_publishing_compliance(paragraphs: Sequence[Dict[str, Any]]) -> Dict[str, Any]:
    """Check front matter, ISBN, legal notice, and chapter heading consistency."""

    findings: List[Dict[str, Any]] = []
    nonempty = [item for item in paragraphs if normalize_spaces(item["text"])]
    full_text = "\n".join(item["text"] for item in paragraphs)
    lower = full_text.lower()

    if not nonempty:
        findings.append(make_finding(SEVERITY_ERROR, "Document contains no readable text."))
    else:
        first_text = normalize_spaces(nonempty[0]["text"])
        first_style = nonempty[0]["style"].lower()
        if "title" not in first_style and len(first_text.split()) > 20:
            findings.append(
                make_finding(
                    SEVERITY_WARNING,
                    "No clear title page detected in the first non-empty paragraph.",
                    nonempty[0]["index"],
                    first_text,
                    {"style": nonempty[0]["style"]},
                )
            )

    required_sections = {
        "copyright page": [r"\bcopyright\b", r"©", r"\ball rights reserved\b"],
        "dedication": [r"\bdedication\b", r"\bdedicated to\b"],
        "table of contents": [r"\btable of contents\b", r"^\s*contents\s*$"],
        "disclaimer/legal notice": [r"\bdisclaimer\b", r"\blegal notice\b", r"\bnot legal advice\b", r"\bfor informational purposes\b"],
    }

    for section, patterns in required_sections.items():
        if not any(re.search(pattern, lower, re.MULTILINE | re.IGNORECASE) for pattern in patterns):
            findings.append(
                make_finding(
                    SEVERITY_WARNING,
                    f"Missing or unclear {section}.",
                    metadata={"section": section},
                )
            )

    if not re.search(r"\bISBN(?:-1[03])?\b", full_text, re.IGNORECASE):
        findings.append(
            make_finding(
                SEVERITY_WARNING,
                "Missing ISBN or ISBN placeholder.",
                metadata={"expected": "ISBN, ISBN-10, or ISBN-13"},
            )
        )

    chapter_headings = find_chapter_headings(paragraphs)
    styles = sorted({item["style"] for item in chapter_headings if item["style"]})
    if len(styles) > 1:
        findings.append(
            make_finding(
                SEVERITY_WARNING,
                "Inconsistent chapter heading styles detected.",
                metadata={"styles": styles, "heading_count": len(chapter_headings)},
            )
        )
    if chapter_headings and any(item["style"].lower() == "normal" for item in chapter_headings):
        findings.append(
            make_finding(
                SEVERITY_WARNING,
                "One or more chapter headings use Normal style instead of a heading style.",
                metadata={"heading_count": len(chapter_headings)},
            )
        )

    return {
        "status": check_status(findings, fail_on_warning=True),
        "summary": f"{len(findings)} publishing-compliance finding(s).",
        "findings": findings,
    }


def parse_chapter_number(value: str) -> Optional[int]:
    """Parse a chapter number from Arabic digits or simple English words."""

    value = normalize_spaces(value).lower()
    if value.isdigit():
        return int(value)
    return CHAPTER_NUMBER_WORDS.get(value)


def chapter_heading_regex(chapter_prefix: str = DEFAULT_CHAPTER_PREFIX) -> re.Pattern[str]:
    """Build the expected chapter heading pattern, such as Chapter 1."""

    prefix = re.escape(chapter_prefix.strip() or DEFAULT_CHAPTER_PREFIX)
    return re.compile(
        rf"^\s*{prefix}\s+(\d+|{'|'.join(CHAPTER_NUMBER_WORDS.keys())})"
        r"(?:\s*[:.\-–—]\s*(.+)|\s+(.+))?\s*$",
        re.IGNORECASE,
    )


def looks_like_toc_chapter_line(text: str) -> bool:
    """Detect chapter references that are probably table-of-contents entries."""

    compact = normalize_spaces(text)
    if "\t" in text or re.search(r"\.{2,}\s*\d+\s*$", text):
        return True
    if len(compact.split()) >= 5 and re.search(r"\s\d{1,4}\s*$", compact):
        return True
    return False


def heading_level(style_name: str) -> Optional[int]:
    """Return the numeric Word heading level for styles like Heading 1."""

    match = re.search(r"\bheading\s+([1-9])\b", style_name or "", re.IGNORECASE)
    return int(match.group(1)) if match else None


def find_numbered_chapter_headings(
    paragraphs: Sequence[Dict[str, Any]],
    chapter_prefix: str = DEFAULT_CHAPTER_PREFIX,
) -> List[Dict[str, Any]]:
    """Find manuscript chapter headings using Chapter 1, Chapter 2, etc."""

    pattern = chapter_heading_regex(chapter_prefix)
    candidates: List[Dict[str, Any]] = []
    for paragraph in paragraphs:
        text = normalize_spaces(paragraph["text"])
        if not text or looks_like_toc_chapter_line(paragraph["text"]):
            continue
        match = pattern.match(text)
        if not match:
            continue
        number = parse_chapter_number(match.group(1))
        if not number:
            continue
        title = normalize_spaces(match.group(2) or match.group(3) or "")
        candidates.append(
            {
                "number": number,
                "title": title,
                "heading": text,
                "paragraph_index": paragraph["index"],
                "style": paragraph["style"] or "",
                "uses_heading_style": "heading" in (paragraph["style"] or "").lower(),
                "heading_level": heading_level(paragraph["style"] or ""),
            }
        )

    styled_candidates = [item for item in candidates if item["uses_heading_style"]]
    if styled_candidates:
        levels = [item["heading_level"] for item in styled_candidates if item["heading_level"] is not None]
        if levels:
            best_level = min(levels)
            return [item for item in styled_candidates if item["heading_level"] == best_level]
        return styled_candidates
    return candidates


def count_words(text: str) -> int:
    """Count rough word tokens across English and Unicode manuscripts."""

    return len(re.findall(r"\b[\w'-]+\b", text, flags=re.UNICODE))


def build_chapter_map(
    paragraphs: Sequence[Dict[str, Any]],
    headings: Sequence[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    """Build chapter ranges and rough word counts from detected headings."""

    chapters: List[Dict[str, Any]] = []
    heading_indices = [item["paragraph_index"] for item in headings]
    for index, heading in enumerate(headings):
        start = heading["paragraph_index"]
        end = heading_indices[index + 1] if index + 1 < len(heading_indices) else (paragraphs[-1]["index"] + 1 if paragraphs else start + 1)
        body_paragraphs = [
            paragraph
            for paragraph in paragraphs
            if start < paragraph["index"] < end and normalize_spaces(paragraph["text"])
        ]
        word_count = sum(count_words(paragraph["text"]) for paragraph in body_paragraphs)
        chapters.append(
            {
                "number": heading["number"],
                "title": heading["title"],
                "heading": heading["heading"],
                "heading_paragraph_index": start,
                "style": heading["style"],
                "uses_heading_style": heading["uses_heading_style"],
                "heading_level": heading.get("heading_level"),
                "body_paragraph_count": len(body_paragraphs),
                "word_count": word_count,
                "estimated_reading_minutes": max(1, round(word_count / 200)) if word_count else 0,
            }
        )
    return chapters


def check_full_book_chapter_structure(
    paragraphs: Sequence[Dict[str, Any]],
    chapter_prefix: str = DEFAULT_CHAPTER_PREFIX,
    min_chapter_words: int = 50,
) -> Dict[str, Any]:
    """Validate a full-book manuscript with chronological Chapter 1, Chapter 2 headings."""

    findings: List[Dict[str, Any]] = []
    headings = find_numbered_chapter_headings(paragraphs, chapter_prefix)
    chapters = build_chapter_map(paragraphs, headings)

    if not chapters:
        findings.append(
            make_finding(
                SEVERITY_ERROR,
                f"No numbered chapter headings found. Use headings like '{chapter_prefix} 1', '{chapter_prefix} 2', etc.",
            )
        )
        return {
            "status": "FAIL",
            "summary": "No full-book chapter structure detected.",
            "findings": findings,
            "chapters": [],
        }

    numbers = [chapter["number"] for chapter in chapters]
    expected = list(range(1, max(numbers) + 1))
    duplicates = sorted({number for number in numbers if numbers.count(number) > 1})
    missing = [number for number in expected if number not in numbers]

    if numbers[0] != 1:
        findings.append(
            make_finding(
                SEVERITY_ERROR,
                "First detected chapter is not Chapter 1.",
                chapters[0]["heading_paragraph_index"],
                chapters[0]["heading"],
                {"first_detected_number": numbers[0]},
            )
        )

    if numbers != sorted(numbers):
        findings.append(
            make_finding(
                SEVERITY_ERROR,
                "Chapter headings are not in chronological order.",
                metadata={"detected_order": numbers},
            )
        )

    if duplicates:
        findings.append(
            make_finding(
                SEVERITY_ERROR,
                "Duplicate chapter numbers detected.",
                metadata={"duplicates": duplicates, "detected_order": numbers},
            )
        )

    if missing:
        findings.append(
            make_finding(
                SEVERITY_ERROR,
                "Missing chapter numbers detected.",
                metadata={"missing": missing, "detected_order": numbers},
            )
        )

    for chapter in chapters:
        if not chapter["uses_heading_style"]:
            findings.append(
                make_finding(
                    SEVERITY_WARNING,
                    "Chapter heading is detected by text but is not using a Word heading style.",
                    chapter["heading_paragraph_index"],
                    chapter["heading"],
                    {"chapter_number": chapter["number"], "style": chapter["style"] or "unknown"},
                )
            )
        if chapter["word_count"] == 0:
            findings.append(
                make_finding(
                    SEVERITY_ERROR,
                    "Chapter heading has no body text before the next chapter.",
                    chapter["heading_paragraph_index"],
                    chapter["heading"],
                    {"chapter_number": chapter["number"]},
                )
            )
        elif chapter["word_count"] < min_chapter_words:
            findings.append(
                make_finding(
                    SEVERITY_WARNING,
                    "Chapter body is very short; verify that the full chapter content is present.",
                    chapter["heading_paragraph_index"],
                    chapter["heading"],
                    {"chapter_number": chapter["number"], "word_count": chapter["word_count"], "minimum": min_chapter_words},
                )
            )

    if not findings:
        findings.append(
            make_finding(
                SEVERITY_INFO,
                "Full-book chapter sequence is chronological and complete.",
                metadata={"chapter_count": len(chapters), "detected_order": numbers},
            )
        )

    return {
        "status": check_status(findings),
        "summary": f"{len(chapters)} chapter(s) detected using '{chapter_prefix} N' headings.",
        "findings": findings,
        "chapters": chapters,
    }


def find_chapter_headings(paragraphs: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Find likely chapter headings from style names and heading-like text."""

    headings: List[Dict[str, Any]] = []
    pattern = re.compile(r"^\s*(chapter|part|section)\s+([0-9ivxlcdm]+|one|two|three|four|five|six|seven|eight|nine|ten)\b", re.IGNORECASE)
    for paragraph in paragraphs:
        text = normalize_spaces(paragraph["text"])
        style = paragraph["style"] or ""
        if not text:
            continue
        if pattern.search(text) or ("heading" in style.lower() and len(text.split()) <= 20):
            headings.append(paragraph)
    return headings


def local_name(tag: str) -> str:
    """Return the XML local name for a namespaced tag."""

    return tag.rsplit("}", 1)[-1] if "}" in tag else tag


def get_xml_attr(element: Any, name: str) -> Optional[str]:
    """Fetch an XML attribute by local name, ignoring namespace prefixes."""

    for key, value in element.attrib.items():
        if local_name(key) == name:
            return value
    return None


def paragraph_has_image(paragraph: Any) -> bool:
    """Return True when a paragraph contains a DOCX drawing/image."""

    return any(local_name(node.tag) == "drawing" for node in paragraph._element.iter())


def image_extension(content_type: str) -> str:
    """Map image content types to safe file extensions."""

    mapping = {
        "image/jpeg": ".jpg",
        "image/jpg": ".jpg",
        "image/png": ".png",
        "image/webp": ".webp",
        "image/gif": ".gif",
        "image/tiff": ".tif",
        "image/bmp": ".bmp",
    }
    return mapping.get((content_type or "").lower(), ".img")


def extract_images(document: Document) -> List[Dict[str, Any]]:
    """Extract embedded DOCX images and layout metadata from drawing XML."""

    images: List[Dict[str, Any]] = []
    seen = set()

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        for drawing in [node for node in paragraph._element.iter() if local_name(node.tag) == "drawing"]:
            layout = "floating" if any(local_name(node.tag) == "anchor" for node in drawing.iter()) else "inline"
            extent = next((node for node in drawing.iter() if local_name(node.tag) == "extent"), None)
            cx = int(get_xml_attr(extent, "cx") or 0) if extent is not None else 0
            cy = int(get_xml_attr(extent, "cy") or 0) if extent is not None else 0

            for blip in [node for node in drawing.iter() if local_name(node.tag) == "blip"]:
                rel_id = get_xml_attr(blip, "embed") or get_xml_attr(blip, "link")
                if not rel_id:
                    continue
                key = (paragraph_index, rel_id, cx, cy, layout)
                if key in seen:
                    continue
                seen.add(key)
                part = document.part.related_parts.get(rel_id)
                if not part:
                    continue
                blob = getattr(part, "blob", b"")
                images.append(
                    {
                        "paragraph_index": paragraph_index,
                        "relationship_id": rel_id,
                        "layout": layout,
                        "extent_cx": cx,
                        "extent_cy": cy,
                        "content_type": getattr(part, "content_type", ""),
                        "partname": str(getattr(part, "partname", "")),
                        "blob": blob,
                    }
                )

    return images


def check_inline_images(document: Document, previews_dir: Path) -> Dict[str, Any]:
    """Validate embedded image rendering characteristics and create thumbnails."""

    previews_dir.mkdir(parents=True, exist_ok=True)
    findings: List[Dict[str, Any]] = []
    image_reports: List[Dict[str, Any]] = []
    images = extract_images(document)

    if not images:
        findings.append(make_finding(SEVERITY_INFO, "No embedded DOCX images found."))
        return {
            "status": "PASS",
            "summary": "No embedded images found.",
            "findings": findings,
            "images": [],
        }

    findings.append(
        make_finding(
            SEVERITY_INFO,
            "Generated embedded-image thumbnails. Exact DOCX page rendering is not available from python-docx alone.",
        )
    )

    for index, item in enumerate(images, start=1):
        report, image_findings = analyze_image(item, index, previews_dir)
        image_reports.append(report)
        findings.extend(image_findings)

    return {
        "status": check_status(findings),
        "summary": f"{len(images)} embedded image(s) checked.",
        "findings": findings,
        "images": image_reports,
    }


def analyze_image(item: Dict[str, Any], index: int, previews_dir: Path) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
    """Analyze one embedded image and create a preview thumbnail."""

    findings: List[Dict[str, Any]] = []
    blob = item["blob"]
    digest = hashlib.sha256(blob).hexdigest()[:12]
    preview_path = previews_dir / f"image_{index:03d}_paragraph_{item['paragraph_index']}_{digest}.jpg"
    report: Dict[str, Any] = {
        "image_index": index,
        "paragraph_index": item["paragraph_index"],
        "relationship_id": item["relationship_id"],
        "layout": item["layout"],
        "content_type": item["content_type"],
        "partname": item["partname"],
        "preview_path": str(preview_path),
    }

    if item["layout"] == "floating":
        findings.append(
            make_finding(
                SEVERITY_WARNING,
                "Floating image detected; this can shift or overlap in ebook and responsive layouts.",
                item["paragraph_index"],
                metadata={"image_index": index},
            )
        )

    try:
        with Image.open(BytesIO(blob)) as image:
            width, height = image.size
            dpi = image.info.get("dpi") or (0, 0)
            report.update({"width_px": width, "height_px": height, "metadata_dpi": list(dpi)})

            effective_ppi = calculate_effective_ppi(width, height, item["extent_cx"], item["extent_cy"])
            if effective_ppi:
                report["effective_ppi"] = {"x": round(effective_ppi[0], 2), "y": round(effective_ppi[1], 2)}
                min_ppi = min(effective_ppi)
            else:
                min_ppi = min(value for value in dpi if isinstance(value, (int, float))) if dpi and all(dpi) else 0

            if min_ppi and min_ppi < 72:
                findings.append(
                    make_finding(
                        SEVERITY_ERROR,
                        "Image resolution appears below 72 DPI/PPI for screen rendering.",
                        item["paragraph_index"],
                        metadata={"image_index": index, "ppi": round(min_ppi, 2)},
                    )
                )
            elif min_ppi and min_ppi < 300:
                findings.append(
                    make_finding(
                        SEVERITY_WARNING,
                        "Image resolution appears below 300 DPI/PPI for print use.",
                        item["paragraph_index"],
                        metadata={"image_index": index, "ppi": round(min_ppi, 2)},
                    )
                )
            elif not min_ppi:
                findings.append(
                    make_finding(
                        SEVERITY_WARNING,
                        "Image has no usable DPI metadata or display extent; print quality must be reviewed manually.",
                        item["paragraph_index"],
                        metadata={"image_index": index},
                    )
                )

            aspect_warning = validate_aspect_ratio(width, height, item["extent_cx"], item["extent_cy"])
            if aspect_warning:
                findings.append(
                    make_finding(
                        SEVERITY_WARNING,
                        "Displayed image aspect ratio differs from the source image and may be distorted.",
                        item["paragraph_index"],
                        metadata={"image_index": index, **aspect_warning},
                    )
                )

            save_image_preview(image, preview_path)
    except UnidentifiedImageError:
        findings.append(
            make_finding(
                SEVERITY_ERROR,
                "Embedded image could not be opened by Pillow.",
                item["paragraph_index"],
                metadata={"image_index": index, "content_type": item["content_type"]},
            )
        )

    return report, findings


def check_cover_page_structure(document: Document) -> Dict[str, Any]:
    """Validate the convention: first DOCX image is front cover and last image is back cover."""

    findings: List[Dict[str, Any]] = []
    images = extract_images(document)
    content_headings = find_upload_section_starts(extract_paragraphs(document), 0, len(document.paragraphs))
    first_content_index = content_headings[0][0] if content_headings else None

    if not images:
        findings.append(
            make_finding(
                SEVERITY_WARNING,
                "No embedded images found. Expected the first page to contain the front cover and the last page to contain the back cover.",
            )
        )
        return {
            "status": check_status(findings, fail_on_warning=True),
            "summary": "No front/back cover images detected.",
            "findings": findings,
            "covers": {},
        }

    front = images[0]
    back = images[-1] if len(images) > 1 else None
    covers: Dict[str, Any] = {
        "front_cover_paragraph_index": front["paragraph_index"],
        "front_cover_content_type": front["content_type"],
    }

    if first_content_index is not None and front["paragraph_index"] > first_content_index:
        findings.append(
            make_finding(
                SEVERITY_WARNING,
                "First embedded image appears after manuscript content begins; verify front cover placement.",
                front["paragraph_index"],
                metadata={"first_content_paragraph_index": first_content_index},
            )
        )
    else:
        findings.append(
            make_finding(
                SEVERITY_INFO,
                "Front cover candidate detected before manuscript content.",
                front["paragraph_index"],
                metadata={"content_type": front["content_type"]},
            )
        )

    if not back or back["paragraph_index"] == front["paragraph_index"]:
        findings.append(
            make_finding(
                SEVERITY_WARNING,
                "Back cover image was not detected separately. Place the back cover image as the final page/image in the DOCX.",
            )
        )
    else:
        covers.update(
            {
                "back_cover_paragraph_index": back["paragraph_index"],
                "back_cover_content_type": back["content_type"],
            }
        )
        nonempty_after_back = [
            paragraph.text
            for paragraph in document.paragraphs[back["paragraph_index"] + 1 :]
            if normalize_spaces(paragraph.text)
        ]
        if nonempty_after_back:
            findings.append(
                make_finding(
                    SEVERITY_WARNING,
                    "Text appears after the back cover candidate; verify that the last page is only the back cover.",
                    back["paragraph_index"],
                    metadata={"text_after_back_cover_count": len(nonempty_after_back)},
                )
            )
        else:
            findings.append(
                make_finding(
                    SEVERITY_INFO,
                    "Back cover candidate detected at the end of the manuscript.",
                    back["paragraph_index"],
                    metadata={"content_type": back["content_type"]},
                )
            )

    return {
        "status": check_status(findings, fail_on_warning=True),
        "summary": f"{len(images)} embedded image(s) inspected for front/back cover placement.",
        "findings": findings,
        "covers": covers,
    }


def calculate_effective_ppi(width: int, height: int, cx: int, cy: int) -> Optional[Tuple[float, float]]:
    """Calculate effective PPI from DOCX display extents when available."""

    if not cx or not cy:
        return None
    display_width_inches = cx / EMU_PER_INCH
    display_height_inches = cy / EMU_PER_INCH
    if display_width_inches <= 0 or display_height_inches <= 0:
        return None
    return width / display_width_inches, height / display_height_inches


def validate_aspect_ratio(width: int, height: int, cx: int, cy: int) -> Optional[Dict[str, float]]:
    """Compare source and displayed aspect ratios."""

    if not width or not height or not cx or not cy:
        return None
    source_ratio = width / height
    display_ratio = cx / cy
    if source_ratio == 0:
        return None
    difference = abs(source_ratio - display_ratio) / source_ratio
    if difference > 0.05:
        return {
            "source_ratio": round(source_ratio, 4),
            "display_ratio": round(display_ratio, 4),
            "difference_percent": round(difference * 100, 2),
        }
    return None


def save_image_preview(image: Image.Image, path: Path) -> None:
    """Save a small thumbnail preview for a DOCX embedded image."""

    preview = ImageOps.contain(image.convert("RGBA"), (640, 640))
    background = Image.new("RGB", preview.size, "white")
    background.paste(preview, mask=preview.split()[-1])
    background.save(path, "JPEG", quality=88)


def safe_filename(value: str, fallback: str = "section") -> str:
    """Create a portable lowercase filename slug."""

    value = re.sub(r"[^a-zA-Z0-9]+", "-", value or "").strip("-").lower()
    value = re.sub(r"-+", "-", value)
    return value or fallback


def truncate_words(value: str, max_words: int) -> str:
    """Trim text to a maximum word count without changing the original source."""

    words = normalize_spaces(value).split()
    if len(words) <= max_words:
        return " ".join(words)
    return " ".join(words[:max_words]).rstrip(".,;:") + "..."


def delete_paragraph(paragraph: Any) -> None:
    """Remove a paragraph from a python-docx document."""

    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def save_paragraph_range(source_path: Path, start: int, end: int, destination: Path) -> None:
    """Save a DOCX containing only a paragraph range from the source document."""

    document = Document(str(source_path))
    for index in range(len(document.paragraphs) - 1, -1, -1):
        if index < start or index >= end:
            delete_paragraph(document.paragraphs[index])
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))


def save_optimized_cover(blob: bytes, destination: Path, max_size: Tuple[int, int] = (1600, 2400)) -> Optional[Dict[str, Any]]:
    """Save a web-friendly cover image without distorting its aspect ratio."""

    try:
        with Image.open(BytesIO(blob)) as image:
            optimized = ImageOps.contain(image.convert("RGB"), max_size)
            optimized.save(destination, "JPEG", quality=88, optimize=True)
            return {
                "width_px": optimized.size[0],
                "height_px": optimized.size[1],
                "path": str(destination),
            }
    except UnidentifiedImageError:
        return None


def save_cover_assets(images: Sequence[Dict[str, Any]], package_dir: Path) -> Dict[str, Any]:
    """Extract first and last DOCX images as front/back cover assets."""

    covers_dir = package_dir / "covers"
    covers_dir.mkdir(parents=True, exist_ok=True)
    result: Dict[str, Any] = {}

    candidates = []
    if images:
        candidates.append(("front", images[0]))
    if len(images) > 1:
        candidates.append(("back", images[-1]))

    for kind, image in candidates:
        ext = image_extension(image["content_type"])
        original_path = covers_dir / f"{kind}-cover-original{ext}"
        optimized_path = covers_dir / f"{kind}-cover-optimized.jpg"
        original_path.write_bytes(image["blob"])
        optimized = save_optimized_cover(image["blob"], optimized_path)
        result[f"{kind}_cover"] = {
            "paragraph_index": image["paragraph_index"],
            "content_type": image["content_type"],
            "original_path": str(original_path),
            "optimized_path": str(optimized_path) if optimized else "",
            "optimized": optimized,
        }

    return result


def content_bounds_from_covers(document: Document) -> Tuple[int, int, List[Dict[str, Any]]]:
    """Return paragraph bounds for manuscript content between cover images."""

    images = extract_images(document)
    start = 0
    end = len(document.paragraphs)
    if images:
        start = min(len(document.paragraphs), images[0]["paragraph_index"] + 1)
    if len(images) > 1 and images[-1]["paragraph_index"] > start:
        end = images[-1]["paragraph_index"]
    return start, end, images


def skip_upload_heading(text: str) -> bool:
    """Skip source front matter headings that the package regenerates."""

    normalized = normalize_spaces(text).lower()
    return normalized in {
        "index",
        "contents",
        "table of contents",
        "title page",
        "copyright",
        "copyright page",
        "disclaimer",
        "legal notice",
    }


def is_upload_section_heading(text: str, style: str) -> bool:
    """Decide whether a paragraph starts an uploadable reader section."""

    if "heading 1" not in (style or "").lower():
        return False
    normalized = normalize_spaces(text)
    if not normalized or skip_upload_heading(normalized):
        return False
    if re.match(r"^chapter\s+(\d+|[a-z]+)\b", normalized, re.IGNORECASE):
        return True
    return normalized.lower() in {
        "introduction",
        "prologue",
        "preface",
        "foreword",
        "closing note",
        "conclusion",
        "epilogue",
        "endnotes",
        "notes",
        "appendix",
        "bibliography",
    }


def find_upload_section_starts(
    paragraphs: Sequence[Dict[str, Any]],
    content_start: int,
    content_end: int,
) -> List[Tuple[int, str]]:
    """Find sections that should become reader upload DOCX files."""

    starts: List[Tuple[int, str]] = []
    for paragraph in paragraphs:
        index = paragraph["index"]
        if index < content_start or index >= content_end:
            continue
        text = normalize_spaces(paragraph["text"])
        if is_upload_section_heading(text, paragraph["style"]):
            starts.append((index, text))

    if starts:
        return starts

    for paragraph in paragraphs:
        index = paragraph["index"]
        if content_start <= index < content_end and normalize_spaces(paragraph["text"]):
            return [(index, "Manuscript")]
    return []


def section_display_title(start_index: int, heading: str, paragraphs: Sequence[Dict[str, Any]]) -> str:
    """Create a readable title for a packaged reader section."""

    text = normalize_spaces(heading)
    if re.match(r"^chapter\s+\d+\b", text, re.IGNORECASE):
        next_paragraph = next((item for item in paragraphs if item["index"] == start_index + 1), None)
        if next_paragraph and "heading 2" in (next_paragraph["style"] or "").lower():
            subtitle = normalize_spaces(next_paragraph["text"])
            if subtitle:
                return f"{text} - {subtitle}"
    if text.lower() == "introduction":
        next_paragraph = next((item for item in paragraphs if item["index"] == start_index + 1), None)
        if next_paragraph and normalize_spaces(next_paragraph["text"]):
            return f"Introduction - {normalize_spaces(next_paragraph['text'])}"
    if text.lower() == "closing note":
        next_paragraph = next((item for item in paragraphs if item["index"] == start_index + 1), None)
        if next_paragraph and normalize_spaces(next_paragraph["text"]):
            return f"Closing Note - {normalize_spaces(next_paragraph['text'])}"
    return text


def derive_category_slug(full_text: str) -> str:
    """Infer the closest existing site category from keyword signals."""

    lowered = full_text.lower()
    scores = {
        category: sum(1 for keyword in keywords if keyword.lower() in lowered)
        for category, keywords in KNOWN_CATEGORY_KEYWORDS.items()
    }
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "business"


def first_paragraph_matching(
    paragraphs: Sequence[Dict[str, Any]],
    start: int,
    end: int,
    predicate: Any,
) -> Optional[Dict[str, Any]]:
    """Find the first paragraph in a range that matches a predicate."""

    for paragraph in paragraphs:
        if start <= paragraph["index"] < end and normalize_spaces(paragraph["text"]) and predicate(paragraph):
            return paragraph
    return None


def derive_upload_form(
    docx_path: Path,
    document: Document,
    paragraphs: Sequence[Dict[str, Any]],
    content_start: int,
    content_end: int,
    sections: Sequence[Dict[str, Any]],
    publisher: str,
    author_override: Optional[str],
) -> Dict[str, Any]:
    """Derive all fields required by the existing admin book upload form."""

    first_section_index = sections[0]["source_start_paragraph"] if sections else content_end
    front_matter_end = max(content_start, first_section_index)
    title_source = first_paragraph_matching(
        paragraphs,
        content_start,
        front_matter_end,
        lambda item: "heading 1" in (item["style"] or "").lower() and not skip_upload_heading(item["text"]),
    )
    subtitle_source = first_paragraph_matching(
        paragraphs,
        (title_source["index"] + 1 if title_source else content_start),
        front_matter_end,
        lambda item: "heading 2" in (item["style"] or "").lower() or "heading 3" in (item["style"] or "").lower(),
    )
    descriptive_source = first_paragraph_matching(
        paragraphs,
        content_start,
        min(content_end, (sections[1]["source_start_paragraph"] if len(sections) > 1 else content_end)),
        lambda item: 12 <= len(normalize_spaces(item["text"]).split()) <= 90 and "heading" not in (item["style"] or "").lower(),
    )

    title = normalize_spaces(title_source["text"]) if title_source else normalize_spaces(document.core_properties.title or "")
    if not title:
        title = docx_path.stem
    subtitle = normalize_spaces(subtitle_source["text"]) if subtitle_source else ""
    author = author_override or normalize_spaces(document.core_properties.author or "") or publisher
    full_text = "\n".join(item["text"] for item in paragraphs[content_start:content_end])
    word_count = sum(item.get("word_count", 0) for item in sections)
    reading_minutes = max(1, round(word_count / 200)) if word_count else 1

    short_description = truncate_words(descriptive_source["text"], 38) if descriptive_source else f"A reader-ready edition of {title}."
    description_candidates = [
        normalize_spaces(item["text"])
        for item in paragraphs
        if content_start <= item["index"] < content_end
        and "heading" not in (item["style"] or "").lower()
        and len(normalize_spaces(item["text"]).split()) >= 18
    ]
    description = truncate_words(" ".join(description_candidates[:3]), 95) if description_candidates else short_description

    category_slug = derive_category_slug(f"{title}\n{subtitle}\n{full_text[:12000]}")
    is_business = category_slug == "business"
    benefits = [
        "Read a clean structured edition prepared for the digital reader",
        "Move through the book in chronological sections",
        "Use the generated index to navigate the manuscript",
    ]
    who_for = ["Readers who want the complete book in a clear digital format"]
    learnings = ["The core ideas and chapters presented in the manuscript"]
    if is_business:
        benefits = [
            "Clarify the central promise and operating discipline behind the work",
            "Review practical lessons in a chapter-by-chapter reader format",
            "Use structured sections for easier preview and author review",
        ]
        who_for = [
            "Founders and operators",
            "Creative professionals and small-business owners",
            "Readers interested in practical enterprise building",
        ]
        learnings = [
            "How the manuscript develops its main argument across chapters",
            "How practical examples connect to operating decisions",
            "How to review each section before publication",
        ]

    return {
        "title": title,
        "subtitle": subtitle,
        "author": author,
        "category_slug": category_slug,
        "short_description": short_description,
        "description": description,
        "cover_image_url": "",
        "back_cover_image_url": "",
        "estimated_reading_time": f"{reading_minutes} min",
        "price_paperback": "",
        "price_ebook": "",
        "buy_url": "",
        "formats": ["Ebook"],
        "benefits": benefits,
        "who_for": who_for,
        "learnings": learnings,
        "about_author": f"Published by {publisher}.",
        "is_published": False,
    }


def write_metadata_import_docx(form: Dict[str, Any], destination: Path) -> None:
    """Write a DOCX metadata template accepted by the existing admin importer."""

    document = Document()
    document.add_heading("Earnalism Book Import Template", level=1)
    labels = [
        ("Title", "title"),
        ("Subtitle", "subtitle"),
        ("Author", "author"),
        ("Category", "category_slug"),
        ("Estimated Reading Time", "estimated_reading_time"),
        ("Formats", "formats"),
        ("Short Description", "short_description"),
        ("Description", "description"),
        ("Benefits", "benefits"),
        ("Who This Is For", "who_for"),
        ("What You Will Learn", "learnings"),
        ("About Author", "about_author"),
        ("Buy URL", "buy_url"),
        ("Paperback Price", "price_paperback"),
        ("Ebook Price", "price_ebook"),
    ]
    for label, key in labels:
        value = form.get(key, "")
        if isinstance(value, list):
            document.add_paragraph(f"{label}:")
            for item in value:
                document.add_paragraph(f"- {item}")
        else:
            document.add_paragraph(f"{label}: {value}")
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))


def write_generated_index_docx(
    form: Dict[str, Any],
    sections: Sequence[Dict[str, Any]],
    destination: Path,
    pagination: Optional[Dict[str, Any]] = None,
) -> None:
    """Write an index page with estimated reader page numbers."""

    document = Document()
    document.add_heading("Index", level=1)
    document.add_paragraph(form["title"])
    if form.get("subtitle"):
        document.add_paragraph(form["subtitle"])
    if pagination and pagination.get("status") == "ready":
        viewport = pagination.get("viewport", {})
        document.add_paragraph(
            "Page numbers are calculated from rendered Earnalism reader screens "
            f"at {viewport.get('width')}x{viewport.get('height')} using the default reader typography."
        )
    else:
        document.add_paragraph("Reader-screen page numbers could not be calculated. See manifest.json for pagination errors.")
    for section in sections:
        page = section.get("reader_start_page")
        page_label = str(page) if page else "unavailable"
        document.add_paragraph(f"{section['title']} ...... {page_label}")
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))


def write_preview_markdown(
    form: Dict[str, Any],
    sections: Sequence[Dict[str, Any]],
    covers: Dict[str, Any],
    pagination: Dict[str, Any],
    destination: Path,
) -> None:
    """Write a simple human-readable upload preview."""

    lines = [
        f"# {form['title']}",
        "",
        form.get("subtitle", ""),
        "",
        f"Author/Publisher: {form['author']}",
        f"Category: {form['category_slug']}",
        f"Estimated reading time: {form['estimated_reading_time']}",
        f"Draft status: {not form['is_published']}",
        "",
        "## Covers",
        "",
        f"- Front cover: {covers.get('front_cover', {}).get('optimized_path') or covers.get('front_cover', {}).get('original_path') or 'missing'}",
        f"- Back cover: {covers.get('back_cover', {}).get('optimized_path') or covers.get('back_cover', {}).get('original_path') or 'missing'}",
        "",
        "## Generated Index",
        "",
    ]
    if pagination.get("status") == "ready":
        viewport = pagination.get("viewport", {})
        lines.append(
            f"Reader pages measured at {viewport.get('width')}x{viewport.get('height')} "
            f"with font size {pagination.get('font_size')}."
        )
        lines.append("")
    else:
        lines.append(f"Reader pagination unavailable: {pagination.get('error', 'unknown error')}")
        lines.append("")
    for section in sections:
        page = section.get("reader_start_page")
        page_label = page if page else "unavailable"
        lines.append(f"- Page {page_label}: {section['title']}")
    destination.write_text("\n".join(lines), encoding="utf-8")


def section_filename(order: int, title: str) -> str:
    """Create a stable section upload filename."""

    return f"{order:02d}-{safe_filename(title)}.docx"


def find_chrome_executable() -> Optional[str]:
    """Find a local Chrome/Chromium executable for Playwright."""

    env_path = os.getenv("CHROME_EXECUTABLE_PATH") or os.getenv("PLAYWRIGHT_CHROME_EXECUTABLE")
    candidates = [
        env_path,
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/Applications/Chromium.app/Contents/MacOS/Chromium",
        "/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge",
        "/usr/bin/google-chrome",
        "/usr/bin/chromium",
        "/usr/bin/chromium-browser",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return candidate
    return None


def fallback_docx_to_reader_html(path: Path) -> str:
    """Convert a DOCX section to simple HTML when mammoth is unavailable."""

    document = Document(str(path))
    parts: List[str] = []
    for paragraph in document.paragraphs:
        text = normalize_spaces(paragraph.text)
        if not text:
            continue
        escaped = html_lib.escape(text)
        style = paragraph.style.name if paragraph.style else ""
        if "Heading 1" in style:
            parts.append(f"<h2>{escaped}</h2>")
        elif "Heading 2" in style:
            parts.append(f"<h3>{escaped}</h3>")
        elif "Heading 3" in style:
            parts.append(f"<h4>{escaped}</h4>")
        else:
            parts.append(f"<p>{escaped}</p>")
    return "".join(parts)


def docx_to_reader_html(path: Path) -> str:
    """Convert a section DOCX to HTML similar to the backend upload pipeline."""

    try:
        import mammoth  # type: ignore

        with path.open("rb") as handle:
            result = mammoth.convert_to_html(handle, style_map=READER_STYLE_MAP)
        html = result.value or ""
        if html.strip():
            return html
    except Exception:
        pass
    return fallback_docx_to_reader_html(path)


def reader_harness_html(title: str, body_html: str, font_size: str, is_bengali: bool) -> str:
    """Build an HTML harness mirroring the Earnalism reader screen."""

    font_family = "'Noto Serif Bengali', 'Crimson Pro', Georgia, serif" if is_bengali else "'Crimson Pro', 'Noto Serif Bengali', Georgia, serif"
    title_font = "'Noto Serif Bengali', 'Crimson Pro', Georgia, serif" if is_bengali else "'Cormorant Garamond', 'Noto Serif Bengali', serif"
    line_height = "1.9" if is_bengali else "1.75"
    title_line_height = "1.55" if is_bengali else "1.4"
    content_class = "reader-content reader-content--bengali" if is_bengali else "drop-cap reader-content"

    return f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=Crimson+Pro:ital,wght@0,300;0,400;0,500;0,600;1,300;1,400&family=Cormorant+Garamond:ital,wght@0,400;0,500;0,600;1,400;1,500&family=Inter:wght@300;400;500;600&family=Noto+Serif+Bengali:wght@400;500;600&family=Noto+Sans+Bengali:wght@400;500&display=swap" rel="stylesheet">
  <style>
    :root {{
      --beige-canvas: #FAF7F0;
      --beige-surface: #F5F0E8;
      --burgundy-deep: #6B1020;
      --text-primary: #1C0A0E;
      --text-secondary: #7A5C62;
      --border-soft: #E8DDD8;
    }}
    * {{ box-sizing: border-box; }}
    html, body {{ margin: 0; padding: 0; background: var(--beige-canvas); color: var(--text-primary); }}
    .reader-scroll {{ position: relative; height: 100vh; min-height: 100vh; overflow-y: auto; background: var(--beige-canvas); }}
    .reader-header {{ position: fixed; top: 2px; left: 0; right: 0; z-index: 40; height: 57px; background: rgba(250,247,240,0.93); border-bottom: 1px solid var(--border-soft); }}
    .reader-footer {{ position: fixed; bottom: 0; left: 0; right: 0; z-index: 40; height: 78px; background: rgba(250,247,240,0.96); border-top: 1px solid var(--border-soft); }}
    main {{ min-height: 100vh; padding: 80px 20px 144px; }}
    .reader-canvas {{
      font-family: 'Crimson Pro', 'Noto Serif Bengali', Georgia, serif;
      font-size: 17px;
      font-weight: 400;
      line-height: 1.75;
      text-align: justify;
      hyphens: auto;
      overflow-wrap: break-word;
      word-break: normal;
      line-break: auto;
      max-width: 58ch;
      margin: 0 auto;
      user-select: none;
    }}
    @media (max-width: 480px) {{
      .reader-canvas {{ text-align: left; hyphens: none; }}
    }}
    .chapter-title {{
      font-family: {title_font};
      font-size: 28px;
      font-weight: 500;
      text-align: center;
      color: var(--burgundy-deep);
      letter-spacing: 0;
      line-height: {title_line_height};
      margin: 0 0 24px;
      overflow-wrap: break-word;
    }}
    .chapter-divider {{ display: flex; align-items: center; gap: 12px; margin-bottom: 40px; justify-content: center; }}
    .chapter-divider span:first-child, .chapter-divider span:last-child {{ flex: 1; height: 1px; background: var(--border-soft); }}
    .chapter-divider .mark {{ color: var(--burgundy-deep); font-size: 20px; flex: 0; }}
    .reader-content {{
      max-width: 100%;
      overflow-wrap: break-word;
      word-break: normal;
      font-family: {font_family};
      font-size: {font_size};
      line-height: {line_height};
      color: var(--text-primary);
    }}
    .reader-content p {{ margin: 0 0 1.6em; }}
    .reader-content h2 {{
      font-family: 'Cormorant Garamond', 'Noto Serif Bengali', serif;
      font-size: clamp(20px, 4vw, 26px);
      font-weight: 500;
      color: var(--burgundy-deep);
      margin: 2em 0 0.8em;
      line-height: 1.3;
    }}
    .reader-content h3 {{
      font-family: 'Cormorant Garamond', 'Noto Serif Bengali', serif;
      font-size: clamp(17px, 3vw, 21px);
      font-weight: 500;
      margin: 1.6em 0 0.6em;
    }}
    .reader-content blockquote {{ border-left: 3px solid #8B1A2A; padding: 0.5em 0 0.5em 1.5em; font-style: italic; color: var(--text-secondary); }}
    .reader-content ul, .reader-content ol {{ padding-left: 1.5em; margin-bottom: 1.4em; }}
    .reader-content li {{ margin-bottom: 0.5em; }}
    .reader-content table {{ width: 100%; border-collapse: collapse; font-size: 15px; overflow-x: auto; display: block; margin: 2em 0; }}
    .reader-content img {{ max-width: 100%; height: auto; object-fit: contain; pointer-events: none; display: block; margin: 2.5em auto; border-radius: 8px; }}
    .drop-cap::first-letter {{
      font-family: 'Cormorant Garamond', 'Noto Serif Bengali', serif;
      font-size: 4.5em;
      font-weight: 600;
      line-height: 0.8;
      float: left;
      margin: 0.05em 0.1em 0 0;
      color: var(--burgundy-deep);
    }}
    .reader-content--bengali {{ font-family: 'Noto Serif Bengali', 'Noto Sans Bengali', 'Crimson Pro', Georgia, serif; line-height: 1.9; text-align: left; }}
  </style>
</head>
<body>
  <div class="reader-scroll" id="reader-scroll">
    <div class="reader-header" aria-hidden="true"></div>
    <main>
      <div class="reader-canvas">
        <h2 class="chapter-title">{html_lib.escape(title)}</h2>
        <div class="chapter-divider"><span></span><span class="mark">❧</span><span></span></div>
        <div id="reader-content" class="{content_class}">{body_html}</div>
      </div>
    </main>
    <div class="reader-footer" aria-hidden="true"></div>
  </div>
</body>
</html>"""


def generated_index_body_html(form: Dict[str, Any], sections: Sequence[Dict[str, Any]]) -> str:
    """Build the generated index content as reader HTML for page measurement."""

    items = []
    for section in sections:
        page = section.get("reader_start_page") or 1
        items.append(f"<li>{html_lib.escape(section['title'])} ...... {page}</li>")
    subtitle = f"<p><em>{html_lib.escape(form.get('subtitle', ''))}</em></p>" if form.get("subtitle") else ""
    return (
        f"<h2>{html_lib.escape(form['title'])}</h2>"
        f"{subtitle}"
        "<p>Page numbers are calculated from rendered Earnalism reader screens.</p>"
        f"<ol>{''.join(items)}</ol>"
    )


def measure_reader_html(
    page: Any,
    title: str,
    body_html: str,
    font_size: str,
    is_bengali: bool,
) -> Dict[str, int]:
    """Measure one reader-rendered HTML body in the Playwright page."""

    page.set_content(reader_harness_html(title, body_html, font_size, is_bengali), wait_until="networkidle")
    page.evaluate("() => document.fonts && document.fonts.ready ? document.fonts.ready : Promise.resolve()")
    return page.evaluate(
        """() => {
          const scroller = document.getElementById('reader-scroll');
          const main = document.querySelector('main');
          const canvas = document.querySelector('.reader-canvas');
          return {
            scrollHeight: Math.ceil(scroller.scrollHeight),
            clientHeight: Math.ceil(scroller.clientHeight),
            mainHeight: Math.ceil(main.getBoundingClientRect().height),
            canvasHeight: Math.ceil(canvas.getBoundingClientRect().height)
          };
        }"""
    )


def measure_reader_render_pages(
    sections: Sequence[Dict[str, Any]],
    form: Dict[str, Any],
    viewport: Tuple[int, int],
    font_size_index: int,
) -> Dict[str, Any]:
    """Measure chapter page counts using a browser-rendered Earnalism reader harness."""

    try:
        from playwright.sync_api import sync_playwright  # type: ignore
    except Exception as exc:
        return {
            "status": "unavailable",
            "error": f"Playwright is not installed: {exc}",
            "install": "python3 -m pip install playwright",
        }

    executable = find_chrome_executable()
    font_size = READER_FONT_SIZES[min(max(font_size_index, 0), len(READER_FONT_SIZES) - 1)]
    width, height = viewport
    measurements: List[Dict[str, Any]] = []

    try:
        with sync_playwright() as playwright:
            launch_options: Dict[str, Any] = {"headless": True}
            if executable:
                launch_options["executable_path"] = executable
            browser = playwright.chromium.launch(**launch_options)
            page = browser.new_page(viewport={"width": width, "height": height}, device_scale_factor=1)
            for section in sections:
                html = docx_to_reader_html(Path(section["file"]))
                is_bengali = bool(re.search(r"[\u0980-\u09FF]", f"{section['title']} {html}"))
                metrics = measure_reader_html(page, section["title"], html, font_size, is_bengali)
                rendered_pages = max(1, int((metrics["scrollHeight"] + metrics["clientHeight"] - 1) // metrics["clientHeight"]))
                measurements.append(
                    {
                        "order": section["order"],
                        "title": section["title"],
                        "file": section["file"],
                        "rendered_screen_count": rendered_pages,
                        "scroll_height_px": metrics["scrollHeight"],
                        "client_height_px": metrics["clientHeight"],
                        "main_height_px": metrics["mainHeight"],
                        "canvas_height_px": metrics["canvasHeight"],
                    }
                )

            current_page = 1
            for measurement in measurements:
                measurement["reader_start_page"] = current_page
                current_page += measurement["rendered_screen_count"]

            index_metrics = measure_reader_html(
                page,
                "Index",
                generated_index_body_html(form, measurements),
                font_size,
                bool(re.search(r"[\u0980-\u09FF]", f"{form.get('title', '')} {form.get('subtitle', '')}")),
            )
            index_screen_count = max(
                1,
                int((index_metrics["scrollHeight"] + index_metrics["clientHeight"] - 1) // index_metrics["clientHeight"]),
            )
            browser.close()
    except Exception as exc:
        return {
            "status": "unavailable",
            "error": f"Reader rendering failed: {exc}",
            "chrome_executable": executable,
        }

    current_page = 1 + index_screen_count
    for measurement in measurements:
        measurement["reader_start_page"] = current_page
        current_page += measurement["rendered_screen_count"]

    return {
        "status": "ready",
        "method": "browser_rendered_earnalism_reader_screen",
        "viewport": {"width": width, "height": height},
        "font_size": font_size,
        "font_size_index": font_size_index,
        "chrome_executable": executable or "playwright-managed chromium",
        "generated_index": {
            "reader_start_page": 1,
            "rendered_screen_count": index_screen_count,
            "scroll_height_px": index_metrics["scrollHeight"],
            "client_height_px": index_metrics["clientHeight"],
            "main_height_px": index_metrics["mainHeight"],
            "canvas_height_px": index_metrics["canvasHeight"],
        },
        "measurements": measurements,
        "total_reader_screens": max(0, current_page - 1),
    }


def apply_reader_pagination(sections: List[Dict[str, Any]], pagination: Dict[str, Any]) -> None:
    """Apply rendered reader pagination data to section manifest entries."""

    by_order = {item["order"]: item for item in pagination.get("measurements", [])}
    for section in sections:
        measurement = by_order.get(section["order"])
        if not measurement:
            section["reader_start_page"] = None
            section["rendered_screen_count"] = None
            continue
        section["reader_start_page"] = measurement["reader_start_page"]
        section["rendered_screen_count"] = measurement["rendered_screen_count"]
        section["reader_scroll_height_px"] = measurement["scroll_height_px"]
        section["reader_client_height_px"] = measurement["client_height_px"]


def build_upload_sections(
    source_path: Path,
    paragraphs: Sequence[Dict[str, Any]],
    content_start: int,
    content_end: int,
    package_dir: Path,
) -> List[Dict[str, Any]]:
    """Split manuscript content into safe, uploadable DOCX sections."""

    chapters_dir = package_dir / "chapters"
    starts = find_upload_section_starts(paragraphs, content_start, content_end)
    sections: List[Dict[str, Any]] = []

    for position, (start, heading) in enumerate(starts, start=1):
        end = starts[position][0] if position < len(starts) else content_end
        title = section_display_title(start, heading, paragraphs)
        word_count = sum(count_words(item["text"]) for item in paragraphs if start <= item["index"] < end)
        filename = section_filename(position, title)
        destination = chapters_dir / filename
        save_paragraph_range(source_path, start, end, destination)
        sections.append(
            {
                "order": position,
                "title": title,
                "source_heading": heading,
                "source_start_paragraph": start,
                "source_end_paragraph_exclusive": end,
                "word_count": word_count,
                "estimated_reading_minutes": max(1, round(word_count / 200)) if word_count else 0,
                "reader_start_page": None,
                "rendered_screen_count": None,
                "file": str(destination),
            }
        )

    return sections


def prepare_upload_package(
    docx_path: Path,
    document: Document,
    paragraphs: Sequence[Dict[str, Any]],
    output_dir: Path,
    publisher: str,
    author_override: Optional[str],
    reader_viewport: Tuple[int, int],
    reader_font_size_index: int,
) -> Dict[str, Any]:
    """Create a complete local upload package from a single full-book DOCX."""

    package_dir = output_dir / "upload_package"
    covers_dir = package_dir / "covers"
    chapters_dir = package_dir / "chapters"
    for directory in [package_dir, covers_dir, chapters_dir]:
        directory.mkdir(parents=True, exist_ok=True)

    content_start, content_end, images = content_bounds_from_covers(document)
    covers = save_cover_assets(images, package_dir)
    sections = build_upload_sections(docx_path, paragraphs, content_start, content_end, package_dir)
    form = derive_upload_form(docx_path, document, paragraphs, content_start, content_end, sections, publisher, author_override)
    pagination = measure_reader_render_pages(sections, form, reader_viewport, reader_font_size_index)
    apply_reader_pagination(sections, pagination)

    index_path = package_dir / "generated-index.docx"
    metadata_path = package_dir / "metadata-import-template.docx"
    form_json_path = package_dir / "book-upload-form.json"
    manifest_path = package_dir / "manifest.json"
    preview_path = package_dir / "reader-preview.md"
    source_without_covers_path = package_dir / "reader-manuscript-without-covers.docx"

    write_generated_index_docx(form, sections, index_path, pagination)
    write_metadata_import_docx(form, metadata_path)
    form_json_path.write_text(json.dumps(form, indent=2, ensure_ascii=False), encoding="utf-8")
    write_preview_markdown(form, sections, covers, pagination, preview_path)
    if content_start < content_end:
        save_paragraph_range(docx_path, content_start, content_end, source_without_covers_path)

    upload_order = [
        {
            "order": 0,
            "title": "Generated Index",
            "file": str(index_path),
            "reader_start_page": 1,
        }
    ] + list(sections)

    blockers = []
    if not covers.get("front_cover"):
        blockers.append("Front cover image was not detected on the first page/image of the DOCX.")
    if not covers.get("back_cover"):
        blockers.append("Back cover image was not detected as the last page/image of the DOCX.")
    if pagination.get("status") != "ready":
        blockers.append("Reader-screen pagination could not be calculated; generated index page numbers are unavailable.")

    manifest = {
        "source_file": str(docx_path),
        "package_dir": str(package_dir),
        "book_upload_form_fields": BOOK_UPLOAD_FORM_FIELDS,
        "book_upload_form_json": str(form_json_path),
        "metadata_import_template": str(metadata_path),
        "generated_index": str(index_path),
        "reader_preview": str(preview_path),
        "reader_manuscript_without_covers": str(source_without_covers_path) if source_without_covers_path.exists() else "",
        "covers": covers,
        "reader_pagination": pagination,
        "content_bounds": {
            "start_paragraph_index": content_start,
            "end_paragraph_index_exclusive": content_end,
        },
        "upload_order": upload_order,
        "book": form,
        "draft_publish_state": "draft",
        "notes": [
            "The original DOCX was not modified.",
            "Page numbers in generated-index.docx are measured against the Earnalism reader screen, not DOCX word count.",
            "Reader page numbers vary if the user changes font size, viewport size, or reader theme/device.",
            "Upload the extracted optimized cover images through the admin import form when available.",
        ],
        "publish_blockers": blockers,
    }
    manifest_path.write_text(json.dumps(manifest, indent=2, ensure_ascii=False), encoding="utf-8")

    steps = [
        f"# {form['title']} Upload Package",
        "",
        "## Admin Upload Steps",
        "",
        "1. Open Admin Dashboard > Books > New book.",
        f"2. Upload `{metadata_path.name}` in the DOCX import section.",
        "3. Upload the extracted front/back cover images from `covers/` when present.",
        "4. Click Import and prefill.",
        "5. Review every field and save as Draft. The package sets `is_published` to false.",
        "6. Reopen the saved book and upload the DOCX files below in order.",
        "7. Preview the reader rendering before publishing.",
        "",
        "## Upload Order",
        "",
        f"00. `{index_path.name}` - Generated Index",
    ]
    for section in sections:
        steps.append(
            f"{section['order']:02d}. `{Path(section['file']).name}` - {section['title']} "
            f"({section['word_count']} words, reader page {section.get('reader_start_page') or 'unavailable'})"
        )
    if blockers:
        steps.extend(["", "## Publish Blockers", ""])
        steps.extend([f"- {blocker}" for blocker in blockers])
    (package_dir / "UPLOAD_STEPS.md").write_text("\n".join(steps), encoding="utf-8")

    return manifest


def build_report(
    docx_path: Path,
    output_dir: Path,
    chapter_prefix: str = DEFAULT_CHAPTER_PREFIX,
    min_chapter_words: int = 50,
    prepare_package: bool = True,
    publisher: str = DEFAULT_PUBLISHER,
    author_override: Optional[str] = None,
    reader_viewport: Tuple[int, int] = (390, 844),
    reader_font_size_index: int = DEFAULT_READER_FONT_SIZE_INDEX,
) -> Dict[str, Any]:
    """Run all checks and return the structured report."""

    document = load_document(docx_path)
    paragraphs = extract_paragraphs(document)
    full_text = "\n".join(paragraph["text"] for paragraph in paragraphs)
    previews_dir = output_dir / "previews"

    checks = {
        "ai_authorship_detection": check_ai_authorship(paragraphs),
        "classified_sensitive_information": check_sensitive_information(paragraphs),
        "defamation_risk": check_defamation_risk(paragraphs),
        "plagiarism_surface_check": check_plagiarism_surface(full_text),
        "full_book_chapter_structure": check_full_book_chapter_structure(paragraphs, chapter_prefix, min_chapter_words),
        "publishing_compliance": check_publishing_compliance(paragraphs),
        "cover_page_structure": check_cover_page_structure(document),
        "inline_image_rendering_validation": check_inline_images(document, previews_dir),
    }

    overall_status = "FAIL" if any(check["status"] == "FAIL" for check in checks.values()) else "PASS"
    upload_package = (
        prepare_upload_package(
            docx_path,
            document,
            paragraphs,
            output_dir,
            publisher,
            author_override,
            reader_viewport,
            reader_font_size_index,
        )
        if prepare_package
        else None
    )
    return {
        "source_file": str(docx_path),
        "generated_at": dt.datetime.now(dt.timezone.utc).isoformat(),
        "overall_status": overall_status,
        "checks": checks,
        "upload_package": upload_package,
        "summary_recommendation": build_summary_recommendation(overall_status, checks),
    }


def build_summary_recommendation(status: str, checks: Dict[str, Dict[str, Any]]) -> str:
    """Create a short recommendation from check statuses."""

    failed = [name for name, result in checks.items() if result["status"] == "FAIL"]
    if status == "PASS":
        return "No blocking issues were detected by the automated review. Complete a human editorial, legal, and layout review before publication."
    return (
        "Do not publish until the failed checks are reviewed and resolved: "
        + ", ".join(failed)
        + ". Automated results are heuristic and should be confirmed by a qualified reviewer."
    )


def write_json_report(report: Dict[str, Any], path: Path) -> None:
    """Write the structured JSON report."""

    path.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")


def write_markdown_summary(report: Dict[str, Any], path: Path) -> None:
    """Write a human-readable Markdown summary."""

    lines = [
        "# DOCX Review Report",
        "",
        f"**Source:** `{report['source_file']}`",
        f"**Generated:** {report['generated_at']}",
        f"**Overall status:** **{report['overall_status']}**",
        "",
        "## Summary Recommendation",
        "",
        report["summary_recommendation"],
        "",
        "## Checks",
        "",
    ]

    for name, result in report["checks"].items():
        lines.extend(
            [
                f"### {name.replace('_', ' ').title()}",
                "",
                f"- Status: **{result['status']}**",
                f"- Summary: {result['summary']}",
                "",
            ]
        )
        if result.get("chapters"):
            lines.append("| Chapter | Heading | Paragraph | Words | Reading Min | Style |")
            lines.append("|---:|---|---:|---:|---:|---|")
            for chapter in result["chapters"]:
                title = chapter["heading"].replace("|", "\\|")
                style = (chapter.get("style") or "unknown").replace("|", "\\|")
                lines.append(
                    f"| {chapter['number']} | {title} | {chapter['heading_paragraph_index']} | "
                    f"{chapter['word_count']} | {chapter['estimated_reading_minutes']} | {style} |"
                )
            lines.append("")
        findings = result.get("findings", [])
        if not findings:
            lines.append("No findings.")
            lines.append("")
            continue
        for finding in findings:
            location = f" paragraph {finding['paragraph_index']}" if "paragraph_index" in finding else ""
            lines.append(f"- **{finding['severity']}**{location}: {finding['message']}")
            if finding.get("excerpt"):
                lines.append(f"  - Excerpt: {finding['excerpt']}")
            if finding.get("metadata"):
                lines.append(f"  - Metadata: `{json.dumps(finding['metadata'], ensure_ascii=False)[:500]}`")
        lines.append("")

    upload_package = report.get("upload_package")
    if upload_package:
        lines.extend(
            [
                "## Upload Package",
                "",
                f"- Package directory: `{upload_package['package_dir']}`",
                f"- Metadata import template: `{upload_package['metadata_import_template']}`",
                f"- Book upload form JSON: `{upload_package['book_upload_form_json']}`",
                f"- Generated index: `{upload_package['generated_index']}`",
                f"- Draft publish state: `{upload_package['draft_publish_state']}`",
                "",
                "### Reader Pagination",
                "",
            ]
        )
        pagination = upload_package.get("reader_pagination") or {}
        if pagination.get("status") == "ready":
            viewport = pagination.get("viewport", {})
            lines.extend(
                [
                    f"- Method: `{pagination.get('method')}`",
                    f"- Viewport: `{viewport.get('width')}x{viewport.get('height')}`",
                    f"- Font size: `{pagination.get('font_size')}`",
                    f"- Total reader screens: `{pagination.get('total_reader_screens')}`",
                    "",
                ]
            )
        else:
            lines.extend([f"- Status: `{pagination.get('status', 'unavailable')}`", f"- Error: {pagination.get('error', 'unknown')}", ""])
        lines.extend(
            [
                "### Upload Order",
                "",
            ]
        )
        for item in upload_package.get("upload_order", []):
            page = item.get("reader_start_page") or "unavailable"
            lines.append(f"- {item['order']:02d}. Page {page}: {item['title']} - `{item['file']}`")
        blockers = upload_package.get("publish_blockers") or []
        if blockers:
            lines.extend(["", "### Publish Blockers", ""])
            lines.extend([f"- {blocker}" for blocker in blockers])
        lines.append("")

    path.write_text("\n".join(lines), encoding="utf-8")


def parse_args(argv: Optional[Sequence[str]] = None) -> argparse.Namespace:
    """Parse CLI arguments."""

    parser = argparse.ArgumentParser(description="Validate and review a DOCX manuscript.")
    parser.add_argument("docx_path", help="Path to the .DOCX file to review.")
    parser.add_argument(
        "--output-dir",
        default="output",
        help="Directory for report.json, report.md, and previews/ (default: output).",
    )
    parser.add_argument(
        "--chapter-prefix",
        default=DEFAULT_CHAPTER_PREFIX,
        help="Expected chapter heading prefix for full-book manuscripts (default: Chapter).",
    )
    parser.add_argument(
        "--min-chapter-words",
        type=int,
        default=50,
        help="Warn when a detected chapter has fewer words than this threshold (default: 50).",
    )
    parser.add_argument(
        "--skip-upload-package",
        action="store_true",
        help="Only run validation reports; do not create upload_package/ assets.",
    )
    parser.add_argument(
        "--publisher",
        default=DEFAULT_PUBLISHER,
        help=f"Publisher name used in generated metadata (default: {DEFAULT_PUBLISHER}).",
    )
    parser.add_argument(
        "--author",
        default=None,
        help="Override author value in the generated book upload form.",
    )
    parser.add_argument(
        "--reader-viewport",
        default=DEFAULT_READER_VIEWPORT,
        help=f"Viewport used for rendered reader pagination, WIDTHxHEIGHT (default: {DEFAULT_READER_VIEWPORT}).",
    )
    parser.add_argument(
        "--reader-font-size-index",
        type=int,
        default=DEFAULT_READER_FONT_SIZE_INDEX,
        choices=range(len(READER_FONT_SIZES)),
        metavar=f"0-{len(READER_FONT_SIZES) - 1}",
        help="Earnalism reader font size index: 0=XS, 1=S, 2=M, 3=L (default: 1).",
    )
    return parser.parse_args(argv)


def main(argv: Optional[Sequence[str]] = None) -> int:
    """CLI entry point."""

    args = parse_args(argv)
    docx_path = Path(args.docx_path).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve()
    reader_viewport = parse_viewport(args.reader_viewport)
    output_dir.mkdir(parents=True, exist_ok=True)

    try:
        report = build_report(
            docx_path,
            output_dir,
            args.chapter_prefix,
            args.min_chapter_words,
            prepare_package=not args.skip_upload_package,
            publisher=args.publisher,
            author_override=args.author,
            reader_viewport=reader_viewport,
            reader_font_size_index=args.reader_font_size_index,
        )
        write_json_report(report, output_dir / "report.json")
        write_markdown_summary(report, output_dir / "report.md")
    except Exception as exc:
        error_report = {
            "source_file": str(docx_path),
            "generated_at": dt.datetime.now(dt.timezone.utc).isoformat(),
            "overall_status": "FAIL",
            "checks": {},
            "summary_recommendation": "The review could not be completed.",
            "error": str(exc),
        }
        write_json_report(error_report, output_dir / "report.json")
        write_markdown_summary(error_report, output_dir / "report.md")
        print(f"Review failed: {exc}", file=sys.stderr)
        return 1

    print(f"Wrote JSON report: {output_dir / 'report.json'}")
    print(f"Wrote Markdown summary: {output_dir / 'report.md'}")
    return 0 if report["overall_status"] == "PASS" else 2


if __name__ == "__main__":
    raise SystemExit(main())
